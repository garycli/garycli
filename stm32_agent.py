#!/usr/bin/env python3
"""
Gary Dev Agent
===============
融合 Gary（编译/烧录/调试闭环）与 ClaudeTerminal（对话 UI + 工具框架）的 STM32 专属 AI 助手。

硬件后端：pyocd（比 OpenOCD 更易用，支持 ST-Link / CMSIS-DAP / J-Link USB 探针）
AI 前端  ：流式对话 + 函数调用工具链

功能：
  - 自然语言 → 生成完整 STM32 HAL C 代码
  - 一键编译（arm-none-eabi-gcc）
  - 一键烧录（pyocd）
  - 读寄存器 / 串口监控 / HardFault 分析
  - 对话式修改历史项目代码
  - 完整自动调试闭环

用法：
  python3 stm32_agent.py               # 启动（不连接硬件）
  python3 stm32_agent.py --connect     # 启动并自动连接第一个可用探针
  python3 stm32_agent.py --chip STM32F407VET6   # 指定芯片
"""

import sys, os, json, re, time, shutil, subprocess, threading, shlex
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Any
from stm32_extra_tools import EXTRA_TOOLS_MAP, EXTRA_TOOL_SCHEMAS
from gary_skills import (
    init_skills, handle_skill_command,
    SKILL_TOOLS_MAP, SKILL_TOOL_SCHEMAS,
    _get_manager,
)
# ─────────────────────────────────────────────────────────────
# 将本文件所在目录加入路径，使能 import compiler / config
# ─────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent.resolve()
sys.path.insert(0, str(_HERE))

# TUI 依赖
from rich.console import Console
from rich.panel import Panel
from rich.text import Text
from rich.table import Table
from rich.markdown import Markdown
from rich.rule import Rule
from rich import box

from prompt_toolkit import PromptSession
from prompt_toolkit.formatted_text import HTML
from prompt_toolkit.styles import Style
from prompt_toolkit.history import InMemoryHistory

from openai import OpenAI

# Flash 项目内部模块
from config import (
    AI_API_KEY, AI_BASE_URL, AI_MODEL, AI_TEMPERATURE,
    WORKSPACE, BUILD_DIR, PROJECTS_DIR,
    DEFAULT_CHIP, DEFAULT_CLOCK,
    SERIAL_PORT, SERIAL_BAUD,
    POST_FLASH_DELAY, REGISTER_READ_DELAY,
)
from compiler import Compiler

# ─────────────────────────────────────────────────────────────
# AI 接口管理（动态读写 config.py，支持运行时切换）
# ─────────────────────────────────────────────────────────────
_AI_PRESETS = [
    # (显示名,                base_url,                                                    默认 model)
    ("OpenAI",               "https://api.openai.com/v1",                                 "gpt-4o"),
    ("DeepSeek",             "https://api.deepseek.com/v1",                               "deepseek-chat"),
    ("Kimi / Moonshot",      "https://api.moonshot.cn/v1",                                "kimi-k2.5"),
    ("Google Gemini",        "https://generativelanguage.googleapis.com/v1beta/openai/",  "gemini-2.0-flash"),
    ("通义千问 (阿里云)",     "https://dashscope.aliyuncs.com/compatible-mode/v1",          "qwen-plus"),
    ("智谱 GLM",             "https://open.bigmodel.cn/api/paas/v4/",                     "glm-4-flash"),
    ("Ollama (本地)",         "http://127.0.0.1:11434/v1",                                 "qwen2.5-coder:14b"),
    ("自定义 / Other",        "",                                                           ""),
]

def _read_ai_config() -> tuple:
    """从 config.py 读取 (api_key, base_url, model)"""
    p = _HERE / "config.py"
    if not p.exists():
        return "", "", ""
    text = p.read_text(encoding="utf-8")
    def _get(pat):
        m = re.search(pat, text, re.MULTILINE)
        return m.group(1).strip() if m else ""
    return (
        _get(r'^AI_API_KEY\s*=\s*["\']([^"\']*)["\']'),
        _get(r'^AI_BASE_URL\s*=\s*["\']([^"\']*)["\']'),
        _get(r'^AI_MODEL\s*=\s*["\']([^"\']*)["\']'),
    )

def _write_ai_config(api_key: str, base_url: str, model: str) -> bool:
    """原地修改 config.py 的三行 AI 配置"""
    p = _HERE / "config.py"
    if not p.exists():
        return False
    text = p.read_text(encoding="utf-8")
    text = re.sub(r'^(AI_API_KEY\s*=\s*).*$',  f'AI_API_KEY = "{api_key}"',  text, flags=re.MULTILINE)
    text = re.sub(r'^(AI_BASE_URL\s*=\s*).*$', f'AI_BASE_URL = "{base_url}"', text, flags=re.MULTILINE)
    text = re.sub(r'^(AI_MODEL\s*=\s*).*$',    f'AI_MODEL = "{model}"',       text, flags=re.MULTILINE)
    p.write_text(text, encoding="utf-8")
    return True

def _reload_ai_globals():
    """写入 config.py 后重新加载 AI 相关全局变量"""
    import importlib, config as _cfg
    importlib.reload(_cfg)
    for name in ("AI_API_KEY", "AI_BASE_URL", "AI_MODEL", "AI_TEMPERATURE"):
        if hasattr(_cfg, name):
            globals()[name] = getattr(_cfg, name)

def _mask_key(key: str) -> str:
    if not key:
        return "(未设置)"
    return key[:6] + "..." + key[-4:] if len(key) > 12 else "***"

# ─────────────────────────────────────────────────────────────
# UI 常量
# ─────────────────────────────────────────────────────────────
CONSOLE = Console()
THEME = "cyan"
MAX_CONTEXT_TOKENS = 128000
MAX_TOOL_RESULT_LEN = 8000

# ─────────────────────────────────────────────────────────────
# 寄存器地址表（按系列）
# ─────────────────────────────────────────────────────────────
_REG_F1 = {
    "RCC_CR": 0x40021000, "RCC_CFGR": 0x40021004,
    "RCC_APB1ENR": 0x4002101C, "RCC_APB2ENR": 0x40021018,
    "GPIOA_CRL": 0x40010800, "GPIOA_CRH": 0x40010804,
    "GPIOA_IDR": 0x40010808, "GPIOA_ODR": 0x4001080C,
    "GPIOB_CRL": 0x40010C00, "GPIOB_CRH": 0x40010C04,
    "GPIOB_IDR": 0x40010C08, "GPIOB_ODR": 0x40010C0C,
    "GPIOC_CRL": 0x40011000, "GPIOC_CRH": 0x40011004,
    "TIM1_CR1": 0x40012C00, "TIM1_CCER": 0x40012C20,
    "TIM2_CR1": 0x40000000, "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400, "TIM3_CCER": 0x40000420,
    "ADC1_SR": 0x40012400, "ADC1_CR2": 0x40012408,
    "I2C1_CR1": 0x40005400, "I2C1_SR1": 0x40005414,
    "I2C2_CR1": 0x40005800, "I2C2_SR1": 0x40005814,
    "USART1_SR": 0x40013800, "USART1_BRR": 0x40013808,
}
_REG_F4 = {
    "RCC_CR": 0x40023800, "RCC_CFGR": 0x40023808,
    "RCC_AHB1ENR": 0x40023830, "RCC_APB1ENR": 0x40023840, "RCC_APB2ENR": 0x40023844,
    "GPIOA_MODER": 0x40020000, "GPIOA_IDR": 0x40020010, "GPIOA_ODR": 0x40020014,
    "GPIOB_MODER": 0x40020400, "GPIOB_IDR": 0x40020410, "GPIOB_ODR": 0x40020414,
    "GPIOC_MODER": 0x40020800, "GPIOC_IDR": 0x40020810, "GPIOC_ODR": 0x40020814,
    "TIM2_CR1": 0x40000000, "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400, "TIM3_CCER": 0x40000420,
    "I2C1_CR1": 0x40005400, "I2C1_SR1": 0x40005414,
    "USART1_SR": 0x40011000, "USART1_BRR": 0x40011008,
}
_REG_F0F3 = {
    "RCC_CR": 0x40021000, "RCC_CFGR": 0x40021004,
    "RCC_AHBENR": 0x40021014, "RCC_APB2ENR": 0x40021018, "RCC_APB1ENR": 0x4002101C,
    "GPIOA_MODER": 0x48000000, "GPIOA_IDR": 0x48000010, "GPIOA_ODR": 0x48000014,
    "GPIOB_MODER": 0x48000400, "GPIOB_IDR": 0x48000410, "GPIOB_ODR": 0x48000414,
    "I2C1_CR1": 0x40005400,
    "USART1_BRR": 0x40013808, "USART1_CR1": 0x4001380C,
}
_REG_COMMON = {
    "SCB_CFSR": 0xE000ED28, "SCB_HFSR": 0xE000ED2C,
    "SCB_BFAR": 0xE000ED38, "NVIC_ISER0": 0xE000E100,
}


def _reg_map(family: str) -> dict:
    base = {"f1": _REG_F1, "f4": _REG_F4, "f0": _REG_F0F3, "f3": _REG_F0F3}
    regs = dict(base.get(family.lower(), _REG_F1))
    regs.update(_REG_COMMON)
    return regs

# ─────────────────────────────────────────────────────────────
# PyOCDBridge（替换 OpenOCD）
# ─────────────────────────────────────────────────────────────
class PyOCDBridge:
    """
    使用 pyocd Python 库直接控制 STM32（无需启动 openocd 进程）。
    支持所有 CMSIS-DAP / ST-Link / J-Link USB 探针。
    安装：pip install pyocd
    """

    def __init__(self):
        self._session = None
        self._target = None
        self.connected = False
        self.chip_info: dict = {}
        self._family = "f1"
        self._reg_map: dict = _reg_map("f1")

    # ---- 内部工具 ----
    def _chip_to_pyocd_target(self, chip: str) -> str:
        """将 STM32F103C8T6 → stm32f103c8（去掉封装+温度后缀）"""
        import re
        name = chip.lower().strip()
        # STM32 命名末尾：封装字母(T/U/H/Y) + 温度等级数字(3/6/7)，如 T6/U3/H7
        name = re.sub(r'[a-z]\d$', '', name)
        return name

    _pyocd_target_cache: Optional[tuple] = None  # (float, set)
    _CACHE_TTL = 60.0  # 秒

    @classmethod
    def _get_all_pyocd_targets(cls) -> set:
        """获取所有可用 pyocd 目标，60秒内复用缓存"""
        now = time.time()
        if cls._pyocd_target_cache is not None:
            ts, cached = cls._pyocd_target_cache
            if now - ts < cls._CACHE_TTL:
                return cached

        try:
            result = subprocess.run(
                [sys.executable, "-m", "pyocd", "list", "--targets"],
                capture_output=True, text=True, timeout=15
            )
            known = set()
            for line in result.stdout.splitlines():
                parts = line.split()
                if parts and parts[0].startswith("stm32"):
                    known.add(parts[0].lower())
            cls._pyocd_target_cache = (now, known)
            return known
        except Exception:
            return set()

    def _resolve_best_target(self, target_name: str) -> str:
        """在 pyocd 所有可用目标（builtin + pack）中找最佳匹配（精确→前缀→截短）"""
        known = self._get_all_pyocd_targets()

        # 若子进程查询失败，降级到 TARGET 字典
        if not known:
            try:
                from pyocd.target import TARGET
                known = {k.lower() for k in TARGET}
            except ImportError:
                return target_name

        # 1. 精确匹配
        if target_name in known:
            return target_name

        # 2. 前缀匹配：target_name 是 known 中某目标的前缀（不太可能，但防御性保留）
        candidates = [k for k in known if k.startswith(target_name)]
        if candidates:
            best = max(candidates, key=lambda k: len(os.path.commonprefix([target_name, k])))
            CONSOLE.print(f"[yellow]  目标映射: {target_name} → {best}[/]")
            return best

        # 3. 截短搜索：逐位去尾，找相同系列最接近的目标
        for trim in range(1, 4):
            prefix = target_name[:-trim]
            if len(prefix) < 8:  # stm32fXX 最短前缀
                break
            candidates = [k for k in known if k.startswith(prefix)]
            if candidates:
                best = max(candidates, key=lambda k: len(os.path.commonprefix([target_name, k])))
                CONSOLE.print(f"[yellow]  目标近似匹配: {target_name} → {best}[/]")
                return best

        return target_name  # 未找到，原样返回让 pyocd 报错

    def _auto_install_pack(self, target_name: str) -> bool:
        """自动安装 pyocd CMSIS pack，返回是否成功"""
        CONSOLE.print(f"[yellow]  未找到目标 {target_name}，正在自动安装支持包...[/]")
        try:
            result = subprocess.run(
                [sys.executable, "-m", "pyocd", "pack", "install", target_name],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode == 0:
                CONSOLE.print(f"[green]  支持包安装成功[/]")
                return True
            else:
                CONSOLE.print(f"[red]  支持包安装失败: {result.stderr.strip()}[/]")
                return False
        except Exception as e:
            CONSOLE.print(f"[red]  支持包安装出错: {e}[/]")
            return False

    def _detect_family(self, chip: str) -> str:
        chip_up = chip.upper()
        if "F0" in chip_up:   return "f0"
        if "F3" in chip_up:   return "f3"
        if "F4" in chip_up or "F7" in chip_up or "H7" in chip_up: return "f4"
        return "f1"

    def set_family(self, family: str):
        self._family = family
        self._reg_map = _reg_map(family)

    # ---- 连接 / 断开 ----
    def start(self, chip: str = DEFAULT_CHIP) -> bool:
        """连接第一个可用探针，成功返回 True"""
        self.stop()
        try:
            from pyocd.core.helpers import ConnectHelper
        except ImportError:
            CONSOLE.print("[red]pyocd 未安装，请运行: pip install pyocd[/]")
            return False

        # chip 为 None 时让 pyocd 自动检测目标
        explicit_chip = chip and chip.upper() != "AUTO"
        family = self._detect_family(chip) if explicit_chip else "f1"
        self.set_family(family)

        if explicit_chip:
            raw_target = self._chip_to_pyocd_target(chip)
            # 在已知目标库中找最佳匹配（精确→前缀→截短）
            target_name = self._resolve_best_target(raw_target)
        else:
            target_name = None

        probe_hint = f"目标: {target_name}" if target_name else "自动检测目标"
        CONSOLE.print(f"[dim]  连接探针（{probe_hint}）...[/]")

        def _do_connect(t_name):
            return ConnectHelper.session_with_chosen_probe(
                target_override=t_name,
                auto_unlock=True,
                connect_mode="halt",
                blocking=False,
                return_first=True,
                options={"frequency": 1000000},
            )

        try:
            self._session = _do_connect(target_name)
        except Exception as e:
            err_str = str(e)
            # 目标不认识 → 自动安装 pack 后重试一次
            if explicit_chip and ("not recognized" in err_str or "Target type" in err_str):
                raw_target = self._chip_to_pyocd_target(chip)
                if self._auto_install_pack(raw_target):
                    # pack 安装后重新解析（新 session 会扫描已安装 pack）
                    target_name = self._resolve_best_target(raw_target)
                    CONSOLE.print(f"[dim]  重新连接（{target_name}）...[/]")
                    try:
                        self._session = _do_connect(target_name)
                    except Exception as e2:
                        CONSOLE.print(f"[red]  连接失败: {e2}[/]")
                        self._session = None
                        self._target = None
                        return False
                else:
                    CONSOLE.print(f"[red]  连接失败: {e}[/]")
                    self._session = None
                    self._target = None
                    return False
            else:
                CONSOLE.print(f"[red]  连接失败: {e}[/]")
                self._session = None
                self._target = None
                return False

        try:
            if self._session is None:
                CONSOLE.print("[red]  未找到调试探针，请检查 USB 连接[/]")
                return False
            self._session.open()
            self._target = self._session.board.target

            # 读取 pyocd 实际识别到的目标型号
            detected = getattr(self._target, "target_type", None) or (
                target_name or "unknown"
            )
            # 若是自动检测，用检测到的型号更新 chip 变量
            resolved_chip = chip.upper() if explicit_chip else detected.upper()
            resolved_family = self._detect_family(resolved_chip)
            self.set_family(resolved_family)

            self.chip_info = {
                "device": resolved_chip,
                "pyocd_target": detected,
                "family": resolved_family,
                "probe": self._session.board.description,
            }
            self.connected = True
            CONSOLE.print(
                f"[green]  已连接: {resolved_chip} "
                f"| 探针: {self._session.board.description}[/]"
            )
            # Warmup：halt→读CPUID，稳定 SWD 会话，保持 halt 状态以便烧录
            try:
                self._target.halt()
                time.sleep(0.1)
                self._target.read32(0xE000ED00)  # CPUID，只读安全寄存器
                time.sleep(0.05)
                # 不 resume，保持 halt——烧录前必须 halt，reconnect 后也不例外
            except Exception:
                pass
            return True
        except Exception as e:
            CONSOLE.print(f"[red]  连接后处理失败: {e}[/]")
            self._session = None
            self._target = None
            return False

    def stop(self):
        if self._session:
            try:
                self._target.resume()
            except Exception:
                pass
            try:
                self._session.close()
            except Exception:
                pass
            self._session = None
            self._target = None
        self.connected = False

    # ---- 烧录 ----
    def flash(self, bin_path: str) -> dict:
        if not self.connected:
            return {"ok": False, "msg": "探针未连接，请先 connect"}
        p = Path(bin_path)
        if not p.exists():
            return {"ok": False, "msg": f"文件不存在: {bin_path}"}

        try:
            from pyocd.flash.file_programmer import FileProgrammer
        except ImportError:
            return {"ok": False, "msg": "pyocd 未安装"}

        size = p.stat().st_size
        t0 = time.time()
        CONSOLE.print(f"[dim]  烧录 {size} 字节...[/]")
        # 烧录前 reset_and_halt，将 MCU 恢复到干净复位状态再写 flash
        # 仅 halt() 不够——若前一个固件开了 IWDG 或 I2C 卡死，flash 算法也会被影响
        try:
            self._target.reset_and_halt()
            time.sleep(0.1)
        except Exception:
            try:
                self._target.halt()
                time.sleep(0.05)
            except Exception:
                pass
        try:
            programmer = FileProgrammer(self._session)
            programmer.program(str(p), base_address=0x08000000)
        except Exception as e:
            return {"ok": False, "msg": f"烧录异常: {e}"}

        dt = time.time() - t0
        spd = size / dt / 1024 if dt > 0 else 0
        # 烧录后复位并运行（失败不致命，固件已写入）
        try:
            self._target.reset_and_halt()
            time.sleep(0.1)
            self._target.resume()
        except Exception as e:
            CONSOLE.print(f"[yellow]  复位警告（固件已烧录）: {e}[/]")
        return {"ok": True, "msg": f"烧录成功 {size}B / {dt:.1f}s ({spd:.1f} KB/s)"}

    # ---- 寄存器读取 ----
    def read_registers(self, names: Optional[list] = None) -> Optional[dict]:
        if not self.connected:
            return None
        try:
            self._target.halt()
            time.sleep(REGISTER_READ_DELAY)

            targets = names if names else list(self._reg_map.keys())
            regs = {}
            for name in targets:
                addr = self._reg_map.get(name)
                if addr is None:
                    continue
                try:
                    val = self._target.read32(addr)
                    regs[name] = f"0x{val:08X}"
                except Exception:
                    pass

            # 读 PC
            try:
                pc = self._target.read_core_register("pc")
                regs["PC"] = f"0x{pc:08X}"
            except Exception:
                pass

            self._target.resume()
            return regs
        except Exception as e:
            CONSOLE.print(f"[red]  寄存器读取异常: {e}[/]")
            return None

    def read_all_for_debug(self) -> Optional[dict]:
        return self.read_registers()

    def analyze_fault(self, regs: dict) -> str:
        cfsr_str = regs.get("SCB_CFSR", "0x00000000")
        try:
            cfsr = int(cfsr_str, 16)
        except ValueError:
            return "CFSR 格式错误"
        if cfsr == 0:
            return "无故障"
        checks = [
            (0x01,      "IACCVIOL: 指令访问违规"),
            (0x02,      "DACCVIOL: 数据访问违规"),
            (0x100,     "IBUSERR: 指令总线错误"),
            (0x200,     "PRECISERR: 精确总线错误（外设未使能时钟）"),
            (0x400,     "IMPRECISERR: 非精确总线错误"),
            (0x10000,   "UNDEFINSTR: 未定义指令"),
            (0x20000,   "INVSTATE: 无效 EPSR 状态"),
            (0x1000000, "UNALIGNED: 非对齐访问"),
            (0x2000000, "DIVBYZERO: 除零"),
        ]
        faults = [desc for mask, desc in checks if cfsr & mask]
        return "; ".join(faults) if faults else f"未知故障 CFSR=0x{cfsr:08X}"

    def list_probes(self) -> list:
        """列出所有可用探针"""
        try:
            from pyocd.core.session import Session
            from pyocd.probe.aggregator import DebugProbeAggregator
            probes = DebugProbeAggregator.get_all_connected_probes()
            return [{"uid": p.unique_id, "description": p.product_name} for p in probes]
        except Exception:
            return []


# ─────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────
# 串口自动检测
# ─────────────────────────────────────────────────────────────
def detect_serial_ports(verbose: bool = False) -> list:
    """
    跨平台扫描可用串口，按 STM32 使用优先级排序。
    Windows : CH340/CP210x → COMx（数字小的优先）
    macOS   : tty.usbserial-* / tty.usbmodem* → cu.* 备选
    Linux   : ttyUSB* / ttyACM* → ttyAMA* → ttyS[4+]
    只返回当前用户有读写权限的端口。
    """
    import glob, platform, re

    plat = platform.system()   # 'Windows' / 'Darwin' / 'Linux'
    found = []                 # 有序去重列表，USB优先

    def _add(port: str, usb: bool = False):
        """加入列表，usb=True 时插到所有非usb端口之前"""
        if port in found:
            return
        if usb:
            # 找到第一个非usb端口的位置，插入其前
            idx = next((i for i, p in enumerate(found) if p not in _usb_set), len(found))
            found.insert(idx, port)
            _usb_set.add(port)
        else:
            found.append(port)

    _usb_set: set = set()

    # ── 1. pyserial list_ports（所有平台最可靠的来源）────────────
    try:
        import serial.tools.list_ports as lp
        skip_kw = ("bluetooth", "virtual", "rfcomm", "modem")
        for info in lp.comports():
            port = info.device
            desc = (info.description or "").lower()
            hwid = (info.hwid or "n/a").lower()
            if any(k in desc for k in skip_kw):
                continue
            # 跳过 hwid='n/a' 的 ttyS：是内核注册的幽灵串口，没有实际硬件
            if hwid == "n/a" and re.search(r"ttyS\d+$", port):
                continue
            is_usb = "usb" in hwid or "ch34" in hwid or "cp21" in hwid or "ft23" in hwid
            _add(port, usb=is_usb)
    except Exception:
        pass

    # ── 2. 平台专属补充扫描（pyserial 有时漏掉设备）───────────────
    if plat == "Windows":
        # Windows: 枚举 COM1-COM256，跳过已找到的
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                                 r"HARDWARE\DEVICEMAP\SERIALCOMM")
            i = 0
            while True:
                try:
                    _, port, _ = winreg.EnumValue(key, i)
                    _add(port)
                    i += 1
                except OSError:
                    break
        except Exception:
            pass

    elif plat == "Darwin":
        # macOS: tty.usbserial / tty.usbmodem 优先，cu.* 备选
        for pattern, is_usb in [
            ("/dev/tty.usbserial*", True),
            ("/dev/tty.usbmodem*",  True),
            ("/dev/tty.SLAB*",      True),   # CP210x
            ("/dev/tty.wchusbserial*", True), # CH340
            ("/dev/cu.usbserial*",  True),
            ("/dev/cu.usbmodem*",   True),
            ("/dev/tty.*",          False),
        ]:
            for p in sorted(glob.glob(pattern)):
                if os.access(p, os.R_OK | os.W_OK):
                    _add(p, usb=is_usb)

    else:  # Linux / FreeBSD / 其他 POSIX
        # USB 转串口（最高优先）
        for pattern in ["/dev/ttyUSB*", "/dev/ttyACM*"]:
            for p in sorted(glob.glob(pattern)):
                real = os.path.realpath(p)
                if os.access(real, os.R_OK | os.W_OK):
                    _add(real, usb=True)
        # by-id 符号链接（解析后去重，仍属 USB 优先）
        for p in sorted(glob.glob("/dev/serial/by-id/*")):
            real = os.path.realpath(p)
            if os.access(real, os.R_OK | os.W_OK) and real not in found:
                _add(real, usb=True)
        # SBC 硬件 UART 补充（ttyAMA 始终可靠；ttyS 用 sysfs 验证有无实际硬件）
        def _has_sysfs_device(port: str) -> bool:
            """Linux: /sys/class/tty/ttySx/device 存在 = 真实硬件 UART"""
            name = os.path.basename(port)
            return os.path.exists(f"/sys/class/tty/{name}/device")

        for p in sorted(glob.glob("/dev/ttyAMA*")):
            if os.access(p, os.R_OK | os.W_OK) and p not in found:
                _add(p, usb=False)

        for p in sorted(glob.glob("/dev/ttyS*")):
            if p in found or not os.access(p, os.R_OK | os.W_OK):
                continue
            if _has_sysfs_device(p):   # 有 sysfs device 条目 = 真实 UART
                _add(p, usb=False)

    if verbose:
        CONSOLE.print(f"[dim]  检测到串口: {found if found else '无'}[/]")
    return found


def auto_open_serial(baud: int = SERIAL_BAUD) -> tuple:
    """
    自动检测并尝试打开第一个可用串口，返回 (port, serial_obj) 或 (None, None)。
    """
    try:
        import serial as pyserial
    except ImportError:
        return None, None

    candidates = detect_serial_ports()
    for port in candidates:
        try:
            s = pyserial.Serial(port, baud, timeout=0.3)
            s.close()
            return port, None   # 可以打开，返回 port
        except Exception:
            continue
    return None, None


# ─────────────────────────────────────────────────────────────
# SerialMonitor（与 hardware.py 保持一致）
# ─────────────────────────────────────────────────────────────
class SerialMonitor:
    def __init__(self):
        self._serial = None
        self._port = None
        self._buffer = ""
        self._lock = threading.Lock()
        self._thread = None
        self._running = False

    def open(self, port: str = None, baud: int = SERIAL_BAUD) -> bool:
        try:
            import serial as pyserial
        except ImportError:
            CONSOLE.print("[yellow]  pyserial 未安装: pip install pyserial[/]")
            return False

        # 确定要尝试的端口列表
        if port:
            # 指定了端口：先试指定的，失败再自动扫描
            candidates = [port] + [p for p in detect_serial_ports() if p != port]
        else:
            # 未指定：完全自动检测
            candidates = detect_serial_ports()
            if not candidates:
                CONSOLE.print("[yellow]  串口: 未检测到任何可用串口[/]")
                return False

        for p in candidates:
            try:
                self._serial = pyserial.Serial(p, baud, timeout=0.5)
                self._serial.reset_input_buffer()
                self._running = True
                self._thread = threading.Thread(target=self._reader, daemon=True)
                self._thread.start()
                self._port = p
                CONSOLE.print(f"[green]  串口: {p} @ {baud}[/]")
                return True
            except Exception as e:
                if p == candidates[-1]:
                    # 最后一个也失败了，给出有用提示（跨平台）
                    import platform as _plt
                    if _plt.system() == "Linux" and not os.access(p, os.R_OK | os.W_OK):
                        try:
                            import grp as _grp
                            grp_name = _grp.getgrgid(os.stat(p).st_gid).gr_name
                        except Exception:
                            grp_name = "dialout"
                        CONSOLE.print(
                            f"[yellow]  串口: {p} 权限不足 → "
                            f"sudo usermod -aG {grp_name} $USER && newgrp {grp_name}[/]"
                        )
                    else:
                        CONSOLE.print(f"[yellow]  串口打开失败: {e}[/]")
                # 继续尝试下一个
                continue
        return False

    def _reader(self):
        try:
            import serial as _pyserial
            _SerialException = _pyserial.SerialException
        except ImportError:
            _SerialException = OSError

        consecutive_errors = 0
        while self._running and self._serial:
            try:
                data = self._serial.read(1024)
                if data:
                    consecutive_errors = 0
                    with self._lock:
                        self._buffer += data.decode("utf-8", errors="ignore")
                        if len(self._buffer) > 8192:
                            self._buffer = self._buffer[-8192:]
            except _SerialException:
                # 串口物理断开
                CONSOLE.print("[yellow]  ⚠ 串口断开[/]")
                self._running = False
                break
            except Exception:
                consecutive_errors += 1
                if consecutive_errors > 10:
                    CONSOLE.print("[yellow]  ⚠ 串口持续异常，停止读取[/]")
                    self._running = False
                    break
                time.sleep(0.1)

    def read_and_clear(self) -> str:
        with self._lock:
            out = self._buffer
            self._buffer = ""
            return out

    def clear(self):
        with self._lock:
            self._buffer = ""

    def wait_for(self, keyword: str, timeout: float = 5.0, clear_first: bool = True) -> str:
        if clear_first:
            self.clear()
        t0 = time.time()
        while time.time() - t0 < timeout:
            with self._lock:
                if keyword in self._buffer:
                    break
            time.sleep(0.1)
        time.sleep(0.3)
        return self.read_and_clear()

    def close(self):
        self._running = False
        if self._serial:
            try:
                self._serial.close()
            except Exception:
                pass
            self._serial = None

def _wait_serial_adaptive(
    serial,
    keyword: str,
    min_wait: float = 0.5,
    max_wait: float = 8.0,
) -> str:
    """
    自适应串口等待：
    - 先等 min_wait 秒（给 MCU 复位时间）
    - 之后每 200ms 采样一次，检测到 keyword 或有内容即停
    - 超过 max_wait 后强制返回
    """
    time.sleep(min_wait)
    t0 = time.time()
    accumulated = ""
    while time.time() - t0 < (max_wait - min_wait):
        chunk = serial.read_and_clear()
        if chunk:
            accumulated += chunk
            if keyword in accumulated:
                break
        time.sleep(0.2)
    time.sleep(0.3)
    accumulated += serial.read_and_clear()
    return accumulated


# ─────────────────────────────────────────────────────────────
# 全局硬件状态（工具函数直接访问）
# ─────────────────────────────────────────────────────────────
_compiler: Optional[Compiler] = None
_bridge: Optional[PyOCDBridge] = None
_serial: Optional[SerialMonitor] = None
_hw_connected = False
_serial_connected = False
_current_chip = DEFAULT_CHIP
_last_bin_path: Optional[str] = None
_last_code: Optional[str] = None

# 调试闭环计数器（每次新任务由 AI 调用 stm32_reset_debug_attempts 重置）
_debug_attempt = 0
MAX_DEBUG_ATTEMPTS = 8  # 提高上限，一个任务最多 8 轮（含修改迭代）


def _get_compiler() -> Compiler:
    global _compiler
    if _compiler is None:
        _compiler = Compiler()
        _compiler.set_chip(_current_chip)
    return _compiler


def _get_bridge() -> PyOCDBridge:
    global _bridge
    if _bridge is None:
        _bridge = PyOCDBridge()
    return _bridge


def _get_serial() -> SerialMonitor:
    global _serial
    if _serial is None:
        _serial = SerialMonitor()
    return _serial


# ─────────────────────────────────────────────────────────────
# STM32 专属工具实现
# ─────────────────────────────────────────────────────────────

def stm32_generate_font(text: str, size: int = 16) -> dict:
    """
    将任意文字（含中文）渲染为 STM32 OLED 用的 C 点阵数组。
    固定使用「横向取模·高位在前（row-major, MSB=left）」格式——
    这是最直观、与 SSD1306 逐行刷新最匹配的格式，配套显示函数一并生成。
    返回 c_code（字模数组 + 完整显示函数），直接粘贴进 main.c 使用。
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return {"success": False, "message": "需要安装 Pillow: pip install Pillow"}

    import platform as _plat
    # 跨平台字体候选
    def _find_cjk_font() -> Optional[str]:
        """动态查找系统 CJK 字体路径"""
        # 优先用 fc-match（Linux/macOS）
        try:
            r = subprocess.run(
                ["fc-match", "--format=%{file}", ":lang=zh"],
                capture_output=True, text=True, timeout=5
            )
            if r.returncode == 0 and r.stdout.strip():
                path = r.stdout.strip()
                if os.path.exists(path):
                    return path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 回退到硬编码候选列表
        candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None
        # 原来的 for fp in font_candidates: ... 全部替换为：
    font_path = _find_cjk_font()
    if font_path is None:
        return {"success": False, "message": "未找到中文字体，请安装 fonts-noto-cjk"}
    try:
        font = ImageFont.truetype(font_path, size)
    except Exception as e:
        return {"success": False, "message": f"字体加载失败 ({font_path}): {e}"}
        
    def _render_char(char: str) -> list:
        """渲染单个字符到 size×size 位图，返回 0/1 列表（行优先）"""
        img = Image.new("L", (size, size), 0)
        draw = ImageDraw.Draw(img)
        try:
            bbox = font.getbbox(char)
            char_w = bbox[2] - bbox[0]
            char_h = bbox[3] - bbox[1]
            # 水平居中，垂直顶对齐（避免因字体 metrics 差异被裁底部）
            ox = (size - char_w) // 2 - bbox[0]
            oy = -bbox[1]  # 让字符顶部与图像顶部对齐
        except Exception:
            ox, oy = 0, 0
        draw.text((ox, oy), char, fill=255, font=font)
        return [1 if p > 127 else 0 for p in img.getdata()]

    def _to_row_msb(pixels: list) -> list:
        """横向取模·高位在前：每行从左到右，bit7=最左列"""
        data = []
        bytes_per_row = (size + 7) // 8
        for row in range(size):
            for b in range(bytes_per_row):
                byte = 0
                for bit in range(8):
                    col = b * 8 + bit
                    if col < size and pixels[row * size + col]:
                        byte |= (1 << (7 - bit))
                data.append(byte)
        return data

    def _ascii_preview(pixels: list) -> str:
        lines = []
        for row in range(size):
            lines.append("".join("█" if pixels[row * size + col] else "." for col in range(size)))
        return "\n".join(lines)

    chars_data = []
    previews = []
    char_list = []
    for char in text:
        pixels = _render_char(char)
        data = _to_row_msb(pixels)
        chars_data.append(data)
        previews.append(_ascii_preview(pixels))
        char_list.append(char)

    bytes_per_char = size * ((size + 7) // 8)
    fname = f"FONT_{size}x{size}"

    # ── 字模数组 ──────────────────────────────────────────────
    char_entries = []
    for i, (char, data) in enumerate(zip(char_list, chars_data)):
        preview_lines = previews[i].split("\n")
        preview_comment = "  /* " + "  ".join(preview_lines[:4]) + " ... */"
        hex_str = ", ".join(f"0x{b:02X}" for b in data)
        char_repr = char if ord(char) < 128 else f"{char}(U+{ord(char):04X})"
        char_entries.append(f"    /* [{i}] '{char_repr}' */\n    {{{hex_str}}}")

    array_code = (
        f"/* ═══ 字模数据：横向取模·高位在前 {size}x{size}px ═══\n"
        f"   格式：每行 {(size+7)//8} 字节，bit7=最左列，共 {bytes_per_char} 字节/字符\n"
        f"   字符表: {' '.join(repr(c) for c in char_list)} */\n"
        f"static const uint8_t {fname}[][{bytes_per_char}] = {{\n"
        + ",\n".join(char_entries)
        + "\n};\n"
    )

    # ── 配套显示函数（与字模格式严格匹配）──────────────────────
    display_func = f"""
/* ═══ 配套显示函数（必须与上面字模数据一起使用）═══ */
/* idx: 字符在 {fname} 中的下标（按字符表顺序） */
/* x,y: OLED 列(0-127)和页起始行(0-63)         */
void OLED_ShowFont{size}(uint8_t x, uint8_t y, uint8_t idx) {{
    const uint8_t *p = {fname}[idx];
    uint8_t bytes_per_row = {(size+7)//8};
    for (uint8_t row = 0; row < {size}; row++) {{
        OLED_SetCursor(x, y + row);   /* 设置到目标行 */
        for (uint8_t b = 0; b < bytes_per_row; b++) {{
            uint8_t byte = p[row * bytes_per_row + b];
            for (int8_t bit = 7; bit >= 0; bit--) {{
                uint8_t col = b * 8 + (7 - bit);
                if (col < {size}) {{
                    OLED_DrawPixel(x + col, y + row, (byte >> bit) & 1);
                }}
            }}
        }}
    }}
}}
/* 用法示例：显示字符表第0个字符在 (0,0) 位置
   OLED_ShowFont{size}(0, 0, 0);  // 显示 '{char_list[0] if char_list else "?"}' */
"""

    # ── ASCII 预览（调试用）────────────────────────────────────
    preview_block = "\n\n".join(
        f"/* '{c}':\n{p} */"
        for c, p in zip(char_list, previews)
    )

    c_code = array_code + display_func
    return {
        "success": True,
        "c_code": c_code,
        "preview": preview_block,        # ASCII 预览，可用于肉眼验证字形
        "char_count": len(text),
        "bytes_per_char": bytes_per_char,
        "font_size": size,
        "mode": "row_msb",
        "char_order": char_list,
    }


def stm32_list_probes() -> dict:
    """列出所有可用的调试探针（ST-Link / CMSIS-DAP / J-Link）"""
    probes = _get_bridge().list_probes()
    if not probes:
        return {"success": True, "probes": [], "message": "未检测到任何探针，请检查 USB 连接"}
    return {"success": True, "probes": probes}


def stm32_connect(chip: str = None) -> dict:
    """连接 STM32 硬件（pyocd 探针 + 串口监控）"""
    global _hw_connected, _serial_connected, _current_chip
    # chip=None → 使用当前已知芯片或默认芯片，避免 pyocd 退化到 generic cortex_m
    target_chip = chip or _current_chip or DEFAULT_CHIP
    bridge = _get_bridge()
    serial = _get_serial()

    if bridge.start(target_chip):
        _hw_connected = True
        # 用 pyocd 实际识别到的型号（自动检测场景下会与传入值不同）
        _current_chip = bridge.chip_info.get("device", target_chip)
        _get_compiler().set_chip(_current_chip)
        _serial_connected = serial.open()
        return {
            "success": True,
            "chip": _current_chip,
            "probe": bridge.chip_info.get("probe", ""),
            "serial_connected": _serial_connected,
            "message": f"硬件已连接: {_current_chip}",
        }
    _hw_connected = False
    return {"success": False, "message": "连接失败，请检查探针 USB 连接和驱动"}


def stm32_serial_connect(port: str = None, baud: int = None) -> dict:
    """
    单独连接/重连 UART 串口（不影响 pyocd 探针连接）。
    用于更换串口设备或在 stm32_connect 之后补充连接串口。
    """
    global _serial_connected
    serial = _get_serial()
    # 若已连接先关闭
    serial.close()
    _serial_connected = False

    use_baud = baud or SERIAL_BAUD
    # port=None → 自动检测；否则先试指定端口，失败时扫描其他
    _serial_connected = serial.open(port or None, use_baud)
    actual_port = getattr(serial, "_port", port or "自动检测")
    if _serial_connected:
        return {"success": True, "port": actual_port, "baud": use_baud,
                "message": f"串口已连接: {actual_port} @ {use_baud}"}
    candidates = detect_serial_ports()
    return {"success": False, "port": port, "baud": use_baud,
            "message": f"串口打开失败，可用端口: {candidates if candidates else '无'}"}


def stm32_serial_disconnect() -> dict:
    """断开串口（保留 pyocd 探针连接）"""
    global _serial_connected
    _get_serial().close()
    _serial_connected = False
    return {"success": True, "message": "串口已断开"}


def stm32_disconnect() -> dict:
    """断开硬件连接（释放探针和串口）"""
    global _hw_connected, _serial_connected
    _get_bridge().stop()
    _get_serial().close()
    _hw_connected = False
    _serial_connected = False
    return {"success": True, "message": "已断开"}


def stm32_set_chip(chip: str) -> dict:
    """切换目标芯片型号（如 STM32F103C8T6 / STM32F407VET6）"""
    global _current_chip
    _current_chip = chip.strip().upper()
    ci = _get_compiler().set_chip(_current_chip)
    if _hw_connected:
        _get_bridge().set_family(ci.get("family", "f1"))
    return {"success": True, "chip": _current_chip, "family": ci.get("family", "f1")}


def stm32_hardware_status() -> dict:
    """获取当前硬件连接状态和工具链可用性"""
    ci = _get_compiler().check(_current_chip)
    return {
        "chip": _current_chip,
        "hw_connected": _hw_connected,
        "serial_connected": _serial_connected,
        "gcc_ok": ci.get("gcc", False),
        "gcc_version": ci.get("gcc_version", "未找到"),
        "hal_ok": ci.get("hal", False),
        "hal_lib_ok": ci.get("hal_lib", False),
        "workspace": str(WORKSPACE),
    }


def stm32_compile(code: str, chip: str = None) -> dict:
    """编译 STM32 C 代码（完整 main.c）"""
    global _last_bin_path, _last_code
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile(code)
    if result["ok"]:
        _last_code = code
        _last_bin_path = result.get("bin_path")
        # 自动保存到 latest_workspace
        try:
            latest = Path.home() / ".stm32agent" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    return {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }


def stm32_compile_rtos(code: str, chip: str = None) -> dict:
    """编译带 FreeRTOS 内核的完整 main.c 代码"""
    global _last_bin_path, _last_code
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile_rtos(code)
    if result["ok"]:
        _last_code = code
        _last_bin_path = result.get("bin_path")
        try:
            latest = Path.home() / ".stm32agent" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    return {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }


def stm32_flash(bin_path: str = None) -> dict:
    """烧录固件到 STM32（需要先 connect + compile）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接，请先调用 stm32_connect"}
    path = bin_path or _last_bin_path
    if not path or not Path(path).exists():
        return {"success": False, "message": f"固件文件不存在: {path}"}
    _get_serial().clear()
    r = _get_bridge().flash(path)
    return {"success": r["ok"], "message": r["msg"]}


def stm32_read_registers(regs: list = None) -> dict:
    """读取 STM32 硬件寄存器（RCC、GPIO、TIM、UART 等）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}
    result = _get_bridge().read_registers(regs) if regs else _get_bridge().read_all_for_debug()
    if result is not None:
        return {"success": True, "registers": result}
    return {"success": False, "message": "寄存器读取失败"}


def stm32_analyze_fault() -> dict:
    """读取并分析 HardFault 寄存器（SCB_CFSR / SCB_HFSR）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}
    regs = _get_bridge().read_registers(["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC"])
    if not regs:
        return {"success": False, "message": "寄存器读取失败"}
    analysis = _get_bridge().analyze_fault(regs)
    return {"success": True, "registers": regs, "analysis": analysis}


def stm32_serial_read(timeout: float = 3.0, wait_for: str = None) -> dict:
    """读取 UART 串口输出（调试日志）"""
    if not _serial_connected:
        return {"success": False, "message": "串口未连接"}
    serial = _get_serial()
    if wait_for:
        output = serial.wait_for(wait_for, timeout=timeout)
    else:
        time.sleep(min(timeout, 2.0))
        output = serial.read_and_clear()
    return {"success": True, "output": output, "has_output": bool(output.strip())}


def stm32_reset_debug_attempts() -> dict:
    """重置调试轮次计数器。开始一个全新需求时调用，确保计数从 1 开始。"""
    global _debug_attempt
    _debug_attempt = 0
    return {"success": True, "message": "计数器已重置"}


def stm32_auto_flash_cycle(code: str, request: str = "") -> dict:
    """
    完整开发闭环（自动计轮次，最多 MAX_DEBUG_ATTEMPTS 轮）：
      编译 → 烧录（若已连接硬件）→ 等待启动 → 读串口 → 读寄存器
    返回每步结果 + 当前轮次 + 是否应放弃。
    """
    global _last_code, _last_bin_path, _debug_attempt
    _debug_attempt += 1
    attempt = _debug_attempt
    steps = []

    # 超出最大轮次 → 直接告知 AI 放弃
    if attempt > MAX_DEBUG_ATTEMPTS:
        return {
            "success": False,
            "give_up": True,
            "attempt": attempt,
            "message": f"已达到最大调试轮次 ({MAX_DEBUG_ATTEMPTS})，自动调试无法解决，请检查硬件接线或手动排查",
            "steps": [],
        }

    remaining = MAX_DEBUG_ATTEMPTS - attempt
    CONSOLE.print(f"[dim]  第 {attempt}/{MAX_DEBUG_ATTEMPTS} 轮[/]")

    # 1. 编译
    comp = stm32_compile(code)
    steps.append({"step": "compile", "success": comp["success"], "msg": comp["message"][:300]})
    if not comp["success"]:
        return {
            "success": False, "attempt": attempt, "remaining": remaining,
            "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
            "steps": steps,
            "error": "编译失败，请根据错误信息修改代码",
            "compile_errors": comp["message"],
        }

    # 2. 烧录（失败时延迟重试，最多2次）
    if _hw_connected and comp.get("bin_path"):
        fr = stm32_flash(comp["bin_path"])
        if not fr["success"]:
            for retry in range(2):
                wait_sec = 1.5 * (retry + 1)
                CONSOLE.print(f"[yellow]  烧录失败（{fr['message'][:60]}），{wait_sec:.0f}s 后重连重试...[/]")
                time.sleep(wait_sec)
                stm32_connect(_current_chip)
                fr = stm32_flash(comp["bin_path"])
                if fr["success"]:
                    break
        steps.append({"step": "flash", "success": fr["success"], "msg": fr["message"]})
        if not fr["success"]:
            return {
                "success": False, "attempt": attempt, "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps, "error": f"烧录失败（重试2次后仍失败）: {fr['message']}",
            }
        # 3. 串口监控
        uart_out = ""
        sensor_errors = []
        if _serial_connected:
            CONSOLE.print("[dim]  等待启动...[/]")
            uart_out = _wait_serial_adaptive(
                _get_serial(),
                keyword="Gary:BOOT",
                min_wait=0.5,
                max_wait=POST_FLASH_DELAY + 4.0,
            )
            boot_ok = "Gary:BOOT" in uart_out
            # 检测传感器错误关键词
            sensor_errors = [line.strip() for line in uart_out.splitlines()
                             if "ERR:" in line or "Error" in line or "not found" in line.lower()]
            # 打印串口输出到终端供用户查看
            if uart_out.strip():
                CONSOLE.print(f"[dim]  串口输出:[/]\n[cyan]{uart_out.strip()[:400]}[/]")
            else:
                CONSOLE.print("[yellow]  串口无输出（程序未启动或卡死）[/]")
            steps.append({"step": "uart", "output": uart_out[:500], "boot_ok": boot_ok,
                          "sensor_errors": sensor_errors})
        else:
            time.sleep(POST_FLASH_DELAY + 1.0)
            boot_ok = True  # 无串口跳过验证

        # 4. 读寄存器（打印到终端）
        regs = _get_bridge().read_all_for_debug()
        has_fault = False
        if regs:
            KEY_SET = ("SCB_CFSR", "PC", "RCC_APB2ENR", "RCC_APB1ENR",
                       "GPIOA_CRL", "GPIOA_CRH", "GPIOA_ODR", "GPIOA_IDR",
                       "GPIOB_CRL", "GPIOB_CRH", "GPIOB_ODR", "GPIOB_IDR",
                       "GPIOC_CRL", "GPIOC_CRH", "GPIOC_ODR", "GPIOC_IDR",
                       "GPIOA_MODER", "GPIOB_MODER", "GPIOC_MODER",
                       "TIM2_CR1",
                       "I2C1_CR1", "I2C1_SR1",
                       "I2C2_CR1", "I2C2_SR1")
            key_regs = {k: v for k, v in regs.items() if k in KEY_SET}
            has_fault = regs.get("SCB_CFSR", "0x00000000") not in ("0x00000000", "0x0")
            # 打印寄存器到终端
            CONSOLE.print("[dim]  关键寄存器:[/]")
            for k, v in key_regs.items():
                color = "red" if (k == "SCB_CFSR" and has_fault) else "dim"
                CONSOLE.print(f"[{color}]    {k} = {v}[/]")
            steps.append({
                "step": "registers", "key_regs": key_regs,
                "has_hardfault": has_fault,
                "fault_analysis": _get_bridge().analyze_fault(regs) if has_fault else "",
            })
        else:
            CONSOLE.print("[yellow]  寄存器读取失败（探针连接问题）[/]")

        # ── 5. 硬件缺失检测（I2C NACK/ARLO = 外设未接）──
        hw_missing = []
        if regs:
            for i2c_name in ("I2C1", "I2C2"):
                sr1_str = regs.get(f"{i2c_name}_SR1", "0x00000000")
                try:
                    sr1 = int(sr1_str, 16)
                except ValueError:
                    continue
                if sr1 & 0x0400:  # bit10 = AF (Acknowledge Failure / NACK)
                    hw_missing.append(f"{i2c_name}: AF(NACK)——设备未应答，很可能未接或地址错误")
                if sr1 & 0x0200:  # bit9 = ARLO (Arbitration Lost)
                    hw_missing.append(f"{i2c_name}: ARLO(仲裁丢失)——总线无设备响应或接线错误")
        # 串口 ERR 也算硬件问题
        if sensor_errors:
            hw_missing.extend(sensor_errors)

        if hw_missing:
            CONSOLE.print(f"[red bold]  ⚠ 检测到硬件缺失:[/]")
            for m in hw_missing:
                CONSOLE.print(f"[red]    • {m}[/]")

        runtime_ok = boot_ok and not has_fault and not hw_missing

        if runtime_ok:
            if request:
                _stm32_save_project(code, comp, request)
            return {
                "success": True, "attempt": attempt,
                "steps": steps, "bin_size": comp.get("bin_size", 0),
            }
        elif hw_missing:
            # 硬件缺失 → 代码没问题，不要再修代码了，直接告知用户
            return {
                "success": False, "attempt": attempt, "remaining": remaining,
                "give_up": True,
                "hw_missing": hw_missing,
                "steps": steps,
                "error": (
                    "⚠ 硬件未接或接线错误（不是代码问题，停止修改代码）：\n"
                    + "\n".join(f"  • {m}" for m in hw_missing)
                    + "\n请告知用户检查硬件连接后重试。"
                ),
            }
        else:
            err_msg = "HardFault 或程序未正常启动，请根据 steps 中的寄存器和串口信息修复"
            if not boot_ok and not uart_out.strip():
                err_msg = "串口无任何输出——程序在打印 Gary: 之前就卡死了（常见原因：I2C 等待超时/传感器未接/死循环）"
            elif not boot_ok and uart_out.strip():
                err_msg = f"程序有输出但未打印 Gary: 启动标志，串口内容: {uart_out.strip()[:200]}"
            return {
                "success": False, "attempt": attempt, "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps,
                "error": err_msg,
            }

    # 无硬件 → 仅编译
    if request:
        _stm32_save_project(code, comp, request)
    return {
        "success": True, "attempt": attempt,
        "steps": steps,
        "note": "硬件未连接，已完成编译",
        "bin_size": comp.get("bin_size", 0),
        "bin_path": comp.get("bin_path"),
    }


def stm32_save_code(code: str, request: str = "untitled") -> dict:
    """保存代码到项目目录"""
    comp_result = {"bin_path": None, "bin_size": 0}
    path = _stm32_save_project(code, comp_result, request)
    return {"success": True, "path": str(path), "message": f"已保存: {path}"}


def _stm32_save_project(code: str, comp: dict, request: str) -> Path:
    """内部：保存项目文件"""
    global _last_code
    ts = time.strftime("%Y%m%d_%H%M%S")
    safe = "".join(c if c.isalnum() or c in "_- " else "" for c in request[:30]).strip()
    d = PROJECTS_DIR / f"{ts}_{safe}"
    d.mkdir(parents=True, exist_ok=True)
    (d / "main.c").write_text(code, encoding="utf-8")
    if comp.get("bin_path") and Path(comp["bin_path"]).exists():
        shutil.copy2(comp["bin_path"], d / "firmware.bin")
    (d / "config.json").write_text(json.dumps({
        "chip": _current_chip, "request": request,
        "bin_size": comp.get("bin_size", 0),
        "timestamp": ts,
    }, ensure_ascii=False, indent=2))
    _last_code = code
    CONSOLE.print(f"[dim]  已保存: {d}[/]")
    return d


def stm32_list_projects() -> dict:
    """列出最近 15 个历史项目"""
    if not PROJECTS_DIR.exists():
        return {"success": True, "projects": [], "message": "暂无项目"}
    projects = []
    for p in sorted(PROJECTS_DIR.iterdir(), reverse=True)[:15]:
        cf = p / "config.json"
        if cf.exists():
            try:
                c = json.loads(cf.read_text(encoding="utf-8"))
                projects.append({
                    "name": p.name,
                    "chip": c.get("chip", "?"),
                    "request": c.get("request", ""),
                    "timestamp": c.get("timestamp", ""),
                })
            except Exception:
                pass
    return {"success": True, "projects": projects}


def stm32_read_project(project_name: str) -> dict:
    """读取指定项目的 main.c 代码"""
    p = PROJECTS_DIR / project_name / "main.c"
    if not p.exists():
        return {"success": False, "message": f"项目不存在: {project_name}"}
    code = p.read_text(encoding="utf-8")
    return {"success": True, "code": code, "path": str(p), "lines": len(code.splitlines())}


# ─────────────────────────────────────────────────────────────
# 通用文件/命令工具（来自 claude_terminal，简化版）
# ─────────────────────────────────────────────────────────────

def read_file(file_path: str) -> dict:
    try:
        p = Path(file_path).resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        lines = content.splitlines()
        numbered = "\n".join(f"{i+1:4d} | {l}" for i, l in enumerate(lines[:800]))
        return {
            "success": True, "numbered_view": numbered,
            "raw_content": content[:40000], "total_lines": len(lines),
        }
    except Exception as e:
        return {"error": str(e)}


def create_or_overwrite_file(file_path: str, content: str) -> dict:
    try:
        p = Path(file_path).resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return {"success": True, "path": str(p), "lines": len(content.splitlines())}
    except Exception as e:
        return {"error": str(e)}


def str_replace_edit(file_path: str, old_str: str, new_str: str) -> dict:
    try:
        p = Path(file_path).resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        count = content.count(old_str)
        if count == 0:
            return {"error": "未找到 old_str，请检查空格/换行是否完全一致"}
        if count > 1:
            return {"error": f"找到 {count} 个匹配，请增加上下文使其唯一"}
        new_content = content.replace(old_str, new_str, 1)
        p.write_text(new_content, encoding="utf-8")
        return {"success": True, "message": "替换成功", "path": str(p)}
    except Exception as e:
        return {"error": str(e)}


def list_directory(path: str = ".") -> dict:
    try:
        p = Path(path).resolve()
        items = [{"name": x.name, "type": "dir" if x.is_dir() else "file"} for x in p.iterdir()]
        return {"success": True, "path": str(p), "items": sorted(items, key=lambda x: (x["type"], x["name"]))}
    except Exception as e:
        return {"error": str(e)}


def execute_command(command: str) -> dict:
    if any(f in command for f in ["rm -rf /", ":(){ :|:& };:"]):
        return {"error": "命令被安全策略拒绝"}
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=60)
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"error": "命令超时（60s）"}
    except Exception as e:
        return {"error": str(e)}


def search_files(query: str, path: str = ".", file_type: str = None) -> dict:
    try:
        results = []
        for fp in Path(path).resolve().rglob("*"):
            if not fp.is_file():
                continue
            if query.lower() not in fp.name.lower():
                continue
            if file_type and fp.suffix != file_type:
                continue
            results.append(str(fp))
            if len(results) >= 20:
                break
        return {"success": True, "files": results}
    except Exception as e:
        return {"error": str(e)}


def web_search(query: str) -> dict:
    try:
        import requests
        r = requests.get("http://127.0.0.1:8080/search",
                         params={"q": query, "format": "json"}, timeout=8)
        data = r.json()
        results = [
            {"title": x.get("title"), "url": x.get("url"), "snippet": x.get("content", "")[:200]}
            for x in data.get("results", [])[:5]
        ]
        return {"success": True, "results": results}
    except Exception as e:
        return {"error": f"搜索失败（需要本地 SearXNG）: {e}"}


# ─────────────────────────────────────────────────────────────
# 扩展工具（来自 claude_terminal / tool_schemas）
# ─────────────────────────────────────────────────────────────

def append_file_content(file_path: str, content: str) -> dict:
    """向文件末尾追加内容"""
    try:
        p = Path(file_path).resolve()
        mode = 'a' if p.exists() else 'w'
        prefix = ""
        if mode == 'a' and p.stat().st_size > 0:
            with open(p, 'rb') as f:
                f.seek(-1, 2)
                if f.read(1) != b'\n':
                    prefix = "\n"
        with open(p, mode, encoding='utf-8') as f:
            f.write(prefix + content)
        return {"success": True, "path": str(p), "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def grep_search(pattern: str, path: str = ".", include_extension: str = None, recursive: bool = True) -> dict:
    """使用正则搜索文件内容（递归）"""
    try:
        search_path = Path(path).resolve()
        results = []
        count = 0
        max_results = 20
        glob_pattern = '**/*' if recursive else '*'
        for fp in search_path.glob(glob_pattern):
            if not fp.is_file():
                continue
            if include_extension and fp.suffix != include_extension:
                continue
            if fp.stat().st_size > 1024 * 1024:
                continue
            try:
                with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()
                matches = list(re.finditer(pattern, file_content, re.MULTILINE))
                if matches:
                    file_matches = []
                    for m in matches[:5]:
                        line_num = file_content.count('\n', 0, m.start()) + 1
                        line_start = file_content.rfind('\n', 0, m.start()) + 1
                        line_end = file_content.find('\n', m.end())
                        if line_end == -1:
                            line_end = len(file_content)
                        line_content = file_content[line_start:line_end].strip()
                        file_matches.append(f"Line {line_num}: {line_content[:100]}")
                    results.append(f"File: {fp.relative_to(search_path)}\n" + "\n".join(file_matches))
                    count += 1
                    if count >= max_results:
                        break
            except Exception:
                continue
        return {"success": True, "matches_found": count, "results": "\n\n".join(results) if results else "No matches found"}
    except Exception as e:
        return {"error": str(e)}


def execute_batch_commands(commands: list, stop_on_error: bool = True) -> dict:
    """批量顺序执行多条 Shell 命令，默认遇错停止"""
    results = []
    overall_success = True
    for cmd in commands:
        res = execute_command(cmd)
        results.append({"command": cmd, "result": res})
        if not res.get("success", False):
            overall_success = False
            if stop_on_error:
                break
    return {"success": overall_success, "executed_count": len(results), "results": results}


def fetch_url(url: str) -> dict:
    """抓取 URL 页面并返回纯文本内容"""
    try:
        import requests
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return {"error": "beautifulsoup4 未安装: pip install beautifulsoup4"}
        headers = {'User-Agent': 'Mozilla/5.0 (compatible; STM32Agent/1.0)'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return {"success": True, "url": url, "content": text[:5000], "truncated": len(text) > 5000}
    except Exception as e:
        return {"error": f"获取失败: {e}"}


def get_current_time() -> dict:
    """获取当前系统时间、星期和时区"""
    try:
        now = datetime.now()
        return {
            "success": True,
            "current_time": now.strftime("%Y-%m-%d %H:%M:%S"),
            "weekday": now.strftime("%A"),
            "timezone": str(now.astimezone().tzinfo),
        }
    except Exception as e:
        return {"error": str(e)}


def ask_human(question: str) -> dict:
    """向用户提问并等待输入"""
    try:
        CONSOLE.print(f"\n[cyan][❓ AI Question]: {question}[/]")
        answer = input(" > ")
        return {"success": True, "answer": answer}
    except Exception as e:
        return {"error": str(e)}


def git_status() -> dict:
    """执行 git status 查看修改状态"""
    return execute_command("git status")


def git_diff() -> dict:
    """执行 git diff 查看代码变更"""
    return execute_command("git diff")


def git_commit(message: str) -> dict:
    """执行 git commit -m <message>"""
    return execute_command(f"git commit -m {shlex.quote(message)}")


def edit_file_lines(file_path: str, operation: str, start_line: int,
                    end_line: int = None, new_content: str = None) -> dict:
    """基于行号编辑文件（replace/insert/delete）"""
    try:
        p = Path(file_path).resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        total = len(lines)
        if start_line < 1 or start_line > total:
            return {"error": f"start_line {start_line} 超出范围 [1, {total}]"}
        if end_line is None:
            end_line = start_line
        if end_line < start_line or end_line > total:
            return {"error": f"end_line {end_line} 无效"}
        si, ei = start_line - 1, end_line
        if operation == "replace":
            if new_content is None:
                return {"error": "replace 需要 new_content"}
            if not new_content.endswith('\n'):
                new_content += '\n'
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[ei:]
        elif operation == "insert":
            if new_content is None:
                return {"error": "insert 需要 new_content"}
            if not new_content.endswith('\n'):
                new_content += '\n'
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[si:]
        elif operation == "delete":
            new_lines = lines[:si] + lines[ei:]
        else:
            return {"error": f"未知操作: {operation}"}
        with open(p, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
        return {"success": True, "path": str(p), "operation": operation, "new_total_lines": len(new_lines)}
    except Exception as e:
        return {"error": str(e)}


def insert_content_by_regex(file_path: str, regex_pattern: str, content: str) -> dict:
    """在文件第一个正则匹配位置之后插入内容"""
    try:
        p = Path(file_path).resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, 'r', encoding='utf-8', errors='ignore') as f:
            file_content = f.read()
        m = re.search(regex_pattern, file_content, re.MULTILINE)
        if not m:
            return {"error": f"正则 '{regex_pattern}' 未匹配到内容"}
        new_content = file_content[:m.end()] + content + file_content[m.end():]
        with open(p, 'w', encoding='utf-8') as f:
            f.write(new_content)
        return {"success": True, "path": str(p), "match_found": m.group(0)[:50], "message": "内容已插入"}
    except Exception as e:
        return {"error": str(e)}


def check_python_code(file_path: str) -> dict:
    """检查 Python 文件语法和风格（flake8 / ast）"""
    import ast
    try:
        p = Path(file_path).resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        try:
            with open(p, 'r', encoding='utf-8') as f:
                ast.parse(f.read())
        except SyntaxError as e:
            return {"success": False, "error_type": "SyntaxError", "line": e.lineno, "message": str(e)}
        lint_result = ""
        try:
            result = subprocess.run(
                f"flake8 {shlex.quote(str(p))}", shell=True,
                capture_output=True, text=True, timeout=10
            )
            if result.returncode != 0 and result.stdout:
                lint_result = f"Flake8:\n{result.stdout}"
        except Exception:
            pass
        return {"success": True, "message": "语法检查通过", "linter_output": lint_result or "无问题"}
    except Exception as e:
        return {"error": str(e)}


def run_python_code(code: str) -> dict:
    """执行 Python 代码片段（临时文件沙箱）"""
    import tempfile
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as tmp:
            tmp.write(code)
            tmp_path = tmp.name
        result = subprocess.run(
            [sys.executable, tmp_path], capture_output=True, text=True, timeout=30
        )
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except Exception as e:
        return {"error": str(e)}


# ── Word 文档工具 ──────────────────────────────────────────────

def _get_docx_module():
    """懒加载 python-docx，未安装时返回 None"""
    try:
        import docx
        return docx
    except ImportError:
        return None


def read_docx(file_path: str) -> dict:
    """读取 Word 文档(.docx)的文本内容"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        text = '\n'.join(p.text for p in doc.paragraphs)
        return {"success": True, "content": text, "total_paragraphs": len(doc.paragraphs)}
    except Exception as e:
        return {"error": str(e)}


def replace_docx_text(file_path: str, old_text: str, new_text: str, use_regex: bool = False) -> dict:
    """替换 Word 文档中的文本（支持正则）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        count = 0
        for para in doc.paragraphs:
            if use_regex:
                if re.search(old_text, para.text):
                    replaced = re.sub(old_text, new_text, para.text)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = replaced
                    else:
                        para.add_run(replaced)
                    count += 1
            else:
                if old_text in para.text:
                    replaced_in_run = False
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
                            replaced_in_run = True
                    if not replaced_in_run:
                        para.text = para.text.replace(old_text, new_text)
                        count += 1
        doc.save(file_path)
        return {"success": True, "replaced_count": count, "message": f"已替换 {count} 处"}
    except Exception as e:
        return {"error": str(e)}


def append_docx_content(file_path: str, content: str,
                        after_paragraph_index: int = None, style: str = None) -> dict:
    """向 Word 文档追加内容（支持指定位置插入）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        paragraphs_text = [t for t in content.split('\n') if t.strip()]
        if after_paragraph_index is None:
            for p_text in paragraphs_text:
                p = doc.add_paragraph(p_text)
                if style:
                    try:
                        p.style = style
                    except Exception:
                        pass
        else:
            n = len(doc.paragraphs)
            if after_paragraph_index < 0 or after_paragraph_index >= n:
                return {"error": f"索引 {after_paragraph_index} 超出范围 (0-{n-1})"}
            if after_paragraph_index == n - 1:
                for p_text in paragraphs_text:
                    p = doc.add_paragraph(p_text)
                    if style:
                        try:
                            p.style = style
                        except Exception:
                            pass
            else:
                next_para = doc.paragraphs[after_paragraph_index + 1]
                base_style = doc.paragraphs[after_paragraph_index].style
                for p_text in paragraphs_text:
                    new_p = next_para.insert_paragraph_before(p_text)
                    if style:
                        try:
                            new_p.style = style
                        except Exception:
                            pass
                    else:
                        new_p.style = base_style
        doc.save(file_path)
        return {"success": True, "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def inspect_docx_structure(file_path: str, max_paragraphs: int = 50) -> dict:
    """查看 Word 文档段落结构（用于定位插入点）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        structure = []
        for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            if not preview.strip():
                preview = "[空段落]"
            structure.append(f"[{i}] {preview}")
        return {"success": True, "total_paragraphs": len(doc.paragraphs), "structure": "\n".join(structure)}
    except Exception as e:
        return {"error": str(e)}


def insert_docx_content_after_heading(file_path: str, heading_text: str,
                                      content: str, style: str = None) -> dict:
    """在 Word 文档指定标题后插入内容（大小写不敏感）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        target_para = None
        for para in doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                target_para = para
                break
        if not target_para:
            return {"error": f"未找到标题: {heading_text}"}
        index = doc.paragraphs.index(target_para)
        return append_docx_content(file_path, content, index, style)
    except Exception as e:
        return {"error": str(e)}


# ── 电脑控制工具 ──────────────────────────────────────────────

def computer_screenshot() -> dict:
    """截取当前桌面截图并保存为 PNG"""
    try:
        import pyautogui
        ts = int(time.time())
        path = os.path.abspath(f"screenshot_{ts}.png")
        pyautogui.screenshot(path)
        return {"success": True, "image_path": path, "message": f"截图已保存: {path}"}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_move(x: int, y: int) -> dict:
    """移动鼠标到指定坐标"""
    try:
        import pyautogui
        pyautogui.moveTo(x, y)
        return {"success": True, "action": "move", "x": x, "y": y}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_click(button: str = "left") -> dict:
    """鼠标点击（left/right/double）"""
    try:
        import pyautogui
        if button == "double":
            pyautogui.doubleClick()
        else:
            pyautogui.click(button=button)
        return {"success": True, "button": button}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_keyboard_type(text: str) -> dict:
    """向焦点窗口输入文本"""
    try:
        import pyautogui
        pyautogui.write(text)
        return {"success": True, "typed": text}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


# ─────────────────────────────────────────────────────────────
# 工具注册表
# ─────────────────────────────────────────────────────────────
TOOLS_MAP: Dict[str, Any] = {
    # STM32 专属
    "stm32_list_probes":     stm32_list_probes,
    "stm32_connect":          stm32_connect,
    "stm32_disconnect":       stm32_disconnect,
    "stm32_serial_connect":   stm32_serial_connect,
    "stm32_serial_disconnect": stm32_serial_disconnect,
    "stm32_set_chip":         stm32_set_chip,
    "stm32_hardware_status":  stm32_hardware_status,
    "stm32_compile":          stm32_compile,
    "stm32_compile_rtos":     stm32_compile_rtos,
    "stm32_flash":            stm32_flash,
    "stm32_read_registers":   stm32_read_registers,
    "stm32_analyze_fault":    stm32_analyze_fault,
    "stm32_serial_read":      stm32_serial_read,
    "stm32_auto_flash_cycle": stm32_auto_flash_cycle,
    "stm32_reset_debug_attempts": stm32_reset_debug_attempts,
    "stm32_generate_font":    stm32_generate_font,
    "stm32_save_code":        stm32_save_code,
    "stm32_list_projects":    stm32_list_projects,
    "stm32_read_project":     stm32_read_project,
    # 通用
    "read_file":              read_file,
    "create_or_overwrite_file": create_or_overwrite_file,
    "str_replace_edit":       str_replace_edit,
    "list_directory":         list_directory,
    "execute_command":        execute_command,
    "search_files":           search_files,
    "web_search":             web_search,
    # 扩展通用工具
    "append_file_content":    append_file_content,
    "grep_search":            grep_search,
    "execute_batch_commands": execute_batch_commands,
    "fetch_url":              fetch_url,
    "get_current_time":       get_current_time,
    "ask_human":              ask_human,
    "git_status":             git_status,
    "git_diff":               git_diff,
    "git_commit":             git_commit,
    "edit_file_lines":        edit_file_lines,
    "insert_content_by_regex": insert_content_by_regex,
    "check_python_code":      check_python_code,
    "run_python_code":        run_python_code,
    # Word 文档工具
    "read_docx":              read_docx,
    "replace_docx_text":      replace_docx_text,
    "append_docx_content":    append_docx_content,
    "inspect_docx_structure": inspect_docx_structure,
    "insert_docx_content_after_heading": insert_docx_content_after_heading,
    # 电脑控制工具
    "computer_screenshot":    computer_screenshot,
    "computer_mouse_move":    computer_mouse_move,
    "computer_mouse_click":   computer_mouse_click,
    "computer_keyboard_type": computer_keyboard_type,
}
TOOLS_MAP.update(EXTRA_TOOLS_MAP)
TOOLS_MAP.update(SKILL_TOOLS_MAP)

# ─────────────────────────────────────────────────────────────
# Tool Schemas（供 AI 调用）
# ─────────────────────────────────────────────────────────────
TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "stm32_list_probes",
            "description": "列出所有已连接的调试探针（ST-Link、CMSIS-DAP、J-Link）。连接前先调用此函数确认探针存在。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_connect",
            "description": "连接 STM32 硬件（pyocd 探针 + UART 串口）。烧录前必须先调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chip": {"type": "string", "description": "芯片型号，如 STM32F103C8T6（可选，不填用当前设置）"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_disconnect",
            "description": "断开探针和串口连接。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_connect",
            "description": (
                "单独连接/重连 UART 串口监控（不影响 pyocd 探针）。"
                "串口用于接收 Debug_Print 日志和 Gary:BOOT 启动标记，是 AI 判断程序运行状态的关键。"
                "stm32_connect 会自动尝试用默认端口连接，若失败或需要更换端口时调用此函数。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "port": {"type": "string",
                             "description": "串口设备路径，如 /dev/ttyUSB0、/dev/ttyAMA0（不填用 config.py 默认值）"},
                    "baud": {"type": "integer",
                             "description": "波特率，默认 115200"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_disconnect",
            "description": "断开串口（保留探针连接）。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_set_chip",
            "description": "切换目标芯片型号，同时更新寄存器地址表。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chip": {"type": "string", "description": "芯片完整型号，如 STM32F103C8T6 / STM32F407VET6"},
                },
                "required": ["chip"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_hardware_status",
            "description": "查询当前硬件状态：芯片型号、探针/串口连接状态、GCC 版本、HAL 库是否就绪。开始工作前建议先调用。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_compile",
            "description": "使用 arm-none-eabi-gcc + HAL 库编译完整的 main.c 代码，返回编译结果和 bin 路径。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "完整的 main.c 代码（含所有 #include 和函数定义）"},
                    "chip": {"type": "string", "description": "可选：临时指定芯片型号"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_compile_rtos",
            "description": "使用 arm-none-eabi-gcc + HAL + FreeRTOS Kernel 编译带 RTOS 的完整 main.c 代码。"
                           "仅在用户需要 FreeRTOS 多任务时使用，裸机项目使用 stm32_compile。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "完整的 main.c 代码（含 FreeRTOS 头文件和任务定义）"},
                    "chip": {"type": "string", "description": "可选：临时指定芯片型号"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_flash",
            "description": "通过 pyocd 将已编译的固件烧录到 STM32。需要先 connect 和 compile。",
            "parameters": {
                "type": "object",
                "properties": {
                    "bin_path": {"type": "string", "description": "可选：bin 文件路径，不填则用上次编译结果"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_read_registers",
            "description": "读取 STM32 关键寄存器（RCC、GPIO ODR/IDR、TIM、UART、I2C 等）。只在 stm32_auto_flash_cycle 未返回寄存器数据时补充调用一次，禁止循环重复调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "regs": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "可选：指定寄存器名称列表，如 [\"RCC_APB2ENR\", \"GPIOA_CRL\"]。不填读取所有调试寄存器。",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_analyze_fault",
            "description": "读取并分析 HardFault 状态寄存器（SCB_CFSR/HFSR/BFAR + PC），定位故障原因。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_read",
            "description": "读取 UART 串口输出（调试日志、错误信息）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "timeout": {"type": "number", "description": "读取超时（秒），默认 3.0"},
                    "wait_for": {"type": "string", "description": "可选：等待直到出现此字符串（如 Gary:）"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_auto_flash_cycle",
            "description": (
                "完整开发闭环（推荐）：编译 → 烧录 → 读串口 → 读寄存器，一步到位。"
                "代码生成后直接调用此函数，获取完整的验证结果。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "完整的 main.c 代码"},
                    "request": {"type": "string", "description": "需求描述（用于保存项目，可选）"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_reset_debug_attempts",
            "description": (
                "重置调试轮次计数器。必须调用的场景：(1)全新需求 (2)用户要求修改功能/显示内容/引脚/逻辑。"
                "不调用的场景：仅在修复上一轮的编译错误/烧录失败/运行异常时继续重试。"
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_generate_font",
            "description": (
                "将任意文字（含中文）渲染为真实字模 C 数组，使用系统字体精确生成，"
                "固定格式：横向取模·高位在前（row-major MSB=left），并附带配套的 OLED_ShowFontN() 显示函数。"
                "返回的 c_code 包含字模数组 + 显示函数，直接粘贴进 main.c 使用，无需修改。"
                "显示中文/特殊字符时必须先调用此工具，禁止手写字模数据。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "要生成字模的文字，如 '你好世界'"},
                    "size": {"type": "integer", "description": "字体大小（像素），默认 16，常用 8/12/16/24/32", "default": 16},
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_save_code",
            "description": "将代码保存到项目目录（不编译，仅保存源码）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "main.c 代码内容"},
                    "request": {"type": "string", "description": "项目描述（作为目录名）"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_list_projects",
            "description": "列出最近 15 个历史项目（名称、芯片、需求描述、时间）。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_read_project",
            "description": "读取指定历史项目的 main.c 源码，用于查看或修改。",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_name": {"type": "string", "description": "项目目录名（从 stm32_list_projects 获取）"},
                },
                "required": ["project_name"],
            },
        },
    },
    # 通用工具
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "读取文件内容（带行号）。",
            "parameters": {
                "type": "object",
                "properties": {"file_path": {"type": "string"}},
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_or_overwrite_file",
            "description": "创建或完全覆盖一个文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string"},
                    "content": {"type": "string"},
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "str_replace_edit",
            "description": "精确替换文件中的字符串（old_str 必须在文件中唯一，包含 3-5 行上下文）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string"},
                    "old_str": {"type": "string", "description": "要替换的原文（必须唯一，包含足够上下文）"},
                    "new_str": {"type": "string", "description": "替换后的新文"},
                },
                "required": ["file_path", "old_str", "new_str"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": "列出目录内容。",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string"}},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_command",
            "description": "执行 Shell 命令（如查看日志、安装包、运行脚本等）。",
            "parameters": {
                "type": "object",
                "properties": {"command": {"type": "string"}},
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_files",
            "description": "按文件名关键字搜索文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string"},
                    "path": {"type": "string", "description": "搜索目录（默认当前）"},
                    "file_type": {"type": "string", "description": "扩展名过滤，如 .c / .h"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "网络搜索（需要本地 SearXNG 实例）。",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string"}},
                "required": ["query"],
            },
        },
    },
    # ── 扩展通用工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "append_file_content",
            "description": "向文件末尾追加内容（代码/文本文件）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "目标文件路径"},
                    "content": {"type": "string", "description": "要追加的内容"},
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "grep_search",
            "description": "在文件中递归搜索正则表达式模式，返回匹配位置和内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "正则表达式模式"},
                    "path": {"type": "string", "description": "搜索目录（默认 .）"},
                    "include_extension": {"type": "string", "description": "按扩展名过滤，如 .py"},
                    "recursive": {"type": "boolean", "description": "是否递归（默认 True）"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_batch_commands",
            "description": "批量顺序执行多条 Shell 命令，默认遇错停止。",
            "parameters": {
                "type": "object",
                "properties": {
                    "commands": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "命令列表",
                    },
                    "stop_on_error": {"type": "boolean", "description": "遇错是否停止（默认 True）"},
                },
                "required": ["commands"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "抓取 URL 页面并返回纯文本内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "要获取的 URL"},
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "获取当前系统时间、星期和时区。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_human",
            "description": "向用户提问，等待用户在终端输入回答。",
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "要问用户的问题"},
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_status",
            "description": "执行 git status，查看当前仓库的文件修改状态。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_diff",
            "description": "执行 git diff，查看实际代码变更内容。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_commit",
            "description": "执行 git commit -m <message> 提交变更。",
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {"type": "string", "description": "提交信息"},
                },
                "required": ["message"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file_lines",
            "description": (
                "基于行号编辑文件（优先使用 str_replace_edit）。\n"
                "操作类型：replace（替换行范围）、insert（在行前插入）、delete（删除行范围）。\n"
                "仅在无法用 str_replace_edit 时使用（如在空白处插入新代码）。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "文件路径"},
                    "operation": {
                        "type": "string",
                        "enum": ["replace", "insert", "delete"],
                        "description": "操作类型",
                    },
                    "start_line": {"type": "integer", "description": "起始行（1-indexed）"},
                    "end_line": {"type": "integer", "description": "结束行（可选，默认等于 start_line）"},
                    "new_content": {"type": "string", "description": "新内容（replace/insert 时必填）"},
                },
                "required": ["file_path", "operation", "start_line"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "insert_content_by_regex",
            "description": "在文件中第一个正则匹配位置之后插入内容，适合向类/函数后添加新方法。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "文件路径"},
                    "regex_pattern": {"type": "string", "description": "用于定位插入点的正则表达式"},
                    "content": {"type": "string", "description": "要插入的内容"},
                },
                "required": ["file_path", "regex_pattern", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "check_python_code",
            "description": "检查 Python 文件的语法错误和代码风格（flake8/ast）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Python 文件路径（.py）"},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_python_code",
            "description": "执行 Python 代码片段（临时文件沙箱），用于验证逻辑或测试库。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "要执行的 Python 代码"},
                },
                "required": ["code"],
            },
        },
    },
    # ── Word 文档工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "读取 Word 文档(.docx)的文本内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "replace_docx_text",
            "description": "替换 Word 文档中的文本（尽量保留格式，支持正则）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "old_text": {"type": "string", "description": "要查找的文本或正则模式"},
                    "new_text": {"type": "string", "description": "替换为的新文本"},
                    "use_regex": {"type": "boolean", "description": "是否启用正则匹配（默认 False）"},
                },
                "required": ["file_path", "old_text", "new_text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "append_docx_content",
            "description": "向 Word 文档追加内容，可追加到末尾或指定段落索引之后。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "content": {"type": "string", "description": "要追加的内容（\\n 分隔多段落）"},
                    "after_paragraph_index": {"type": "integer", "description": "插入位置段落索引（不填则追加到末尾）"},
                    "style": {"type": "string", "description": "Word 样式名，如 'Heading 1'、'Normal'（可选）"},
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "inspect_docx_structure",
            "description": "查看 Word 文档段落结构（索引和内容预览），用于确定插入位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "max_paragraphs": {"type": "integer", "description": "最多显示段落数（默认 50）"},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "insert_docx_content_after_heading",
            "description": "在 Word 文档指定标题段落之后插入内容（大小写不敏感匹配）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "heading_text": {"type": "string", "description": "目标标题文本（大小写不敏感）"},
                    "content": {"type": "string", "description": "要插入的内容"},
                    "style": {"type": "string", "description": "Word 样式（可选）"},
                },
                "required": ["file_path", "heading_text", "content"],
            },
        },
    },
    # ── 电脑控制工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "computer_screenshot",
            "description": "截取当前桌面截图，返回保存路径。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_mouse_move",
            "description": "移动鼠标到指定坐标（x, y）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "x": {"type": "integer", "description": "X 坐标"},
                    "y": {"type": "integer", "description": "Y 坐标"},
                },
                "required": ["x", "y"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_mouse_click",
            "description": "在当前鼠标位置点击（left/right/double）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "button": {
                        "type": "string",
                        "enum": ["left", "right", "double"],
                        "description": "点击类型（默认 left）",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_keyboard_type",
            "description": "向当前焦点窗口输入文本。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "要输入的文本"},
                },
                "required": ["text"],
            },
        },
    },
]
TOOL_SCHEMAS.extend(EXTRA_TOOL_SCHEMAS)
TOOL_SCHEMAS.extend(SKILL_TOOL_SCHEMAS)
_skill_mgr = init_skills(TOOLS_MAP, TOOL_SCHEMAS)

# ─────────────────────────────────────────────────────────────
# STM32 系统提示词
STM32_SYSTEM_PROMPT = """你是 Gary Dev Agent，专为 STM32 嵌入式开发设计的 AI 助手，深度集成了编译、烧录、调试工具链。

## 核心能力
1. **代码生成**：根据自然语言需求生成完整可编译的 STM32 HAL C 代码
2. **编译验证**：调用 arm-none-eabi-gcc 编译，立即发现并修复错误
3. **固件烧录**：通过 pyocd 将固件烧录到 STM32（支持 ST-Link / CMSIS-DAP / J-Link）
4. **硬件调试**：读取外设寄存器，分析 HardFault，监控 UART 日志
5. **代码修改**：对话式增量修改，保留已有逻辑

## 标准工作流

### 串口监控（AI 判断程序运行状态的唯一来源）
- 串口 = STM32 UART TX → USB-TTL 适配器 → 主机 `/dev/ttyUSBx` 或 `/dev/ttyAMAx`
- `stm32_hardware_status` 返回 `serial_connected: false` 时，**必须提醒用户连接串口**
- 用户可用 `/serial /dev/ttyUSB0` 连接，或告诉 AI 调用 `stm32_serial_connect(port=...)`
- 无串口时 AI 无法看到 `Gary:BOOT`、`Debug_Print` 输出和运行时错误，调试能力严重受限
- 烧录成功但无串口时，在回复末尾加一句：`⚠️ 串口未连接，无法监控运行状态`

### 全新代码生成 / 功能修改
1. 调用 `stm32_reset_debug_attempts` — **以下情况必须调用**：全新需求、修改功能/引脚/内容/逻辑；仅在修复上轮编译/烧录/运行错误时跳过
2. 调用 `stm32_hardware_status` — 了解当前芯片和工具链状态，**检查 serial_connected**
3. 生成完整 main.c（见代码规范）
4. - 代码直接作为参数传入，不需要在对话里额外展示
   - **禁止**只把代码输出在文本里而不调用此工具
5. 读取工具返回值中的关键字段：
   - `success: true` → 读 `steps` 中 `step=registers` 的 `key_regs`，通过寄存器值向用户说明验证结果
   - `give_up: true` → **立即停止**，告知用户已达上限，建议手动排查硬件
   - `hw_missing` 字段存在 → **这是硬件未接/接线错误，不是代码 bug！立即停止修改代码**，将 hw_missing 列表完整告知用户，说明哪个总线/设备有问题，让用户检查接线后重试
   - `success: false, give_up: false` → 根据 `steps` 中的错误修复代码，再次调用（**不要重置计数器**）
   - 若 key_regs 为空，最多补充调用**一次** `stm32_read_registers`，之后无论结果如何直接向用户汇报
6. 寄存器解读规则（**必须向用户说明验证结果**）：
   - GPIO 输出验证：`GPIOA_ODR` 的 bit N 为 1 → PA[N] 已拉高；bit N 为 0 → 已拉低
     例：`GPIOA_ODR=0x00000001` → PA0=HIGH ✓；`GPIOA_ODR=0x00000000` → PA0=LOW ✗
   - GPIO 模式验证（F1）：`GPIOA_CRL` 每 4 bit 控制一个引脚，bit[1:0]=11→输出，bit[3:2]=00→推挽
   - GPIO 模式验证（F4/F7/H7）：`GPIOA_MODER` 每 2 bit 控制一个引脚，01→输出，00→输入
   - RCC 时钟验证：`RCC_APB2ENR` bit2=1→GPIOA 时钟已开；bit3=1→GPIOB 时钟已开
   - 有 HardFault：`SCB_CFSR != 0` → 调用 `stm32_analyze_fault` 分析
7. 修复方向：
   - `compile_errors` 非空 → 修复编译错误
   - `has_hardfault: true` → 调用 `stm32_analyze_fault`，根据 CFSR 修复
   - `boot_ok: false` → 程序未启动，检查 SysTick_Handler 和 UART 初始化
   - 寄存器值不符预期（如 ODR bit 未置位）→ 检查 RCC 时钟是否开启、GPIO 模式配置是否正确

### 增量修改（最重要！）
用户对上一次代码提出修改要求时（如"改成共阳"、"加一个按键"）：
1. 先通过对话上下文或 `stm32_read_project` 获取**上一次完整代码**
2. **只修改用户要求的部分**，其余逻辑原封不动
3. 例：上次是跑马灯共阴 → 用户说"改共阳" → 只改电平逻辑，不重写整个程序
4. 若用户需求与上次完全无关，才从头生成

### 修改历史项目
1. `stm32_list_projects` → `stm32_read_project(name)` 读取源码
2. `str_replace_edit` 精确替换（old_str 必须在文件中唯一，含3-5行上下文）

## STM32 代码规范（严格遵守）

### 必须包含
- 完整 `#include`（stm32xxx_hal.h 及各外设头文件）
- `SystemClock_Config()` — **只用 HSI 内部时钟，禁止 HSE**；根据 chip 型号正确配置 PLL 倍频/分频/Flash 等待周期/APB 分频
- `SysTick_Handler` — **必须定义，否则 HAL_Delay 永久阻塞：**
  ```c
  void SysTick_Handler(void) { HAL_IncTick(); }
  ```

### main() 函数结构（**严格按此顺序，不可调换**）
```c
int main(void) {
    HAL_Init();
    SystemClock_Config();
    // 1. 最先初始化 UART（仅配置 GPIO 和 USART，不涉及外部设备）
    MX_USART1_UART_Init();
    // 2. 紧接着打印启动标记——此时其他外设都还没初始化
    Debug_Print("Gary:BOOT\\r\\n");
    // 3. 然后初始化其他外设（I2C、SPI、TIM、OLED 等）
    MX_I2C1_Init();  // OLED
    MX_I2C2_Init();  // 传感器
    // 4. 检测外部传感器是否在线（必须有超时，不可阻塞）
    if (HAL_I2C_IsDeviceReady(&hi2c2, SENSOR_ADDR<<1, 3, 200) != HAL_OK) {
        Debug_Print("ERR: Sensor not found\\r\\n");
        // 有 OLED 时在屏幕显示错误
    }
    // 5. 主循环
    while (1) { ... }
}
```
**关键**：`Debug_Print("Gary:BOOT")` 必须紧跟 UART 初始化，在 I2C/SPI/TIM 等一切初始化**之前**。
若 I2C 初始化卡死（传感器未接导致总线锁死），至少串口已经打印了启动标志，AI 能正确判断"程序已启动但外设有问题"。
- 轻量调试函数（**不得用 sprintf**，手写整数转字符串）：
  ```c
  void Debug_Print(const char* s) {
      HAL_UART_Transmit(&huartX, (uint8_t*)s, strlen(s), 100);
  }
  void Debug_PrintInt(const char* prefix, int val) {
      // 手写：除法取位 + '0' 偏移，或查表
      char buf[16]; int i = 0, neg = 0;
      if (val < 0) { neg = 1; val = -val; }
      if (val == 0) { buf[i++] = '0'; }
      else { while (val) { buf[i++] = '0' + val % 10; val /= 10; } }
      if (neg) buf[i++] = '-';
      // 反转后发送
      HAL_UART_Transmit(&huartX, (uint8_t*)prefix, strlen(prefix), 100);
      for (int j = i-1; j >= 0; j--) HAL_UART_Transmit(&huartX, (uint8_t*)&buf[j], 1, 100);
      HAL_UART_Transmit(&huartX, (uint8_t*)"\\r\\n", 2, 100);
  }
  ```
- 每个关键外设（I2C、SPI、ADC 等）初始化后检查返回值：
  ```c
  if (HAL_I2C_Init(&hi2c1) != HAL_OK) { Debug_Print("ERR: I2C Init Fail\\r\\n"); }
  ```
- **I2C 传感器必须检测设备是否在线**（不能假设已连接）：
  ```c
  if (HAL_I2C_IsDeviceReady(&hi2c2, SENSOR_ADDR << 1, 3, 100) != HAL_OK) {
      Debug_Print("ERR: Sensor not found\\r\\n");
      // 若有 OLED，显示错误信息并停留：
      OLED_ShowString(0, 0, "Sensor Error");
      while (1) { HAL_Delay(500); }
  }
  ```
  传感器地址用 7-bit 值（代码中左移1位），不确定地址时查数据手册
- **读取传感器数据必须检查每次 HAL 调用返回值**：
  ```c
  if (HAL_I2C_Mem_Read(...) != HAL_OK) { Debug_Print("ERR: Read fail\\r\\n"); continue; }
  ```
- 业务逻辑中出现超时/异常时用 `Debug_PrintInt` 打印错误状态码

### 显示文字/OLED 字模规则
- **必须**先调用 `stm32_generate_font(text="你好世界", size=16)` 获取真实渲染字模
- 将返回的 `c_code` 原样粘贴进代码，**禁止手写或修改字模数据**
- 出现乱码 = 字模数据错误，重新调用 `stm32_generate_font` 生成，不要猜

### 严格禁止（链接必然失败）
- `sprintf / printf / snprintf / sscanf` — 触发 `_sbrk` / `end` 未定义链接错误
- `malloc / calloc / free` — 无堆管理，链接报 `sbrk`
- `float` 格式化输出 — 用整数×10 替代（253 = 25.3°C）

### 引脚复用注意
- PA13/PA14 = SWD，PA15/PB3/PB4 = JTAG
- STM32F1 若复用这些引脚作 GPIO，必须先：`__HAL_AFIO_REMAP_SWJ_NOJTAG()`（保留 SWD）
- STM32F4+ 通过 GPIO AF 配置即可，无需 AFIO
- 重映射必须在 GPIO_Init 之前完成

### GPIO 模式速查
- 输出：`OUTPUT_PP`；PWM：`AF_PP`；ADC：`ANALOG`
- I2C：`AF_OD`（F1）或 `AF_PP`（F4+）；按键：`INPUT + PULLUP/PULLDOWN`

## 常见硬件知识

### 数码管
- 型号 `xx61AS` = 共阳极（段码低有效，位选低有效）
- 型号 `xx61BS` = 共阴极（段码高有效，位选高有效）
- 用户说"共阳"：段码取反（0亮1灭），位选低有效；"共阴"反之
- 动态扫描：每位显示 2-5ms，逐位轮流；用户未说明时在回复中注明假设

### 蜂鸣器
- 有源蜂鸣器：GPIO 高/低电平直接驱动，**不需要 PWM**
- 无源蜂鸣器：需要 PWM 方波，频率决定音调

### I2C
- 必须检查返回值，失败不阻塞
- `SR1 bit10 (AF)` = 无应答，检查设备地址和接线
- `SR2 bit1 (BUSY)` = 总线锁死，需软件复位：先 Deinit 再 Init

## 调试诊断

### 编译失败
- `undefined reference to _sbrk/end` → 用了 sprintf/printf/malloc，换手写函数
- `undefined reference to _init` → 链接脚本问题，不修改代码
- `undefined reference to HAL_xxx` → 缺 HAL 源文件或 #include

### HardFault（读 SCB_CFSR 分析）
- `PRECISERR (bit9)` → 访问了未使能时钟的外设，补 `__HAL_RCC_xxx_CLK_ENABLE()`
- `IACCVIOL (bit0)` → 函数指针/跳转地址非法
- `UNDEFINSTR (bit16)` → Thumb/ARM 模式混乱
- 配合 `PC` 寄存器定位出错位置

### 程序卡死（无 HardFault）
- **首要怀疑**：缺少 `SysTick_Handler`，`HAL_Delay()` 永远不返回
- PC 指向 `Default_Handler`（死循环 `b .`）→ 某中断未定义处理函数

### 外设不工作（无 HardFault）
- 时钟：RCC_APBxENR 对应位为 0 → 补 CLK_ENABLE
- GPIO：F1 看 CRL/CRH（4 位/引脚），F4+ 看 MODER/AFR（2 位/引脚）
- 定时器：CR1 bit0=0 → 未启动；CCER 通道位=0 → 输出未使能；检查 PSC/ARR
- UART：BRR 值是否匹配目标波特率 × 总线时钟
- I2C：见上方 SR1/SR2 分析

### 利用串口日志定位问题
- 每轮修复后仔细阅读工具返回的 `uart_output` 字段
- 通过上一轮埋入的 `Debug_Print`/`Debug_PrintInt` 精准定位逻辑 bug

### 代码缓存与精准增量修改（极其重要）
每次你调用 `stm32_compile` 后，客户端都会自动将代码缓存到本地文件：`~/.stm32agent/workspace/projects下面`。
当用户要求在已有代码基础上修改（如修改引脚、增加逻辑）时，**绝对禁止重写全部代码**！必须按以下闭环操作：
1. 思考要替换的代码片段。
2. 调用 `str_replace_edit` 工具：
   - `file_path` 固定为 `latest = Path.home() / ".stm32agent" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")这段代码保存`
   - `old_str` 填原代码片段（必须完全匹配）
   - `new_str` 填修改后的片段
3. 替换成功后，**必须**调用 `read_file` 读取该文件的最新内容（提取返回值中的 `raw_content`）。
4. 将读出的最新完整源码传给 `stm32_compile` 进行编译。

## PID 自动调参工作流

### 串口数据格式（必须在 PID 代码中埋入）
在 PID 控制循环中每次计算后打印（10-50ms 间隔）：
  PID:t=<毫秒>,sp=<目标值>,pv=<实际值>,out=<输出>,err=<误差>

### 调参闭环（每轮只改 PID 参数）
1. 生成含 PID 调试输出的代码 → stm32_auto_flash_cycle
2. 等 3-5 秒采集数据 → stm32_serial_read(timeout=5)
3. 分析+推荐 → stm32_pid_tune(kp, ki, kd, serial_output=...)
4. 用推荐参数修改代码 → str_replace_edit 替换 Kp/Ki/Kd
5. 重新烧录 → 回到步骤 1
6. 重复直到 diagnosis 显示 "响应质量良好"

### 其他实用工具
- 不确定 I2C 地址 → stm32_i2c_scan 生成扫描代码
- 舵机角度不对 → stm32_servo_calibrate 校准
- 引脚可能冲突 → stm32_pin_conflict 静态检查
- ADC 噪声大 → stm32_signal_capture 分析信号质量
- Flash 快满了 → stm32_memory_map 查看占用

## 回复规范
- **极度简洁**，像命令行工具一样输出，不写大段说明
- 工具调用后只说结论，**禁止**逐条解释代码逻辑、列"代码说明"章节
- 编译/烧录成功：一句话结论即可，如"编译成功，3716B，已烧录"
- 编译/烧录失败：直接说错误原因 + 修复动作，不加前缀废话
- 遇到错误直接修复，不询问"是否需要帮你修改"
- 代码用 ```c 包裹，但**不在代码后加解释**，除非用户主动问
- 中文回复，寄存器名/函数名保持英文
- 用户未说明硬件型号细节时（如共阳/共阴），只在最后一句简单注明假设

## 约束
- 最多5轮，第5轮仍失败 give_up=true
- 每轮只改必要部分
- 永远输出完整可编译 main.c
- user_message 用通俗中文
- 第1轮就要生成能编译通过的代码，不要留 TODO 或占位符
- 永远不要说你的模型型号，说明你是Gary开发的模型
- 每次烧录完成后，必须读寄存器，有问题解决,并且简要说明错在哪里，并且表示你正在修改，没有问题正常输出。
- 有问题优先使用str_replace_edit替换错误位置，而不是重新编写代码。

## STM32F411CEU6 专项说明

### 时钟配置（100 MHz，仅 HSI，禁用 HSE）
```c
void SystemClock_Config(void) {
    RCC_OscInitTypeDef osc = {0};
    RCC_ClkInitTypeDef clk = {0};
    osc.OscillatorType      = RCC_OSCILLATORTYPE_HSI;
    osc.HSIState            = RCC_HSI_ON;
    osc.HSICalibrationValue = RCC_HSICALIBRATION_DEFAULT;
    osc.PLL.PLLState        = RCC_PLL_ON;
    osc.PLL.PLLSource       = RCC_PLLSOURCE_HSI;
    osc.PLL.PLLM            = 16;   /* HSI/16 = 1 MHz VCO input */
    osc.PLL.PLLN            = 200;  /* × 200 = 200 MHz VCO */
    osc.PLL.PLLP            = RCC_PLLP_DIV2;  /* /2 = 100 MHz SYSCLK */
    osc.PLL.PLLQ            = 4;    /* USB/SDIO/RNG: 50 MHz */
    HAL_RCC_OscConfig(&osc);
    clk.ClockType      = RCC_CLOCKTYPE_HCLK | RCC_CLOCKTYPE_SYSCLK
                       | RCC_CLOCKTYPE_PCLK1 | RCC_CLOCKTYPE_PCLK2;
    clk.SYSCLKSource   = RCC_SYSCLKSOURCE_PLLCLK;
    clk.AHBCLKDivider  = RCC_SYSCLK_DIV1;   /* HCLK  = 100 MHz */
    clk.APB1CLKDivider = RCC_HCLK_DIV2;     /* APB1  =  50 MHz（上限 50） */
    clk.APB2CLKDivider = RCC_HCLK_DIV1;     /* APB2  = 100 MHz */
    HAL_RCC_ClockConfig(&clk, FLASH_LATENCY_3);  /* 100 MHz → 3WS */
}
```
**注意**：F411 最高 100 MHz（≠ F407 的 168 MHz），Flash Latency 必须是 3WS。

### UART 波特率计算（APB2 = 100 MHz）
- USART1/USART6 挂 APB2（100 MHz）；USART2 挂 APB1（50 MHz）
- BRR = fCK / baudrate，用于寄存器验证时换算

### pyocd 烧录目标名
- 连接时使用 `STM32F411CE` 或 `stm32f411ceux`

---

## FreeRTOS 开发规范

> 用户要求 RTOS / 多任务 / 任务调度时启用本节。编译改用 `stm32_compile_rtos`。

### 关键差异（vs 裸机）
| 项目 | 裸机 | FreeRTOS |
|------|------|----------|
| 编译工具 | `stm32_compile` | `stm32_compile_rtos` |
| SysTick | 自定义 `SysTick_Handler` | **禁止** 自定义（FreeRTOS 已接管） |
| HAL 时基 | SysTick 直接 | `vApplicationTickHook` 内调用 `HAL_IncTick()` |
| 延时 | `HAL_Delay(ms)` | `vTaskDelay(pdMS_TO_TICKS(ms))` |
| 全局变量共享 | 直接访问 | 必须用 mutex / queue 保护 |

### FreeRTOS Kernel 未下载时的处理
- `stm32_compile_rtos` 会返回错误 "FreeRTOS 内核未下载"
- 告知用户运行：`python setup.py --rtos`

### main.c 模板（FreeRTOS + HAL）
```c
#include "stm32f4xx_hal.h"
#include "FreeRTOS.h"
#include "task.h"
#include "queue.h"
#include "semphr.h"

/* ── UART ──────────────────────────────────────── */
UART_HandleTypeDef huart1;
void MX_USART1_UART_Init(void) { /* ... */ }
void Debug_Print(const char *s) {
    HAL_UART_Transmit(&huart1, (uint8_t*)s, strlen(s), 100);
}

/* ── 任务函数 ───────────────────────────────────── */
void LED_Task(void *pvParam) {
    /* 初始化 GPIO... */
    while (1) {
        HAL_GPIO_TogglePin(GPIOC, GPIO_PIN_13);
        vTaskDelay(pdMS_TO_TICKS(500));
    }
}

/* ── FreeRTOS Hooks ─────────────────────────────── */
/* 维持 HAL_Delay / HAL_GetTick 正常工作 */
void vApplicationTickHook(void)   { HAL_IncTick(); }
void vApplicationIdleHook(void)   {}
void vApplicationStackOverflowHook(TaskHandle_t xTask, char *pcName) {
    Debug_Print("ERR:StackOvf\r\n"); while (1);
}
void vApplicationMallocFailedHook(void) {
    Debug_Print("ERR:MallocFail\r\n"); while (1);
}

/* ── main ───────────────────────────────────────── */
int main(void) {
    HAL_Init();
    SystemClock_Config();
    MX_USART1_UART_Init();
    Debug_Print("Gary:BOOT\r\n");
    /* 其他外设初始化... */

    xTaskCreate(LED_Task, "LED", 128, NULL, 1, NULL);
    vTaskStartScheduler();   /* 启动调度器，不返回 */
    while (1);
}
```

### FreeRTOS 常用 API
- 创建任务：`xTaskCreate(func, "name", stack_words, param, priority, &handle)`
- 任务延时：`vTaskDelay(pdMS_TO_TICKS(ms))` 或 `vTaskDelayUntil(&lastWake, period)`
- 创建队列：`xQueueCreate(length, sizeof(item_type))`
- 发送/接收：`xQueueSend(q, &item, 0)` / `xQueueReceive(q, &item, portMAX_DELAY)`
- ISR 中发送：`xQueueSendFromISR(q, &item, &xHigherPriorityTaskWoken)` + `portYIELD_FROM_ISR()`
- 互斥量：`xSemaphoreCreateMutex()` / `xSemaphoreTake(m, timeout)` / `xSemaphoreGive(m)`
- 二值信号量：`xSemaphoreCreateBinary()`

### 常见 RTOS 编译/运行错误
- `undefined reference to vApplicationTickHook` → 忘记定义 hook 函数
- `vApplicationMallocFailedHook` 被调用 → `configTOTAL_HEAP_SIZE` 不足，减少任务栈或任务数
- 调度器启动后 HardFault → 检查任务栈大小（`stack_words` 太小），最小建议 128 words
- `SysTick_Handler` 重复定义 → 删除自己写的 `SysTick_Handler`，改用 `vApplicationTickHook`
"""
# 追加所有 skill 的 AI 提示词
STM32_SYSTEM_PROMPT += _skill_mgr.get_all_prompt_additions()

# ─────────────────────────────────────────────────────────────
# Gary doctor — 一键诊断所有配置
# ─────────────────────────────────────────────────────────────
def run_doctor():
    """检查 AI 接口、工具链、HAL、硬件探针的完整状态"""
    CONSOLE.print()
    CONSOLE.rule(f"[bold cyan]  Gary Doctor  —  环境诊断[/]")
    CONSOLE.print()
    all_ok = True

    # ── 1. AI 接口 ──────────────────────────────────────────
    CONSOLE.print("[bold]■ AI 接口[/]")
    cur_key, cur_url, cur_model = _read_ai_config()
    placeholders = ("YOUR_API_KEY", "", "sk-YOUR")
    ai_configured = bool(cur_key and not any(cur_key.startswith(p) for p in placeholders))
    if ai_configured:
        CONSOLE.print(f"  [green]✓[/] API Key   {_mask_key(cur_key)}")
        CONSOLE.print(f"  [green]✓[/] Base URL  {cur_url}")
        CONSOLE.print(f"  [green]✓[/] Model     {cur_model}")
        # 尝试连接 AI 服务
        try:
            from openai import OpenAI as _OAI
            c = _OAI(api_key=cur_key, base_url=cur_url, timeout=8.0)
            c.models.list()
            CONSOLE.print("  [green]✓[/] API 连通性  [dim]测试通过[/]")
        except Exception as e:
            err_msg = str(e)[:80]
            # 部分服务不支持 /models，收到 4xx 也算通（至少网络通了）
            if any(code in err_msg for code in ("401", "403", "404", "400")):
                CONSOLE.print(f"  [yellow]⚠[/] API 连通性  [dim]{err_msg}[/]")
            else:
                CONSOLE.print(f"  [red]✗[/] API 连通性  [dim]{err_msg}[/]")
                CONSOLE.print("    [dim]→ 运行 Gary config 重新设置 API Key[/]")
                all_ok = False
    else:
        CONSOLE.print("  [red]✗[/] API Key 未配置")
        CONSOLE.print("    [dim]→ 运行 Gary config 配置 AI 接口[/]")
        all_ok = False
    CONSOLE.print()

    # ── 2. 编译工具链 ────────────────────────────────────────
    CONSOLE.print("[bold]■ 编译工具链[/]")
    gcc = shutil.which("arm-none-eabi-gcc")
    if gcc:
        try:
            r = subprocess.run([gcc, "--version"], capture_output=True, text=True)
            ver = r.stdout.split("\n")[0][:70]
        except Exception:
            ver = gcc
        CONSOLE.print(f"  [green]✓[/] arm-none-eabi-gcc  [dim]{ver}[/]")
    else:
        CONSOLE.print("  [red]✗[/] arm-none-eabi-gcc  未找到")
        CONSOLE.print("    [dim]→ sudo apt install gcc-arm-none-eabi  或  python3 setup.py --auto[/]")
        all_ok = False

    # HAL
    from config import WORKSPACE as _WS
    hal_dir = _WS / "hal"
    hal_found = []
    for fam in ("f0", "f1", "f3", "f4"):
        if (hal_dir / "Inc" / f"stm32{fam}xx_hal.h").exists():
            hal_found.append(f"STM32{fam.upper()}xx")
    cmsis_ok = (hal_dir / "CMSIS" / "Include" / "core_cm3.h").exists()
    if hal_found and cmsis_ok:
        CONSOLE.print(f"  [green]✓[/] HAL 库      {', '.join(hal_found)}")
        CONSOLE.print(f"  [green]✓[/] CMSIS Core")
    elif hal_found:
        CONSOLE.print(f"  [yellow]⚠[/] HAL 库      {', '.join(hal_found)}（CMSIS Core 缺失）")
        CONSOLE.print("    [dim]→ python3 setup.py --hal[/]")
        all_ok = False
    else:
        CONSOLE.print("  [yellow]⚠[/] HAL 库      未下载（仅代码生成模式）")
        CONSOLE.print("    [dim]→ python3 setup.py --hal  下载所需系列[/]")
    CONSOLE.print()

    # ── 3. Python 依赖 ───────────────────────────────────────
    CONSOLE.print("[bold]■ Python 依赖[/]")
    _required = [("openai","openai"),("rich","rich"),("prompt_toolkit","prompt_toolkit")]
    _optional = [("serial","pyserial"),("pyocd","pyocd"),("docx","python-docx"),("PIL","Pillow")]
    for imp, pkg in _required:
        try:
            __import__(imp)
            CONSOLE.print(f"  [green]✓[/] {pkg}")
        except ImportError:
            CONSOLE.print(f"  [red]✗[/] {pkg}  [dim][必须] pip install {pkg}[/]")
            all_ok = False
    for imp, pkg in _optional:
        try:
            __import__(imp)
            CONSOLE.print(f"  [green]✓[/] {pkg}  [dim](可选)[/]")
        except Exception:
            CONSOLE.print(f"  [dim]○[/] {pkg}  [dim](可选，pip install {pkg})[/]")
    CONSOLE.print()

    # ── 4. 硬件探针 ─────────────────────────────────────────
    CONSOLE.print("[bold]■ 硬件探针[/]")
    try:
        import pyocd.probe.usb_probe as _up
        probes = _up.USBProbe.get_all_connected_probes(unique_id=None, is_explicit=False)
        if probes:
            for p in probes:
                CONSOLE.print(f"  [green]✓[/] {p.description}  [dim]({p.unique_id})[/]")
        else:
            CONSOLE.print("  [yellow]⚠[/] 未检测到探针  [dim](连接 ST-Link / CMSIS-DAP 后重试)[/]")
    except Exception:
        CONSOLE.print("  [dim]○[/] pyocd 未安装，无法扫描探针")

    serial_ports = detect_serial_ports(verbose=False)
    if serial_ports:
        for p in serial_ports:
            CONSOLE.print(f"  [green]✓[/] 串口 {p}")
    else:
        CONSOLE.print("  [dim]○[/] 未检测到串口设备")
    CONSOLE.print()

    # ── 总结 ─────────────────────────────────────────────────
    if all_ok:
        CONSOLE.print("[bold green]  ✅  所有核心配置正常，Gary 已就绪！[/]")
    else:
        CONSOLE.print("[bold yellow]  ⚠  存在问题，请按上方提示修复[/]")
    CONSOLE.print()


# ─────────────────────────────────────────────────────────────
# Gary config — CLI 内 AI 接口配置向导
# ─────────────────────────────────────────────────────────────
def configure_ai_cli(agent: "STM32Agent | None" = None):
    """交互式配置 AI 接口（可在 CLI 内调用，也可独立运行）"""
    import getpass as _gp

    CONSOLE.print()
    CONSOLE.rule("[bold cyan]  配置 AI 后端接口[/]")
    CONSOLE.print()

    cur_key, cur_url, cur_model = _read_ai_config()
    placeholders = ("YOUR_API_KEY", "", "sk-YOUR")
    is_configured = bool(cur_key and not any(cur_key.startswith(p) for p in placeholders))

    if is_configured:
        CONSOLE.print(f"  [dim]当前 API Key :[/] {_mask_key(cur_key)}")
        CONSOLE.print(f"  [dim]当前 Base URL:[/] {cur_url}")
        CONSOLE.print(f"  [dim]当前 Model   :[/] {cur_model}")
        CONSOLE.print()

    # 服务商菜单
    CONSOLE.print("[bold cyan]  请选择 AI 服务提供商：[/]")
    for i, (name, url, _) in enumerate(_AI_PRESETS, 1):
        url_hint = f"  [dim]{url[:55]}[/]" if url else ""
        CONSOLE.print(f"    [yellow]{i}[/].  {name:<24}{url_hint}")
    CONSOLE.print()

    valid = [str(i) for i in range(1, len(_AI_PRESETS) + 1)]
    choice = ""
    while choice not in valid:
        try:
            choice = input(f"  输入序号 [1-{len(_AI_PRESETS)}] (回车取消): ").strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]")
            return
        if choice == "":
            CONSOLE.print("[dim]已取消[/]")
            return

    idx = int(choice) - 1
    preset_name, preset_url, preset_model = _AI_PRESETS[idx]

    # Base URL
    if preset_url:
        base_url = preset_url
        CONSOLE.print(f"  [dim]Base URL: {base_url}[/]")
    else:
        try:
            base_url = input("  Base URL: ").strip()
        except (EOFError, KeyboardInterrupt):
            base_url = cur_url

    # Model
    default_model = preset_model or cur_model or ""
    try:
        hint = f" (默认 {default_model})" if default_model else ""
        entered = input(f"  Model 名称{hint}: ").strip()
        model = entered if entered else default_model
    except (EOFError, KeyboardInterrupt):
        model = default_model

    # API Key
    CONSOLE.print()
    if preset_name == "Ollama (本地)":
        api_key = "ollama"
        CONSOLE.print("  [dim]Ollama 本地模式，API Key 自动设为 ollama[/]")
    else:
        CONSOLE.print(f"  [dim]请输入 {preset_name} API Key（不显示输入内容）[/]")
        try:
            api_key = _gp.getpass("  API Key: ")
        except Exception:
            try:
                api_key = input("  API Key: ").strip()
            except (EOFError, KeyboardInterrupt):
                api_key = ""
        if not api_key:
            if is_configured:
                CONSOLE.print("  [dim]未输入，保留原有 Key[/]")
                api_key = cur_key
            else:
                CONSOLE.print("[yellow]  未输入 API Key，配置取消[/]")
                return

    # 写入 config.py
    if _write_ai_config(api_key, base_url, model):
        _reload_ai_globals()
        CONSOLE.print()
        CONSOLE.print("[green]  ✓ 配置已保存到 config.py[/]")
        CONSOLE.print(f"  [green]✓[/] 服务商  {preset_name}")
        CONSOLE.print(f"  [green]✓[/] API Key {_mask_key(api_key)}")
        CONSOLE.print(f"  [green]✓[/] Model   {model}")
        # 重建当前会话的 AI 客户端
        if agent is not None:
            from openai import OpenAI as _OAI
            agent.client = _OAI(api_key=api_key, base_url=base_url, timeout=180.0)
            CONSOLE.print("  [green]✓[/] AI 客户端已热重载，无需重启")
    else:
        CONSOLE.print("[red]  ✗ 写入 config.py 失败[/]")
    CONSOLE.print()


# ─────────────────────────────────────────────────────────────
# STM32 Agent（TUI + 流式对话 + 工具框架）
# ─────────────────────────────────────────────────────────────
class STM32Agent:
    def __init__(self):
        self.messages: List[Dict] = [{"role": "system", "content": STM32_SYSTEM_PROMPT}]
        self.client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL, timeout=180.0)
        self.session = PromptSession(
            history=InMemoryHistory(),
            complete_while_typing=False,
            enable_history_search=True,
        )
        os.environ.setdefault("no_proxy", "localhost,127.0.0.1")
        os.environ.setdefault("NO_PROXY", "localhost,127.0.0.1")

    # ── Token 估算 ──────────────────────────────────────────
    # 替换原来的 _tokens 和 _truncate 方法
    def _tokens(self) -> int:
        return sum(len(str(m.get("content", ""))) // 3 for m in self.messages)

    def _truncate_result(self, s: str, tool_name: str = "") -> str:
        """针对不同工具结果使用不同截断策略"""
        if len(s) <= MAX_TOOL_RESULT_LEN:
            return s
        half = MAX_TOOL_RESULT_LEN // 2
        # 编译/串口结果：错误在末尾，保留末尾更多
        if tool_name in ("stm32_compile", "stm32_serial_read", "stm32_auto_flash_cycle"):
            head = MAX_TOOL_RESULT_LEN // 4
            tail = MAX_TOOL_RESULT_LEN - head
            return s[:head] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-tail:]
        return s[:half] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-half:]

    def _truncate_history(self):
        """滑动窗口：保留 system prompt + 最近消息，总字符不超限"""
        MAX_CHARS = 180_000
        total = sum(len(str(m.get("content", ""))) for m in self.messages)
        removed = 0
        while total > MAX_CHARS and len(self.messages) > 3:
            # 始终保留 messages[0]（system prompt）
            victim = self.messages.pop(1)
            victim_len = len(str(victim.get("content", "")))
            total -= victim_len
            removed += 1
        if removed:
            CONSOLE.print(f"[dim]  📦 历史压缩：移除 {removed} 条旧消息[/]")
   
    # 原来是直接传 self.messages，改为过滤后传
    def _messages_for_api(self) -> list:
        """发送给 API 前处理消息格式：
        - 若对话中出现过 reasoning_content（thinking 模式），则所有 assistant 消息都必须带该字段
        - 否则过滤掉该字段（避免不支持的 API 报错）
        """
        # 检测当前会话是否启用了 thinking 模式
        has_thinking = any(
            "reasoning_content" in m
            for m in self.messages
            if m.get("role") == "assistant"
        )

        result = []
        for m in self.messages:
            if m.get("role") == "assistant" and has_thinking:
                # thinking 模式：确保每条 assistant 消息都有 reasoning_content
                clean = dict(m)
                if "reasoning_content" not in clean:
                    clean["reasoning_content"] = ""
                result.append(clean)
            else:
                # 非 thinking 模式：过滤掉该字段
                clean = {k: v for k, v in m.items() if k != "reasoning_content"}
                result.append(clean)
        return result
   
    # ── 流式响应 + 工具调用 ─────────────────────────────────
    def chat(self, user_input: str):
        self._truncate_history()
        self.messages.append({"role": "user", "content": user_input})

        while True:
            # API 调用
            try:
                stream = self.client.chat.completions.create(
                    model=AI_MODEL,
                    messages=self._messages_for_api(),
                    tools=TOOL_SCHEMAS,
                    tool_choice="auto",
                    temperature=AI_TEMPERATURE,
                    stream=True,
                )
            except Exception as e:
                CONSOLE.print(f"\n[red]API 错误: {e}[/]")
                return

            # 收集流式输出
            content = ""
            tool_calls_raw: Dict[int, dict] = {}
            thinking = ""
            in_think = False

            try:
                for chunk in stream:
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta

                    # Reasoning (deepseek-r1 style)
                    rc = getattr(delta, "reasoning_content", None)
                    if rc:
                        if not in_think:
                            CONSOLE.print(f"\n[dim {THEME}]💭 思考:[/]")
                            in_think = True
                        thinking += rc
                        CONSOLE.print(rc, end="", style="dim")

                    # 文本内容
                    if delta.content:
                        if in_think:
                            CONSOLE.print()
                            in_think = False
                        content += delta.content
                        CONSOLE.print(delta.content, end="", style="white")

                    # 工具调用
                    if delta.tool_calls:
                        for tc in delta.tool_calls:
                            idx = tc.index
                            if idx not in tool_calls_raw:
                                tool_calls_raw[idx] = {"id": "", "name": "", "args": ""}
                            if tc.id:
                                tool_calls_raw[idx]["id"] = tc.id
                            if tc.function and tc.function.name:
                                tool_calls_raw[idx]["name"] = tc.function.name
                            if tc.function and tc.function.arguments:
                                tool_calls_raw[idx]["args"] += tc.function.arguments

                if in_think:
                    CONSOLE.print()
                if content:
                    CONSOLE.print()

            except Exception as e:
                CONSOLE.print(f"\n[red]流式读取错误: {e}[/]")
                return

            # 无工具调用 → 结束
            if not tool_calls_raw:
                assistant_msg = {"role": "assistant", "content": content or ""}
                if thinking:  # 如果有思考内容，带上它
                    assistant_msg["reasoning_content"] = thinking
                self.messages.append(assistant_msg)
                break

            # 构造 assistant tool_calls 消息
            tool_calls_list = []
            for idx in sorted(tool_calls_raw.keys()):
                tc = tool_calls_raw[idx]
                tool_calls_list.append({
                    "id": tc["id"],
                    "type": "function",
                    "function": {"name": tc["name"], "arguments": tc["args"]},
                })
            assistant_tool_msg = {
                "role": "assistant",
                "content": content or "",
                "tool_calls": tool_calls_list,
            }
            if thinking:  # 关键修复：把之前收集的 thinking 塞进来
                assistant_tool_msg["reasoning_content"] = thinking
                
            self.messages.append(assistant_tool_msg)

            # 执行工具
            tool_results = []
            for tc in tool_calls_list:
                func_name = tc["function"]["name"]
                args_str = tc["function"]["arguments"]

                CONSOLE.print(f"[dim]  🔧 {func_name}[/]", end="")
                try:
                    args = json.loads(args_str) if args_str.strip() else {}
                    if func_name in TOOLS_MAP:
                        result = TOOLS_MAP[func_name](**args)
                        result_str = json.dumps(result, ensure_ascii=False, indent=2)
                    else:
                        result_str = f'{{"error": "工具不存在: {func_name}"}}'
                except Exception as e:
                    result_str = f'{{"error": "{e}"}}'

                # 简短预览
                preview = result_str[:120].replace("\n", " ")
                CONSOLE.print(f" → [dim green]{preview}[/]")

                tool_results.append({
                    "role": "tool",
                    "tool_call_id": tc["id"],
                    "content": self._truncate_result(result_str, func_name),
                })

            self.messages.extend(tool_results)
            # 继续循环，把工具结果发回 AI

    # ── 内置命令处理 ────────────────────────────────────────
    def handle_builtin(self, cmd: str) -> bool:
        """处理 /xxx 命令，返回 True 表示已处理"""
        parts = cmd.strip().split(None, 1)
        head = parts[0].lower()
        arg = parts[1] if len(parts) > 1 else ""

        if head in ("/help", "?"):
            self._show_help()
            return True

        if head == "/connect":
            chip = arg.strip() or None
            CONSOLE.print(f"\n[{THEME}]连接硬件...[/]")
            r = stm32_connect(chip)
            CONSOLE.print(f"[{'green' if r['success'] else 'red'}]{r['message']}[/]\n")
            return True

        if head == "/disconnect":
            r = stm32_disconnect()
            CONSOLE.print(f"[{THEME}]{r['message']}[/]\n")
            return True

        if head == "/skill":
            handle_skill_command(arg, agent=self)
            return True

        if head == "/serial":
            # /serial              → 自动检测并连接
            # /serial list         → 列出可用串口
            # /serial /dev/ttyUSB0 → 指定端口
            # /serial /dev/ttyUSB0 9600 → 指定端口+波特率
            tokens = arg.split()
            if tokens and tokens[0] == "list":
                ports = detect_serial_ports()
                if ports:
                    CONSOLE.print(f"[green]  可用串口:[/]")
                    for p in ports:
                        try:
                            import serial.tools.list_ports as lp
                            infos = {i.device: i.description for i in lp.comports()}
                            desc = infos.get(p, "")
                        except Exception:
                            desc = ""
                        CONSOLE.print(f"    {p}  {desc}")
                else:
                    CONSOLE.print("[yellow]  未检测到可用串口[/]")
                CONSOLE.print()
                return True
            port = tokens[0] if tokens and tokens[0].startswith("/dev/") else None
            baud = None
            for t in tokens:
                if t.isdigit():
                    baud = int(t)
                    break
            r = stm32_serial_connect(port, baud)
            color = "green" if r["success"] else "red"
            CONSOLE.print(f"[{color}]{r['message']}[/]\n")
            return True

        if head == "/chip":
            if not arg:
                CONSOLE.print(f"[{THEME}]当前芯片: {_current_chip}[/]\n")
            else:
                r = stm32_set_chip(arg)
                CONSOLE.print(f"[{THEME}]已切换: {r['chip']} ({r['family']})[/]\n")
            return True

        if head == "/status":
            s = stm32_hardware_status()
            table = Table(box=box.SIMPLE, show_header=False)
            table.add_column(style=f"bold {THEME}")
            table.add_column(style="white")
            for k, v in s.items():
                table.add_row(k, str(v))
            CONSOLE.print(table)
            CONSOLE.print()
            return True

        if head == "/probes":
            probes = stm32_list_probes()
            if probes["probes"]:
                for p in probes["probes"]:
                    CONSOLE.print(f"  [{THEME}]{p['description']}[/] ({p['uid']})")
            else:
                CONSOLE.print(f"[yellow]{probes.get('message', '未找到探针')}[/]")
            CONSOLE.print()
            return True

        if head == "/projects":
            r = stm32_list_projects()
            if r["projects"]:
                table = Table(title="历史项目", box=box.SIMPLE)
                table.add_column("项目名", style=f"bold {THEME}")
                table.add_column("芯片", style="cyan")
                table.add_column("描述", style="white")
                for p in r["projects"]:
                    table.add_row(p["name"], p["chip"], p["request"][:40])
                CONSOLE.print(table)
            else:
                CONSOLE.print("[dim]暂无历史项目[/]")
            CONSOLE.print()
            return True

        if head == "/clear":
            self.messages = [{"role": "system", "content": STM32_SYSTEM_PROMPT}]
            CONSOLE.clear()
            self._print_header()
            return True

        if head == "/config":
            configure_ai_cli(agent=self)
            return True

        if head in ("/exit", "/quit"):
            CONSOLE.print(f"\n[{THEME}]断开硬件...[/]")
            stm32_disconnect()
            CONSOLE.print("Goodbye!")
            sys.exit(0)

        return False

    # ── UI ──────────────────────────────────────────────────
    def _print_header(self):
        chip_line = f"芯片: [bold]{_current_chip}[/]  |  模型: [bold]{AI_MODEL}[/]"
        hw_line = (
            f"硬件: [green]已连接[/]  串口: [{'green' if _serial_connected else 'yellow'}]"
            f"{'已连接' if _serial_connected else '未连接'}[/]"
            if _hw_connected else "硬件: [dim]未连接[/]"
        )
        art = (
            "   ██████╗  █████╗ ██████╗ ██╗   ██╗\n"
            "  ██╔════╝ ██╔══██╗██╔══██╗╚██╗ ██╔╝\n"
            "  ██║  ███╗███████║██████╔╝ ╚████╔╝ \n"
            "  ██║   ██║██╔══██║██╔══██╗  ╚██╔╝  \n"
            "  ╚██████╔╝██║  ██║██║  ██║   ██║   \n"
            "   ╚═════╝ ╚═╝  ╚═╝╚═╝  ╚═╝   ╚═╝  "
        )
        panel = Panel(
            f"[bold {THEME}]{art}[/]\n\n"
            f"  {chip_line}\n  {hw_line}\n\n"
            f"  [dim]输入需求即可生成代码 · /help 查看命令 · /connect 连接硬件[/]",
            title=f"[bold {THEME}]Gary Dev Agent[/]",
            border_style=THEME,
            padding=(0, 1),
        )
        CONSOLE.print(panel)
        CONSOLE.print()

    def _show_help(self):
        table = Table(title="内置命令", box=box.SIMPLE)
        table.add_column("命令", style=f"bold {THEME}")
        table.add_column("说明", style="white")
        cmds = [
            ("/connect [芯片]",           "连接探针（如 /connect STM32F103C8T6）"),
            ("/serial [端口] [波特率]",   "连接串口（如 /serial /dev/ttyUSB0 115200）"),
            ("/disconnect",               "断开探针和串口"),
            ("/chip [型号]",      "查看/切换芯片型号"),
            ("/probes",           "列出所有可用探针"),
            ("/status",           "查看硬件+工具链状态"),
            ("/config",            "配置 AI 接口（API Key / Model / Base URL）"),
            ("/projects",         "列出历史项目"),
            ("/clear",            "清空对话历史"),
            ("/exit",             "退出"),
            ("?",                 "显示帮助"),
            ("/skill [子命令]",     "技能管理: list/install/enable/disable/create/export"),
        ]
        for cmd, desc in cmds:
            table.add_row(cmd, desc)
        CONSOLE.print(table)
        CONSOLE.print()

    def _status_bar(self):
        tokens = self._tokens()
        hw = f"[green]●[/] {_current_chip}" if _hw_connected else "[dim]○ 未连接[/]"
        CONSOLE.print(
            f"[dim]{hw}  │  {AI_MODEL}  │  context: ~{tokens} tokens[/]"
        )
        CONSOLE.rule(style="dim")

    # ── 主循环 ──────────────────────────────────────────────
    def run(self):
        CONSOLE.clear()
        self._print_header()
        pt_style = Style.from_dict({"prompt": f"bold {THEME}"})

        while True:
            try:
                self._status_bar()
                user_input = self.session.prompt(
                    HTML(f'<style color="cyan"><b>Gary > </b></style>'),
                    style=pt_style,
                )
                if not user_input.strip():
                    continue

                if user_input.startswith("/") or user_input.strip() == "?":
                    self.handle_builtin(user_input.strip())
                    continue

                self.chat(user_input)

            except KeyboardInterrupt:
                CONSOLE.print("\n[dim]Ctrl+C 中断。/exit 退出。[/]")
            except EOFError:
                stm32_disconnect()
                break


# ─────────────────────────────────────────────────────────────
# 入口
# ─────────────────────────────────────────────────────────────
def main():
    global _current_chip

    # 命令行参数
    args = sys.argv[1:]

    # ── 诊断模式：Gary doctor ────────────────────────────────
    if "--doctor" in args:
        run_doctor()
        sys.exit(0)

    # ── 配置模式：Gary config ────────────────────────────────
    if "--config" in args:
        configure_ai_cli()
        sys.exit(0)

    if "--chip" in args:
        idx = args.index("--chip")
        if idx + 1 < len(args):
            _current_chip = args[idx + 1].upper()

    # 检查依赖
    CONSOLE.print(f"[dim]检查环境...[/]")

    # GCC
    compiler = _get_compiler()
    ci = compiler.check(_current_chip)
    if ci.get("gcc"):
        CONSOLE.print(f"[green]  GCC: {ci['gcc_version']}[/]")
    else:
        CONSOLE.print(f"[yellow]  GCC: 未找到 arm-none-eabi-gcc[/]")

    if ci.get("hal"):
        CONSOLE.print(f"[green]  HAL: 已就绪[/]")
    else:
        CONSOLE.print(f"[yellow]  HAL: 未找到，请运行 setup.sh[/]")

    # pyocd
    try:
        import pyocd
        CONSOLE.print(f"[green]  pyocd: {pyocd.__version__}[/]")
    except ImportError:
        CONSOLE.print(f"[yellow]  pyocd: 未安装（pip install pyocd）[/]")

    # 串口自动扫描
    import platform as _platform, glob as _glob
    serial_candidates = detect_serial_ports(verbose=False)
    if serial_candidates:
        CONSOLE.print(f"[green]  串口: 检测到 {serial_candidates}（连接时自动选择）[/]")
    else:
        _plat = _platform.system()
        # 扫不到 → 检查是否权限问题，给出平台相关修复命令
        if _plat == "Linux":
            import grp as _grp
            no_perm = [p for pat in ["/dev/ttyUSB*", "/dev/ttyACM*", "/dev/ttyS[4-9]*"]
                       for p in _glob.glob(pat)
                       if os.path.exists(p) and not os.access(p, os.R_OK | os.W_OK)]
            if no_perm:
                try:
                    grp_name = _grp.getgrgid(os.stat(no_perm[0]).st_gid).gr_name
                except Exception:
                    grp_name = "dialout"
                CONSOLE.print(f"[yellow]  串口: 发现 {no_perm} 但权限不足[/]")
                CONSOLE.print(f"[yellow]    → sudo usermod -aG {grp_name} $USER && newgrp {grp_name}[/]")
            else:
                CONSOLE.print(f"[dim]  串口: 未检测到串口设备（连接硬件后重试）[/]")
        elif _plat == "Darwin":
            CONSOLE.print(f"[dim]  串口: 未检测到串口设备[/]")
            CONSOLE.print(f"[dim]    插上 USB 转串口后重启程序，或运行 /serial list 查看[/]")
        elif _plat == "Windows":
            CONSOLE.print(f"[dim]  串口: 未检测到 COM 口[/]")
            CONSOLE.print(f"[dim]    请在设备管理器确认驱动已安装（CH340/CP210x），或运行 /serial list[/]")
        else:
            CONSOLE.print(f"[dim]  串口: 未检测到串口设备（连接硬件后重试）[/]")

    CONSOLE.print()

    # 单次执行模式：Gary do "任务"  →  python stm32_agent.py --do "任务"
    if "--do" in args:
        idx = args.index("--do")
        task = args[idx + 1] if idx + 1 < len(args) else ""
        if not task:
            CONSOLE.print("[red]--do 后需要任务描述，例如: Gary do \"让 PA0 LED 闪烁\"[/]")
            sys.exit(1)
        if "--connect" in args:
            chip_arg = None
            if "--chip" in args:
                ci_idx = args.index("--chip")
                if ci_idx + 1 < len(args):
                    chip_arg = args[ci_idx + 1]
            stm32_connect(chip_arg)
        agent = STM32Agent()
        CONSOLE.print(f"\n[cyan]  ▶ Gary do: {task}[/]\n")
        agent.chat(task)
        stm32_disconnect()
        sys.exit(0)

    # 自动连接
    if "--connect" in args:
        chip_arg = None
        if "--chip" in args:
            idx = args.index("--chip")
            if idx + 1 < len(args):
                chip_arg = args[idx + 1]
        stm32_connect(chip_arg)

    agent = STM32Agent()
    agent.run()


if __name__ == "__main__":
    main()
