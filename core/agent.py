"""Gary core runtime, agent loop, and CLI orchestration."""

import argparse
import sys, os, json, re, time, shutil, subprocess, threading, shlex
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Any, Callable
import requests
from gary_skills import _get_manager

# 以项目根目录作为运行锚点，避免模块迁移后相对路径失效。
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_PROJECT_ROOT))

# Flash 项目内部模块
import ai.client as _ai_client_module
import config as _cfg
import compiler as _compiler_module
from ai.client import (
    _AI_PRESETS,
    _ai_is_configured,
    _api_key_is_placeholder,
    _mask_key,
    _read_ai_config,
    _write_ai_config,
    _write_cli_language_config,
    get_ai_client,
    reload_ai_config,
    stream_chat,
)
from ai.tools import TOOL_SCHEMAS, bind_tool_implementations, dispatch_tool_call
from core.memory import (
    _append_member_memory,
    _record_success_memory,
    gary_save_member_memory,
)
from core.state import get_context
from hardware.serial_mon import (
    SerialMonitor,
    connect_serial,
    detect_serial_ports,
    disconnect_serial,
    read_serial_output,
    wait_serial_adaptive,
)
from hardware.swd import (
    PyOCDBridge,
    connect_swd,
    disconnect_swd,
    flash_via_swd,
    list_probes as list_swd_probes,
    read_registers,
)
from hardware.uart_isp import flash_via_uart
from integrations.telegram import (
    configure_telegram_integration,
    handle_telegram_command,
    start_telegram_bot,
    stop_local_telegram_bridge,
    stop_telegram_bot,
    telegram_is_configured,
    telegram_log,
)
from prompts.debug import get_debug_prompt
from prompts.member import get_member_prompt_section
from prompts.system import build_system_prompt
from tui.commands import GaryCompleter, handle_slash_command
from tui.ui import console as CONSOLE, print_banner, run_doctor, run_interactive, run_oneshot

Compiler = _compiler_module.Compiler

AI_API_KEY = _ai_client_module.AI_API_KEY
AI_BASE_URL = _ai_client_module.AI_BASE_URL
AI_MODEL = _ai_client_module.AI_MODEL
AI_TEMPERATURE = _ai_client_module.AI_TEMPERATURE

WORKSPACE = _cfg.WORKSPACE
BUILD_DIR = _cfg.BUILD_DIR
PROJECTS_DIR = _cfg.PROJECTS_DIR
DEFAULT_CHIP = _ai_client_module.DEFAULT_CHIP
DEFAULT_CLOCK = _ai_client_module.DEFAULT_CLOCK
CLI_LANGUAGE = _ai_client_module.CLI_LANGUAGE
SERIAL_PORT = _ai_client_module.SERIAL_PORT
SERIAL_BAUD = _ai_client_module.SERIAL_BAUD
POST_FLASH_DELAY = _ai_client_module.POST_FLASH_DELAY
REGISTER_READ_DELAY = _ai_client_module.REGISTER_READ_DELAY


def _parse_cli_language(value: Any, default: Optional[str] = None) -> Optional[str]:
    raw = str(value or "").strip().lower()
    if not raw:
        return default
    if raw in {"en", "eng", "english", "英文"}:
        return "en"
    if raw in {"zh", "cn", "zh-cn", "zh_cn", "chinese", "中文"}:
        return "zh"
    return None


def _normalize_cli_language(value: Any) -> str:
    return _parse_cli_language(value, default="zh") or "zh"


CLI_LANGUAGE = _normalize_cli_language(CLI_LANGUAGE)
get_context().cli_language = CLI_LANGUAGE


def _is_cli_english() -> bool:
    return get_context().cli_language == "en"


def _cli_text(zh: str, en: str) -> str:
    return en if _is_cli_english() else zh


def _sync_ai_runtime_settings(settings: Optional[dict[str, Any]] = None) -> dict[str, Any]:
    """同步 ai.client 重载后的运行时配置到当前模块。"""

    global AI_API_KEY
    global AI_BASE_URL
    global AI_MODEL
    global AI_TEMPERATURE
    global DEFAULT_CHIP
    global DEFAULT_CLOCK
    global CLI_LANGUAGE
    global SERIAL_PORT
    global SERIAL_BAUD
    global POST_FLASH_DELAY
    global REGISTER_READ_DELAY

    data = settings or reload_ai_config()
    AI_API_KEY = str(data.get("AI_API_KEY", AI_API_KEY))
    AI_BASE_URL = str(data.get("AI_BASE_URL", AI_BASE_URL))
    AI_MODEL = str(data.get("AI_MODEL", AI_MODEL))
    AI_TEMPERATURE = data.get("AI_TEMPERATURE", AI_TEMPERATURE)
    DEFAULT_CHIP = str(data.get("DEFAULT_CHIP", DEFAULT_CHIP))
    DEFAULT_CLOCK = str(data.get("DEFAULT_CLOCK", DEFAULT_CLOCK))
    CLI_LANGUAGE = _normalize_cli_language(data.get("CLI_LANGUAGE", CLI_LANGUAGE))
    SERIAL_PORT = str(data.get("SERIAL_PORT", SERIAL_PORT))
    SERIAL_BAUD = data.get("SERIAL_BAUD", SERIAL_BAUD)
    POST_FLASH_DELAY = data.get("POST_FLASH_DELAY", POST_FLASH_DELAY)
    REGISTER_READ_DELAY = data.get("REGISTER_READ_DELAY", REGISTER_READ_DELAY)
    get_context().cli_language = CLI_LANGUAGE
    return data


# ─────────────────────────────────────────────────────────────
# Gary member.md 经验库（自动记忆 + 系统提示注入）
# ─────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────
# Telegram 机器人配置 / 守护进程 / 长轮询
# ─────────────────────────────────────────────────────────────
GARY_HOME = Path.home() / ".gary"


def _ensure_gary_home():
    GARY_HOME.mkdir(parents=True, exist_ok=True)


def _shutdown_cli_runtime(stop_telegram: bool = True) -> dict:
    """退出 CLI 时统一清理资源。"""
    results = {
        "hardware": stm32_disconnect(),
        "bridge_stopped": False,
        "telegram": {"success": True, "message": "Telegram 保持运行"},
    }
    try:
        stop_local_telegram_bridge()
        results["bridge_stopped"] = True
    except Exception as e:
        results["bridge_stopped"] = False
        results["bridge_error"] = str(e)
    if stop_telegram:
        try:
            results["telegram"] = stop_telegram_bot()
        except Exception as e:
            results["telegram"] = {"success": False, "message": f"停止 Telegram 机器人失败: {e}"}
    return results


def _ensure_cli_telegram_daemon() -> Optional[dict]:
    """在 CLI 启动时按配置自动拉起 Telegram 后台机器人。"""

    ctx = get_context()
    if ctx.telegram_cli_autostart_done:
        return None
    ctx.telegram_cli_autostart_done = True

    configured = telegram_is_configured()
    ai_ready = _ai_is_configured()
    if not configured and not ai_ready:
        return None
    if configured and not ai_ready:
        return {
            "success": False,
            "message": _cli_text(
                "已检测到 Telegram 配置，但 AI 接口未配置，未自动启动",
                "Telegram is configured, but AI is not configured, so it was not started automatically",
            ),
        }
    if not configured:
        return None

    result = start_telegram_bot()
    return {
        "success": bool(result.get("success")),
        "message": str(result.get("message", "")).strip(),
    }


# ─────────────────────────────────────────────────────────────
# UI 常量
# ─────────────────────────────────────────────────────────────
THEME = "cyan"
MAX_CONTEXT_TOKENS = 128000
MAX_TOOL_RESULT_LEN = 8000

# ─────────────────────────────────────────────────────────────
# 寄存器地址表（按系列）
# ─────────────────────────────────────────────────────────────
_REG_F1 = {
    "RCC_CR": 0x40021000,
    "RCC_CFGR": 0x40021004,
    "RCC_APB1ENR": 0x4002101C,
    "RCC_APB2ENR": 0x40021018,
    "GPIOA_CRL": 0x40010800,
    "GPIOA_CRH": 0x40010804,
    "GPIOA_IDR": 0x40010808,
    "GPIOA_ODR": 0x4001080C,
    "GPIOB_CRL": 0x40010C00,
    "GPIOB_CRH": 0x40010C04,
    "GPIOB_IDR": 0x40010C08,
    "GPIOB_ODR": 0x40010C0C,
    "GPIOC_CRL": 0x40011000,
    "GPIOC_CRH": 0x40011004,
    "TIM1_CR1": 0x40012C00,
    "TIM1_CCER": 0x40012C20,
    "TIM2_CR1": 0x40000000,
    "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400,
    "TIM3_CCER": 0x40000420,
    "ADC1_SR": 0x40012400,
    "ADC1_CR2": 0x40012408,
    "I2C1_CR1": 0x40005400,
    "I2C1_SR1": 0x40005414,
    "I2C2_CR1": 0x40005800,
    "I2C2_SR1": 0x40005814,
    "USART1_SR": 0x40013800,
    "USART1_BRR": 0x40013808,
}
_REG_F4 = {
    "RCC_CR": 0x40023800,
    "RCC_CFGR": 0x40023808,
    "RCC_AHB1ENR": 0x40023830,
    "RCC_APB1ENR": 0x40023840,
    "RCC_APB2ENR": 0x40023844,
    "GPIOA_MODER": 0x40020000,
    "GPIOA_IDR": 0x40020010,
    "GPIOA_ODR": 0x40020014,
    "GPIOB_MODER": 0x40020400,
    "GPIOB_IDR": 0x40020410,
    "GPIOB_ODR": 0x40020414,
    "GPIOC_MODER": 0x40020800,
    "GPIOC_IDR": 0x40020810,
    "GPIOC_ODR": 0x40020814,
    "TIM2_CR1": 0x40000000,
    "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400,
    "TIM3_CCER": 0x40000420,
    "I2C1_CR1": 0x40005400,
    "I2C1_SR1": 0x40005414,
    "USART1_SR": 0x40011000,
    "USART1_BRR": 0x40011008,
}
_REG_F0F3 = {
    "RCC_CR": 0x40021000,
    "RCC_CFGR": 0x40021004,
    "RCC_AHBENR": 0x40021014,
    "RCC_APB2ENR": 0x40021018,
    "RCC_APB1ENR": 0x4002101C,
    "GPIOA_MODER": 0x48000000,
    "GPIOA_IDR": 0x48000010,
    "GPIOA_ODR": 0x48000014,
    "GPIOB_MODER": 0x48000400,
    "GPIOB_IDR": 0x48000410,
    "GPIOB_ODR": 0x48000414,
    "I2C1_CR1": 0x40005400,
    "USART1_BRR": 0x40013808,
    "USART1_CR1": 0x4001380C,
}
_REG_COMMON = {
    "SCB_CFSR": 0xE000ED28,
    "SCB_HFSR": 0xE000ED2C,
    "SCB_BFAR": 0xE000ED38,
    "NVIC_ISER0": 0xE000E100,
}


def _reg_map(family: str) -> dict:
    base = {"f1": _REG_F1, "f4": _REG_F4, "f0": _REG_F0F3, "f3": _REG_F0F3}
    regs = dict(base.get(family.lower(), _REG_F1))
    regs.update(_REG_COMMON)
    return regs


# ─────────────────────────────────────────────────────────────
# PyOCDBridge / SerialMonitor 已迁移到 hardware/
# ─────────────────────────────────────────────────────────────
MAX_DEBUG_ATTEMPTS = 8  # 提高上限，一个任务最多 8 轮（含修改迭代）


def _get_compiler() -> Compiler:
    """返回 Compiler 单例；若 compiler/ 包文件已更新则自动热重载。"""
    ctx = get_context()
    import importlib

    try:
        mtime = _compiler_module.get_package_mtime()
    except Exception:
        mtime = 0.0
    if mtime > ctx.compiler_mtime:
        globals()["_compiler_module"] = importlib.reload(_compiler_module)
        if hasattr(_compiler_module, "reload_package"):
            _compiler_module.reload_package()
        globals()["Compiler"] = _compiler_module.Compiler
        ctx.compiler = None  # 旧实例作废，下方重新创建
        ctx.compiler_mtime = mtime
    if ctx.compiler is None:
        ctx.compiler = _compiler_module.Compiler()
        ctx.compiler.check(ctx.chip)  # 探测 GCC/HAL，设置 has_gcc/has_hal
    return ctx.compiler


def _get_bridge() -> PyOCDBridge:
    """Return the SWD bridge singleton stored in the runtime context."""

    ctx = get_context()
    bridge = ctx.bridge
    if not isinstance(bridge, PyOCDBridge):
        ctx.bridge = PyOCDBridge(
            console=CONSOLE,
            reg_map_factory=_reg_map,
            register_read_delay=REGISTER_READ_DELAY,
            default_chip=DEFAULT_CHIP,
        )
    else:
        bridge.configure(
            console=CONSOLE,
            reg_map_factory=_reg_map,
            register_read_delay=REGISTER_READ_DELAY,
            default_chip=DEFAULT_CHIP,
        )
    return ctx.bridge


def _get_serial() -> SerialMonitor:
    """Return the serial monitor singleton stored in the runtime context."""

    ctx = get_context()
    monitor = ctx.serial
    if not isinstance(monitor, SerialMonitor):
        ctx.serial = SerialMonitor(console=CONSOLE)
    else:
        monitor.console = CONSOLE
    return ctx.serial


# ─────────────────────────────────────────────────────────────
# STM32 专属工具实现
# ─────────────────────────────────────────────────────────────


def stm32_generate_font(text: str, size: int = 16) -> dict:
    """
    将任意文字（含中文）渲染为 STM32 OLED 用的 C 点阵数组。
    固定使用「横向取模·高位在前（row-major, MSB=left）」格式——
    这是最直观、与 SSD1306 逐行刷新最匹配的格式，配套显示函数一并生成。
    返回 c_code（字模数组 + 完整显示函数），直接粘贴进 main.c 使用。
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return {"success": False, "message": "需要安装 Pillow: pip install Pillow"}

    import platform as _plat

    # 跨平台字体候选
    def _find_cjk_font() -> Optional[str]:
        """动态查找系统 CJK 字体路径"""
        # 优先用 fc-match（Linux/macOS）
        try:
            r = subprocess.run(
                ["fc-match", "--format=%{file}", ":lang=zh"],
                capture_output=True,
                text=True,
                timeout=5,
            )
            if r.returncode == 0 and r.stdout.strip():
                path = r.stdout.strip()
                if os.path.exists(path):
                    return path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 回退到硬编码候选列表
        candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None
        # 原来的 for fp in font_candidates: ... 全部替换为：

    font_path = _find_cjk_font()
    if font_path is None:
        return {"success": False, "message": "未找到中文字体，请安装 fonts-noto-cjk"}
    try:
        font = ImageFont.truetype(font_path, size)
    except Exception as e:
        return {"success": False, "message": f"字体加载失败 ({font_path}): {e}"}

    def _render_char(char: str) -> list:
        """渲染单个字符到 size×size 位图，返回 0/1 列表（行优先）"""
        img = Image.new("L", (size, size), 0)
        draw = ImageDraw.Draw(img)
        try:
            bbox = font.getbbox(char)
            char_w = bbox[2] - bbox[0]
            char_h = bbox[3] - bbox[1]
            # 水平居中，垂直顶对齐（避免因字体 metrics 差异被裁底部）
            ox = (size - char_w) // 2 - bbox[0]
            oy = -bbox[1]  # 让字符顶部与图像顶部对齐
        except Exception:
            ox, oy = 0, 0
        draw.text((ox, oy), char, fill=255, font=font)
        return [1 if p > 127 else 0 for p in img.getdata()]

    def _to_row_msb(pixels: list) -> list:
        """横向取模·高位在前：每行从左到右，bit7=最左列"""
        data = []
        bytes_per_row = (size + 7) // 8
        for row in range(size):
            for b in range(bytes_per_row):
                byte = 0
                for bit in range(8):
                    col = b * 8 + bit
                    if col < size and pixels[row * size + col]:
                        byte |= 1 << (7 - bit)
                data.append(byte)
        return data

    def _ascii_preview(pixels: list) -> str:
        lines = []
        for row in range(size):
            lines.append("".join("█" if pixels[row * size + col] else "." for col in range(size)))
        return "\n".join(lines)

    chars_data = []
    previews = []
    char_list = []
    for char in text:
        pixels = _render_char(char)
        data = _to_row_msb(pixels)
        chars_data.append(data)
        previews.append(_ascii_preview(pixels))
        char_list.append(char)

    bytes_per_char = size * ((size + 7) // 8)
    fname = f"FONT_{size}x{size}"

    # ── 字模数组 ──────────────────────────────────────────────
    char_entries = []
    for i, (char, data) in enumerate(zip(char_list, chars_data)):
        preview_lines = previews[i].split("\n")
        preview_comment = "  /* " + "  ".join(preview_lines[:4]) + " ... */"
        hex_str = ", ".join(f"0x{b:02X}" for b in data)
        char_repr = char if ord(char) < 128 else f"{char}(U+{ord(char):04X})"
        char_entries.append(f"    /* [{i}] '{char_repr}' */\n    {{{hex_str}}}")

    array_code = (
        f"/* ═══ 字模数据：横向取模·高位在前 {size}x{size}px ═══\n"
        f"   格式：每行 {(size+7)//8} 字节，bit7=最左列，共 {bytes_per_char} 字节/字符\n"
        f"   字符表: {' '.join(repr(c) for c in char_list)} */\n"
        f"static const uint8_t {fname}[][{bytes_per_char}] = {{\n"
        + ",\n".join(char_entries)
        + "\n};\n"
    )

    # ── 配套显示函数（与字模格式严格匹配）──────────────────────
    display_func = f"""
/* ═══ 配套显示函数（必须与上面字模数据一起使用）═══ */
/* idx: 字符在 {fname} 中的下标（按字符表顺序） */
/* x,y: OLED 列(0-127)和页起始行(0-63)         */
void OLED_ShowFont{size}(uint8_t x, uint8_t y, uint8_t idx) {{
    const uint8_t *p = {fname}[idx];
    uint8_t bytes_per_row = {(size+7)//8};
    for (uint8_t row = 0; row < {size}; row++) {{
        OLED_SetCursor(x, y + row);   /* 设置到目标行 */
        for (uint8_t b = 0; b < bytes_per_row; b++) {{
            uint8_t byte = p[row * bytes_per_row + b];
            for (int8_t bit = 7; bit >= 0; bit--) {{
                uint8_t col = b * 8 + (7 - bit);
                if (col < {size}) {{
                    OLED_DrawPixel(x + col, y + row, (byte >> bit) & 1);
                }}
            }}
        }}
    }}
}}
/* 用法示例：显示字符表第0个字符在 (0,0) 位置
   OLED_ShowFont{size}(0, 0, 0);  // 显示 '{char_list[0] if char_list else "?"}' */
"""

    # ── ASCII 预览（调试用）────────────────────────────────────
    preview_block = "\n\n".join(f"/* '{c}':\n{p} */" for c, p in zip(char_list, previews))

    c_code = array_code + display_func
    return {
        "success": True,
        "c_code": c_code,
        "preview": preview_block,  # ASCII 预览，可用于肉眼验证字形
        "char_count": len(text),
        "bytes_per_char": bytes_per_char,
        "font_size": size,
        "mode": "row_msb",
        "char_order": char_list,
    }


def stm32_list_probes() -> dict:
    """列出所有可用的调试探针（ST-Link / CMSIS-DAP / J-Link）"""
    probes = list_swd_probes(console=CONSOLE)
    if not probes:
        return {"success": True, "probes": [], "message": "未检测到任何探针，请检查 USB 连接"}
    return {"success": True, "probes": probes}


def stm32_connect(chip: str = None) -> dict:
    """连接 STM32 硬件（pyocd 探针 + 串口监控）"""
    ctx = get_context()
    swd_result = connect_swd(
        ctx,
        chip,
        default_chip=DEFAULT_CHIP,
        register_read_delay=REGISTER_READ_DELAY,
        reg_map_factory=_reg_map,
        console=CONSOLE,
    )
    bridge = swd_result.get("bridge")
    if isinstance(bridge, PyOCDBridge):
        ctx.bridge = bridge
    target_chip = chip or ctx.chip or DEFAULT_CHIP
    if swd_result["success"]:
        ctx.hw_connected = True
        ctx.chip = bridge.chip_info.get("device", target_chip)
        _get_compiler().set_chip(ctx.chip)
        serial_result = connect_serial(ctx, baud=SERIAL_BAUD, console=CONSOLE)
        monitor = serial_result.get("serial")
        if isinstance(monitor, SerialMonitor):
            ctx.serial = monitor
        ctx.serial_connected = bool(serial_result.get("success"))
        return {
            "success": True,
            "chip": ctx.chip,
            "probe": bridge.chip_info.get("probe", ""),
            "serial_connected": ctx.serial_connected,
            "message": f"硬件已连接: {ctx.chip}",
        }
    ctx.hw_connected = False
    return {"success": False, "message": str(swd_result.get("message", "连接失败"))}


def stm32_serial_connect(port: str = None, baud: int = None) -> dict:
    """
    单独连接/重连 UART 串口（不影响 pyocd 探针连接）。
    用于更换串口设备或在 stm32_connect 之后补充连接串口。
    """
    ctx = get_context()
    use_baud = baud or SERIAL_BAUD
    result = connect_serial(ctx, port or None, use_baud, console=CONSOLE)
    monitor = result.get("serial")
    if isinstance(monitor, SerialMonitor):
        ctx.serial = monitor
    ctx.serial_connected = bool(result.get("success"))
    actual_port = getattr(ctx.serial, "port", None) or port or "自动检测"
    if ctx.serial_connected:
        return {
            "success": True,
            "port": actual_port,
            "baud": use_baud,
            "message": f"串口已连接: {actual_port} @ {use_baud}",
        }
    candidates = detect_serial_ports()
    return {
        "success": False,
        "port": port,
        "baud": use_baud,
        "message": f"串口打开失败，可用端口: {candidates if candidates else '无'}",
    }


def stm32_serial_disconnect() -> dict:
    """断开串口（保留 pyocd 探针连接）"""
    ctx = get_context()
    disconnect_serial(ctx)
    ctx.serial_connected = False
    return {"success": True, "message": "串口已断开"}


def stm32_disconnect() -> dict:
    """断开硬件连接（释放探针和串口）"""
    ctx = get_context()
    disconnect_swd(ctx)
    disconnect_serial(ctx)
    ctx.hw_connected = False
    ctx.serial_connected = False
    return {"success": True, "message": "已断开"}


def stm32_set_chip(chip: str) -> dict:
    """切换目标芯片型号（如 STM32F103C8T6 / STM32F407VET6）"""
    ctx = get_context()
    ctx.chip = chip.strip().upper()
    ci = _get_compiler().set_chip(ctx.chip)
    if ctx.hw_connected:
        _get_bridge().set_family(ci.get("family", "f1"))
    return {"success": True, "chip": ctx.chip, "family": ci.get("family", "f1")}


def stm32_hardware_status() -> dict:
    """获取当前硬件连接状态和工具链可用性"""
    ctx = get_context()
    ci = _get_compiler().check(ctx.chip)
    return {
        "chip": ctx.chip,
        "hw_connected": ctx.hw_connected,
        "serial_connected": ctx.serial_connected,
        "gcc_ok": ci.get("gcc", False),
        "gcc_version": ci.get("gcc_version", "未找到"),
        "hal_ok": ci.get("hal", False),
        "hal_lib_ok": ci.get("hal_lib", False),
        "workspace": str(WORKSPACE),
    }


def stm32_compile(code: str, chip: str = None) -> dict:
    """编译 STM32 C 代码（完整 main.c）"""
    ctx = get_context()
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile(code)
    if result["ok"]:
        ctx.last_code = code
        ctx.last_bin_path = result.get("bin_path")
        # 自动保存到 latest_workspace
        try:
            latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    payload = {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }
    if payload["success"]:
        _record_success_memory(
            "compile_success",
            code=code,
            result=payload,
            chip=get_context().chip,
            log_error=telegram_log,
        )
    return payload


def stm32_compile_rtos(code: str, chip: str = None) -> dict:
    """编译带 FreeRTOS 内核的完整 main.c 代码"""
    ctx = get_context()
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile_rtos(code)
    if result["ok"]:
        ctx.last_code = code
        ctx.last_bin_path = result.get("bin_path")
        try:
            latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    payload = {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }
    if payload["success"]:
        _record_success_memory(
            "compile_success",
            code=code,
            result=payload,
            chip=get_context().chip,
            log_error=telegram_log,
        )
    return payload


def stm32_recompile(mode: str = "auto") -> dict:
    """从 latest_workspace/main.c 直接重编译，无需 Gary 传递代码字符串。
    str_replace_edit 修改文件后调用此函数代替 read_file + stm32_compile。
    mode: "bare"(裸机) | "rtos"(FreeRTOS) | "auto"(自动检测，默认)
    """
    latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
    main_c = latest / "main.c"
    if not main_c.exists():
        return {"success": False, "message": "latest_workspace/main.c 不存在，请先编译一次完整代码"}
    code = main_c.read_text(encoding="utf-8")
    if mode == "auto":
        mode = (
            "rtos"
            if ("FreeRTOS.h" in code or "task.h" in code or "xTaskCreate" in code)
            else "bare"
        )
    if mode == "rtos":
        return stm32_compile_rtos(code)
    return stm32_compile(code)


def stm32_regen_bsp(chip: str = None) -> dict:
    """强制重新生成 startup.s / link.ld / FreeRTOSConfig.h。
    每次修改 compiler/ 包或切换芯片后调用一次，确保构建文件是最新版本。
    自动检查 startup.s 是否包含 FPU 使能代码（Cortex-M4 必须）。
    """
    import importlib

    # 强制重载 compiler 包，绕过进程内缓存
    globals()["_compiler_module"] = importlib.reload(_compiler_module)
    if hasattr(_compiler_module, "reload_package"):
        _compiler_module.reload_package()
    globals()["Compiler"] = _compiler_module.Compiler
    ctx = get_context()
    ctx.compiler = None
    try:
        ctx.compiler_mtime = _compiler_module.get_package_mtime()
    except Exception:
        ctx.compiler_mtime = 0.0

    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())

    # 调用 set_chip 触发 startup.s / link.ld 重新生成
    chip_name = chip.strip().upper() if chip else ctx.chip
    ci = compiler.set_chip(chip_name)
    if ci is None:
        return {"success": False, "message": f"未知芯片: {chip_name}"}

    # 验证 startup.s 包含 FPU 使能
    from config import BUILD_DIR

    startup_path = BUILD_DIR / "startup.s"
    has_fpu_enable = False
    if startup_path.exists():
        content = startup_path.read_text()
        has_fpu_enable = "Enable FPU" in content or "CPACR" in content

    cpu = ci.get("cpu", "?")
    fpu = ci.get("fpu", False)
    ram_k = ci.get("ram_k", 0)
    flash_k = ci.get("flash_k", 0)

    warnings = []
    if fpu and not has_fpu_enable:
        warnings.append("⚠ startup.s 缺少 FPU 使能代码——请更新 compiler/ 包")
    if ram_k < 32 and ci.get("family") in ("f4",):
        warnings.append(f"⚠ RAM 仅 {ram_k}KB，FreeRTOS heap 可能不够")

    return {
        "success": True,
        "chip": chip_name,
        "cpu": cpu,
        "fpu_supported": fpu,
        "startup_fpu_enabled": has_fpu_enable,
        "flash_k": flash_k,
        "ram_k": ram_k,
        "files_regenerated": ["startup.s", "link.ld"],
        "warnings": warnings,
        "message": f"BSP 文件已重新生成 ({cpu}, {'有 FPU' if fpu else '无 FPU'}，{flash_k}K Flash / {ram_k}K RAM)"
        + (("；" + "；".join(warnings)) if warnings else ""),
    }


def stm32_analyze_fault_rtos() -> dict:
    """读取并分析 FreeRTOS 程序的 HardFault 寄存器。
    在普通 stm32_analyze_fault 基础上增加 FreeRTOS 专项诊断：
    - 识别 FPU 未使能导致的 PRECISERR
    - 识别 SysTick_Handler 冲突
    - 识别任务栈溢出引发的故障
    - 检查 startup.s 是否包含 FPU 初始化
    返回明确的根本原因和修复建议。
    """
    if not get_context().hw_connected:
        return {"success": False, "message": "硬件未连接"}

    bridge = _get_bridge()
    regs = bridge.read_registers(
        ["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC", "SCB_MMFAR", "SCB_CPACR_FIELD"]
    )
    if not regs:
        regs = bridge.read_registers(["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC"])
    if not regs:
        return {"success": False, "message": "寄存器读取失败，请确认硬件已连接并已暂停"}

    cfsr = regs.get("SCB_CFSR", "0x0")
    hfsr = regs.get("SCB_HFSR", "0x0")
    bfar = regs.get("SCB_BFAR", "0x0")
    pc = regs.get("PC", "0x0")

    try:
        cfsr_v = int(cfsr, 16)
        hfsr_v = int(hfsr, 16)
        bfar_v = int(bfar, 16)
        pc_v = int(pc, 16)
    except (ValueError, TypeError):
        cfsr_v = hfsr_v = bfar_v = pc_v = 0

    # 检查 startup.s 的 FPU 使能
    from config import BUILD_DIR

    startup_path = BUILD_DIR / "startup.s"
    has_fpu_in_startup = False
    if startup_path.exists():
        txt = startup_path.read_text()
        has_fpu_in_startup = "CPACR" in txt or "Enable FPU" in txt

    # 检查 CPACR 硬件状态（读 0xE000ED88 的 bit[23:20]）
    fpu_hw_enabled = False
    try:
        cpacr = bridge._target.read32(0xE000ED88)
        fpu_hw_enabled = ((cpacr >> 20) & 0xF) == 0xF
    except Exception:
        pass

    # ── 诊断逻辑 ──────────────────────────────────────────
    root_cause = "未知"
    fix = "请读取 SCB_CFSR 后手动分析"
    severity = "unknown"

    preciserr = bool(cfsr_v & (1 << 9))
    bfarvalid = bool(cfsr_v & (1 << 15))
    iaccviol = bool(cfsr_v & (1 << 0))
    undefinstr = bool(cfsr_v & (1 << 16))
    stkovf = bool(cfsr_v & (1 << 12))  # STKOVF: 栈下溢
    unstkovf = bool(cfsr_v & (1 << 11))  # UNSTKOVF: 不稳定栈下溢

    # 判断 BFAR 是否"合理"（STM32 合法地址范围）
    bfar_valid_range = (
        0x08000000 <= bfar_v < 0x08200000
        or 0x20000000 <= bfar_v < 0x20030000
        or 0x40000000 <= bfar_v < 0x60000000
        or 0xE0000000 <= bfar_v
    )
    bfar_garbage = bfarvalid and not bfar_valid_range

    if preciserr and bfar_garbage and not fpu_hw_enabled:
        root_cause = (
            "FPU 未使能：ARM_CM4F port.c 的 PendSV_Handler 执行 vpush/vpop 时触发 PRECISERR"
        )
        fix = (
            "修复方法（已在 compiler/ 包修复，重新编译即可）：\n"
            "  1. 调用 stm32_regen_bsp() 重新生成 startup.s（含 CPACR 初始化）\n"
            "  2. 重新编译：stm32_compile_rtos(code)\n"
            "  不需要在代码里手动写 SCB->CPACR"
        )
        severity = "critical"
    elif undefinstr and not fpu_hw_enabled:
        root_cause = "FPU 指令在 FPU 禁用状态下执行（UNDEFINSTR）"
        fix = "同上，重新生成 startup.s 后重新编译"
        severity = "critical"
    elif stkovf or unstkovf:
        root_cause = "任务栈溢出（FreeRTOS 检测到栈越界）"
        fix = "增大任务的 stack_words 参数：普通任务 ≥128，FPU 任务 ≥256，含 printf 的任务 ≥384"
        severity = "high"
    elif preciserr and bfarvalid:
        root_cause = f"总线错误：精确地址 BFAR={bfar} 非法（可能是外设时钟未开启或访问了无效内存）"
        fix = "检查 PC 指向的函数，确认相关外设 RCC 时钟已 __HAL_RCC_xxx_CLK_ENABLE()"
        severity = "medium"
    elif hfsr_v & (1 << 1):
        root_cause = "向量表读取失败（VECTTBL）：中断向量地址无效"
        fix = "检查链接脚本 FLASH 地址和向量表对齐"
        severity = "critical"
    elif cfsr_v == 0 and hfsr_v == 0:
        root_cause = "无 HardFault（程序正常运行）"
        fix = "无需修复"
        severity = "none"

    # 补充 startup 检查信息
    startup_note = ""
    if not has_fpu_in_startup:
        startup_note = "⚠ 当前 startup.s 不含 FPU 使能代码，调用 stm32_regen_bsp() 后重新编译"
    elif not fpu_hw_enabled:
        startup_note = "⚠ startup.s 含 FPU 使能但硬件 CPACR 当前为 0（可能是旧固件未重新编译）"

    return {
        "success": True,
        "registers": regs,
        "fpu_hw_enabled": fpu_hw_enabled,
        "startup_has_fpu_enable": has_fpu_in_startup,
        "root_cause": root_cause,
        "severity": severity,
        "fix": fix,
        "startup_note": startup_note,
        "cfsr_bits": {
            "PRECISERR": preciserr,
            "BFARVALID": bfarvalid,
            "IACCVIOL": iaccviol,
            "UNDEFINSTR": undefinstr,
            "STKOVF": stkovf,
        },
    }


def stm32_rtos_check_code(code: str) -> dict:
    """FreeRTOS 代码静态检查 —— 编译前捕获常见 RTOS 编程错误。
    检查 SysTick 冲突、HAL_Delay 陷阱、缺少 hook 函数、栈大小、ISR 安全等。
    """
    import re

    errors = []
    warnings = []
    suggestions = []

    # 1. SysTick_Handler 冲突
    if re.search(r"\bvoid\s+SysTick_Handler\b", code):
        errors.append(
            "❌ 禁止自定义 SysTick_Handler —— FreeRTOS 已通过 xPortSysTickHandler 接管 SysTick。"
            "删除 SysTick_Handler，改用 vApplicationTickHook() 维持 HAL_IncTick()"
        )

    # 2. HAL_Delay 在 xTaskCreate 之前
    hal_delay_pos = code.find("HAL_Delay")
    xtask_pos = code.find("xTaskCreate")
    if hal_delay_pos >= 0 and xtask_pos >= 0 and hal_delay_pos < xtask_pos:
        # 排除注释中的情况（简单排除）
        line_with_delay = code[:hal_delay_pos].rfind("\n")
        delay_line = code[line_with_delay:hal_delay_pos]
        if "//" not in delay_line and "/*" not in delay_line:
            errors.append(
                "❌ HAL_Delay() 在 xTaskCreate() 之前调用 —— SysTick 会触发 FreeRTOS tick handler，"
                "访问未初始化的任务列表导致 HardFault。"
                "把含 HAL_Delay 的初始化移到任务函数内部"
            )

    # 3. 必需 hook 函数
    required_hooks = {
        "vApplicationTickHook": "void vApplicationTickHook(void) { HAL_IncTick(); }",
        "vApplicationMallocFailedHook": "void vApplicationMallocFailedHook(void) { while(1); }",
        "vApplicationStackOverflowHook": "void vApplicationStackOverflowHook(TaskHandle_t t, char *n) { while(1); }",
        "vApplicationIdleHook": "void vApplicationIdleHook(void) {}",
    }
    for hook, template in required_hooks.items():
        if hook not in code:
            errors.append(f"❌ 缺少 {hook} —— 请添加: {template}")

    # 4. 任务栈大小检查
    task_creates = re.findall(r'xTaskCreate\s*\(\s*(\w+)\s*,\s*"[^"]*"\s*,\s*(\d+)', code)
    for func_name, stack_str in task_creates:
        stack = int(stack_str)
        # 查找任务函数体
        func_pattern = rf"void\s+{re.escape(func_name)}\s*\("
        func_match = re.search(func_pattern, code)
        if func_match:
            # 提取函数体（简单：从匹配位置到后续 2000 字符）
            func_body = code[func_match.start() : func_match.start() + 2000]
            has_float = any(
                kw in func_body
                for kw in [
                    "float ",
                    "double ",
                    "sinf(",
                    "cosf(",
                    "sqrtf(",
                    "tanf(",
                    "arm_",
                    ".0f",
                    "fabsf(",
                    "powf(",
                    "logf(",
                    "expf(",
                ]
            )
            has_printf = any(kw in func_body for kw in ["snprintf(", "sprintf(", "printf("])
            if has_printf and stack < 384:
                warnings.append(f"⚠ 任务 {func_name} 使用 snprintf 但栈仅 {stack} words，建议 ≥384")
            elif has_float and stack < 256:
                warnings.append(f"⚠ 任务 {func_name} 使用浮点运算但栈仅 {stack} words，建议 ≥256")
            elif stack < 128:
                warnings.append(f"⚠ 任务 {func_name} 栈仅 {stack} words，建议 ≥128")

    # 5. ISR 安全检查 —— 在 IRQHandler 中使用非 FromISR 的 API
    irq_funcs = re.findall(r"void\s+(\w+_IRQHandler)\s*\(void\)\s*\{", code)
    for irq_name in irq_funcs:
        # 提取 IRQ 函数体
        irq_start = code.find(f"void {irq_name}")
        if irq_start >= 0:
            irq_body = code[irq_start : irq_start + 1500]
            # 检查危险 API（非 FromISR 版本）
            unsafe_apis = [
                "xQueueSend(",
                "xQueueReceive(",
                "xSemaphoreTake(",
                "xSemaphoreGive(",
                "vTaskDelay(",
                "vTaskDelayUntil(",
                "xTaskCreate(",
                "vTaskDelete(",
                "printf(",
            ]
            for api in unsafe_apis:
                if api in irq_body and api.replace("(", "FromISR(") not in irq_body:
                    safe_api = api.replace("(", "FromISR(")
                    warnings.append(f"⚠ {irq_name} 中使用了 {api} —— ISR 中必须用 {safe_api}")

    # 6. 缺少头文件
    if "strlen(" in code or "memset(" in code or "memcpy(" in code:
        if "#include <string.h>" not in code and "#include<string.h>" not in code:
            warnings.append("⚠ 使用了 strlen/memset/memcpy 但未 #include <string.h>")
    if "sinf(" in code or "cosf(" in code or "sqrtf(" in code:
        if "#include <math.h>" not in code and "#include<math.h>" not in code:
            warnings.append("⚠ 使用了 sinf/cosf/sqrtf 但未 #include <math.h>")
    if "snprintf(" in code:
        if "#include <stdio.h>" not in code and "#include<stdio.h>" not in code:
            warnings.append("⚠ 使用了 snprintf 但未 #include <stdio.h>")

    # 7. 建议
    if "xSemaphoreCreateBinary" in code and "vTaskNotifyGive" not in code:
        suggestions.append(
            "💡 考虑用任务通知 (vTaskNotifyGive/ulTaskNotifyTake) 替代二值信号量，速度更快且更省内存"
        )
    if (
        "xEventGroupCreate" not in code
        and re.findall(r"xSemaphoreCreateBinary", code).__len__() >= 3
    ):
        suggestions.append("💡 多个二值信号量可能适合用事件组 (xEventGroupCreate) 替代")
    if "vTaskDelayUntil" not in code and "vTaskDelay" in code:
        suggestions.append(
            "💡 需要精确周期执行时用 vTaskDelayUntil 替代 vTaskDelay（避免累积漂移）"
        )

    return {
        "success": True,
        "errors": errors,
        "warnings": warnings,
        "suggestions": suggestions,
        "error_count": len(errors),
        "warning_count": len(warnings),
        "message": (
            f"检查完成: {len(errors)} 个错误, {len(warnings)} 个警告, {len(suggestions)} 个建议"
            + (
                "\n" + "\n".join(errors + warnings + suggestions)
                if errors or warnings or suggestions
                else "\n✅ 代码通过所有 RTOS 检查"
            )
        ),
    }


def stm32_rtos_task_stats() -> dict:
    """通过 pyocd 读取 FreeRTOS 运行时任务统计信息。
    从 ELF 符号表解析全局变量地址，读取任务数、堆使用、当前任务等。
    需要先编译（有 ELF 文件）并连接硬件。
    """
    if not get_context().hw_connected:
        return {"success": False, "message": "硬件未连接"}

    from config import BUILD_DIR

    elf_path = BUILD_DIR / "firmware.elf"
    if not elf_path.exists():
        return {"success": False, "message": "ELF 文件不存在，请先编译"}

    bridge = _get_bridge()

    # 通过 nm 解析 ELF 符号表获取变量地址
    import subprocess

    symbols_to_find = {
        "uxCurrentNumberOfTasks": None,  # uint32_t 任务数
        "xFreeBytesRemaining": None,  # size_t heap_4 剩余
        "xMinimumEverFreeBytesRemaining": None,  # size_t 历史最低
        "pxCurrentTCB": None,  # TCB* 当前任务
    }

    try:
        r = subprocess.run(
            ["arm-none-eabi-nm", "-g", str(elf_path)], capture_output=True, text=True, timeout=5
        )
        if r.returncode != 0:
            return {"success": False, "message": "nm 解析 ELF 失败"}

        for line in r.stdout.split("\n"):
            parts = line.strip().split()
            if len(parts) >= 3:
                addr_str, _typ, name = parts[0], parts[1], parts[2]
                if name in symbols_to_find:
                    symbols_to_find[name] = int(addr_str, 16)
    except Exception as e:
        return {"success": False, "message": f"符号解析异常: {e}"}

    result = {"success": True}

    # 读取各变量值
    try:
        addr = symbols_to_find["uxCurrentNumberOfTasks"]
        if addr:
            result["task_count"] = bridge._target.read32(addr)

        addr = symbols_to_find["xFreeBytesRemaining"]
        if addr:
            result["heap_free_bytes"] = bridge._target.read32(addr)

        addr = symbols_to_find["xMinimumEverFreeBytesRemaining"]
        if addr:
            result["heap_min_ever_free"] = bridge._target.read32(addr)

        addr = symbols_to_find["pxCurrentTCB"]
        if addr:
            tcb_ptr = bridge._target.read32(addr)
            result["current_tcb_addr"] = f"0x{tcb_ptr:08X}"

            # 尝试读取当前任务名（TCB 偏移量 52 处是 pcTaskName，长度 configMAX_TASK_NAME_LEN=16）
            try:
                name_offset = 52  # 标准 FreeRTOS TCB pcTaskName 偏移
                name_bytes = bytearray()
                for i in range(16):
                    b = bridge._target.read8(tcb_ptr + name_offset + i)
                    if b == 0:
                        break
                    name_bytes.append(b)
                result["current_task_name"] = name_bytes.decode("ascii", errors="replace")
            except Exception:
                pass

    except Exception as e:
        result["read_error"] = str(e)

    # 构造摘要消息
    parts = []
    if "task_count" in result:
        parts.append(f"任务数: {result['task_count']}")
    if "heap_free_bytes" in result:
        parts.append(f"堆剩余: {result['heap_free_bytes']}B")
    if "heap_min_ever_free" in result:
        parts.append(f"堆历史最低: {result['heap_min_ever_free']}B")
    if "current_task_name" in result:
        parts.append(f"当前任务: {result['current_task_name']}")
    result["message"] = " | ".join(parts) if parts else "读取完成（部分符号未找到）"

    return result


def stm32_rtos_suggest_config(
    task_count: int, use_fpu: bool = False, use_printf: bool = False, ram_k: int = 0
) -> dict:
    """根据用户的任务需求，计算推荐的 FreeRTOS 配置参数。
    估算堆大小、栈大小、优先级数，并检查 RAM 是否足够。
    """
    compiler = _get_compiler()
    ci = compiler._chip_info
    if ci is None:
        return {"success": False, "message": "芯片未设置，请先 stm32_set_chip"}

    actual_ram_k = ram_k if ram_k > 0 else ci.get("ram_k", 64)
    has_fpu = ci.get("fpu", False)

    # 栈大小推荐
    if use_printf:
        recommended_stack = 384
        stack_reason = "含 snprintf/printf，需要 384+ words"
    elif use_fpu or has_fpu:
        recommended_stack = 256
        stack_reason = "FPU 上下文保存需要 256+ words"
    else:
        recommended_stack = 128
        stack_reason = "普通任务最小 128 words"

    # 堆大小估算
    tcb_size = 92  # TCB 约 92 字节
    stack_bytes = recommended_stack * 4  # words → bytes
    per_task = tcb_size + stack_bytes
    idle_task = per_task  # Idle 任务
    timer_task = tcb_size + (recommended_stack * 2 * 4)  # Timer 任务栈更大

    total_task_mem = per_task * task_count + idle_task + timer_task
    # 额外开销：队列/信号量/事件组 约 20%
    recommended_heap = int(total_task_mem * 1.3)
    # 对齐到 256B
    recommended_heap = ((recommended_heap + 255) // 256) * 256

    max_heap = actual_ram_k * 1024 // 2  # 最多用一半 RAM
    if recommended_heap > max_heap:
        recommended_heap = max_heap

    # RAM 使用估算
    static_overhead = 4096  # .data + .bss + ISR 栈 约 4KB
    total_ram_usage = static_overhead + recommended_heap
    ram_pct = total_ram_usage * 100 // (actual_ram_k * 1024)

    warnings = []
    if ram_pct > 85:
        warnings.append(f"⚠ RAM 使用率预估 {ram_pct}%，建议减少任务数或栈大小")
    if task_count > 7:
        warnings.append("⚠ 任务数较多，注意优先级分配避免优先级反转")
    if actual_ram_k < 16 and task_count > 2:
        warnings.append(f"⚠ RAM 仅 {actual_ram_k}KB，建议最多 2 个任务")

    config = {
        "success": True,
        "recommended_stack_words": recommended_stack,
        "stack_reason": stack_reason,
        "recommended_heap_bytes": recommended_heap,
        "recommended_priorities": min(task_count + 2, 7),
        "estimated_ram_usage_bytes": total_ram_usage,
        "estimated_ram_percent": ram_pct,
        "ram_total_kb": actual_ram_k,
        "per_task_overhead_bytes": per_task,
        "warnings": warnings,
        "message": (
            f"推荐配置: 栈={recommended_stack}words, 堆={recommended_heap}B, "
            f"优先级={min(task_count + 2, 7)}, RAM预估使用{ram_pct}%"
            + (("  " + "; ".join(warnings)) if warnings else "")
        ),
    }
    return config


def stm32_rtos_plan_project(
    description: str, peripherals: list = None, task_hints: list = None
) -> dict:
    """FreeRTOS 项目规划工具 —— 复杂 RTOS 项目的架构规划。

    根据用户需求描述，生成结构化的项目规划：
    - 任务分解（任务名、职责、栈大小、优先级）
    - 通信拓扑（任务间用什么机制通信：Queue/Semaphore/Notification/EventGroup）
    - 中断处理策略（哪些中断需要处理，如何通知任务）
    - 外设分配（哪个任务负责哪些外设）
    - 时序约束（哪些操作有实时性要求）
    - 资源估算（总堆/栈/RAM 使用）

    AI 在规划复杂 RTOS 项目时必须先调用此工具，待用户确认后再写代码。
    """
    import re

    compiler = _get_compiler()
    ci = compiler._chip_info
    if ci is None:
        return {"success": False, "message": "芯片未设置，请先 stm32_set_chip"}

    has_fpu = ci.get("fpu", False)
    ram_k = ci.get("ram_k", 64)
    flash_k = ci.get("flash_k", 256)
    cpu = ci.get("cpu", "cortex-m3")

    # ── 分析需求 ──────────────────────────────────────────
    desc_lower = description.lower()

    # 检测是否涉及浮点
    uses_float = any(
        kw in desc_lower
        for kw in [
            "浮点",
            "float",
            "sin",
            "cos",
            "pid",
            "温度计算",
            "dsp",
            "adc采样",
            "滤波",
            "fft",
            "角度",
            "加速度",
            "陀螺仪",
        ]
    )

    # 检测是否需要 printf/snprintf
    uses_printf = any(
        kw in desc_lower for kw in ["printf", "snprintf", "格式化", "打印浮点", "调试输出"]
    )

    # 检测外设
    detected_peripherals = []
    periph_map = {
        "uart": ["串口", "uart", "usart", "通信", "打印", "日志"],
        "i2c": ["i2c", "oled", "传感器", "bmp", "sht", "mpu", "加速度", "陀螺仪", "屏幕"],
        "spi": ["spi", "sd卡", "flash", "w25q", "tft", "lcd"],
        "adc": ["adc", "模拟", "采样", "电压", "温度", "光照"],
        "tim_pwm": ["pwm", "舵机", "电机", "蜂鸣器", "呼吸灯", "调光"],
        "tim_basic": ["定时器", "timer", "周期", "计时"],
        "gpio_out": ["led", "继电器", "数码管", "指示灯", "开关输出"],
        "gpio_in": ["按键", "按钮", "开关输入", "限位", "光电"],
        "exti": ["外部中断", "中断触发", "边沿检测"],
        "dma": ["dma", "高速传输"],
    }
    for periph, keywords in periph_map.items():
        if any(kw in desc_lower for kw in keywords):
            detected_peripherals.append(periph)

    if peripherals:
        for p in peripherals:
            if p.lower() not in detected_peripherals:
                detected_peripherals.append(p.lower())

    # ── 任务规划 ──────────────────────────────────────────
    planned_tasks = []

    # 根据需求自动推荐任务结构
    task_templates = {
        "sensor": {
            "triggers": [
                "传感器",
                "adc",
                "采样",
                "温度",
                "湿度",
                "加速度",
                "陀螺仪",
                "光照",
                "压力",
            ],
            "name": "SensorTask",
            "purpose": "传感器数据采集与处理",
            "priority": 3,
            "stack": 256 if uses_float else 128,
            "comm_out": "Queue（发送处理后的数据）",
        },
        "display": {
            "triggers": ["显示", "oled", "lcd", "tft", "数码管", "屏幕"],
            "name": "DisplayTask",
            "purpose": "显示刷新（从队列接收数据更新显示）",
            "priority": 1,
            "stack": 256,
            "comm_in": "Queue（接收要显示的数据）",
        },
        "control": {
            "triggers": ["控制", "电机", "舵机", "pid", "调节", "反馈"],
            "name": "ControlTask",
            "purpose": "控制算法执行（PID 等）",
            "priority": 4,
            "stack": 384 if uses_float else 256,
            "comm_in": "Queue（接收传感器数据）",
            "comm_out": "Queue/直接GPIO（输出控制信号）",
        },
        "comm": {
            "triggers": ["通信", "上位机", "蓝牙", "wifi", "发送", "协议"],
            "name": "CommTask",
            "purpose": "通信处理（串口/蓝牙数据收发）",
            "priority": 2,
            "stack": 384 if uses_printf else 256,
            "comm_in": "Queue（待发送的数据）",
        },
        "led": {
            "triggers": ["led", "指示灯", "呼吸灯", "闪烁"],
            "name": "LEDTask",
            "purpose": "LED 状态指示",
            "priority": 1,
            "stack": 128,
        },
        "button": {
            "triggers": ["按键", "按钮", "输入", "用户交互"],
            "name": "ButtonTask",
            "purpose": "按键扫描与事件分发",
            "priority": 2,
            "stack": 128,
            "comm_out": "TaskNotify/EventGroup（按键事件通知其他任务）",
        },
        "alarm": {
            "triggers": ["报警", "蜂鸣器", "警报", "阈值"],
            "name": "AlarmTask",
            "purpose": "报警判断与执行",
            "priority": 3,
            "stack": 128,
            "comm_in": "TaskNotify（由传感器任务触发）",
        },
        "log": {
            "triggers": ["日志", "记录", "sd卡", "存储"],
            "name": "LogTask",
            "purpose": "数据记录与存储",
            "priority": 1,
            "stack": 384,
            "comm_in": "Queue（待记录的数据）",
        },
    }

    if task_hints:
        # 用户提供了任务提示，直接使用
        for hint in task_hints:
            if isinstance(hint, dict):
                planned_tasks.append(hint)
            elif isinstance(hint, str):
                planned_tasks.append(
                    {
                        "name": hint,
                        "purpose": f"用户指定任务: {hint}",
                        "priority": 2,
                        "stack": 256 if uses_float else 128,
                    }
                )
    else:
        # 自动推荐
        for tmpl_name, tmpl in task_templates.items():
            if any(kw in desc_lower for kw in tmpl["triggers"]):
                task = {k: v for k, v in tmpl.items() if k != "triggers"}
                planned_tasks.append(task)

    # 如果没有匹配到任何任务，给一个默认的
    if not planned_tasks:
        planned_tasks.append(
            {
                "name": "MainTask",
                "purpose": "主要业务逻辑",
                "priority": 2,
                "stack": 256 if uses_float else 128,
            }
        )

    # ── 中断规划 ──────────────────────────────────────────
    interrupt_plan = []
    if "exti" in detected_peripherals or "中断" in desc_lower:
        interrupt_plan.append(
            {
                "irq": "EXTIx_IRQHandler",
                "strategy": "vTaskNotifyGiveFromISR → 唤醒处理任务",
                "note": "ISR 中仅做通知，数据处理在任务中完成",
            }
        )
    if "uart" in detected_peripherals:
        interrupt_plan.append(
            {
                "irq": "USARTx_IRQHandler",
                "strategy": "接收中断 → xQueueSendFromISR → CommTask 处理",
                "note": "使用 HAL_UART_Receive_IT 触发中断回调",
            }
        )
    if "tim_basic" in detected_peripherals or "tim_pwm" in detected_peripherals:
        interrupt_plan.append(
            {
                "irq": "TIMx_IRQHandler",
                "strategy": "定时器中断回调 → 设置标志或通知任务",
                "note": "周期性采样可用 vTaskDelayUntil 替代定时器中断",
            }
        )

    # ── 通信拓扑推荐 ──────────────────────────────────────
    comm_topology = []
    task_names = [t["name"] for t in planned_tasks]

    if "SensorTask" in task_names and "DisplayTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "DisplayTask",
                "mechanism": "xQueueSend/xQueueReceive",
                "data": "传感器数据结构体",
            }
        )
    if "SensorTask" in task_names and "ControlTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "ControlTask",
                "mechanism": "xQueueSend/xQueueReceive",
                "data": "传感器原始值",
            }
        )
    if "SensorTask" in task_names and "AlarmTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "AlarmTask",
                "mechanism": "xTaskNotifyGive（阈值触发时通知）",
                "data": "无（AlarmTask 自行读取共享数据）",
            }
        )
    if "ButtonTask" in task_names:
        for t in task_names:
            if t != "ButtonTask" and t != "LEDTask":
                comm_topology.append(
                    {
                        "from": "ButtonTask",
                        "to": t,
                        "mechanism": "xEventGroupSetBits",
                        "data": "按键事件位",
                    }
                )
                break  # 只连接一个示例

    if "CommTask" in task_names:
        data_sources = [t for t in task_names if t in ("SensorTask", "ControlTask", "LogTask")]
        for src in data_sources[:1]:
            comm_topology.append(
                {
                    "from": src,
                    "to": "CommTask",
                    "mechanism": "xQueueSend",
                    "data": "待发送的数据包",
                }
            )

    # ── 资源估算 ──────────────────────────────────────────
    total_stack_bytes = sum(t.get("stack", 128) * 4 for t in planned_tasks)
    tcb_overhead = len(planned_tasks) * 92
    idle_timer_overhead = 92 + 256 * 4 + 92 + 512 * 4  # Idle + Timer 任务
    queue_overhead = len(comm_topology) * 120  # 每个队列约 120B
    recommended_heap = int(
        (total_stack_bytes + tcb_overhead + idle_timer_overhead + queue_overhead) * 1.3
    )
    recommended_heap = ((recommended_heap + 255) // 256) * 256

    total_ram_est = recommended_heap + 4096  # heap + static
    ram_pct = total_ram_est * 100 // (ram_k * 1024)

    resource_check = {
        "total_tasks": len(planned_tasks) + 2,  # +Idle +Timer
        "total_stack_bytes": total_stack_bytes,
        "recommended_heap": recommended_heap,
        "estimated_ram_usage": total_ram_est,
        "ram_percent": ram_pct,
        "flash_k": flash_k,
        "ram_k": ram_k,
    }

    warnings = []
    if ram_pct > 85:
        warnings.append(
            f"⚠ RAM 使用率预估 {ram_pct}%（{total_ram_est}B / {ram_k*1024}B），考虑减少任务数或栈大小"
        )
    if ram_k < 16 and len(planned_tasks) > 2:
        warnings.append(f"⚠ RAM 仅 {ram_k}KB，{len(planned_tasks)} 个任务可能不够用")
    if not has_fpu and uses_float:
        warnings.append("⚠ 当前芯片无 FPU，浮点运算将使用软件模拟（较慢）")

    # ── 构造规划文档 ──────────────────────────────────────
    plan_text = f"📋 FreeRTOS 项目规划\n"
    plan_text += f"芯片: {ci.get('define','')} ({cpu}, {flash_k}K Flash, {ram_k}K RAM, {'FPU' if has_fpu else '无FPU'})\n"
    plan_text += f"需求: {description}\n\n"

    plan_text += "━━━ 任务规划 ━━━\n"
    for i, t in enumerate(planned_tasks, 1):
        plan_text += (
            f"  {i}. {t['name']} (优先级={t.get('priority',2)}, "
            f"栈={t.get('stack',128)}words)\n"
            f"     职责: {t.get('purpose','')}\n"
        )
        if "comm_in" in t:
            plan_text += f"     输入: {t['comm_in']}\n"
        if "comm_out" in t:
            plan_text += f"     输出: {t['comm_out']}\n"
    plan_text += f"  + Idle任务 + Timer任务 (系统自动创建)\n\n"

    if comm_topology:
        plan_text += "━━━ 通信拓扑 ━━━\n"
        for c in comm_topology:
            plan_text += f"  {c['from']} → {c['to']}: {c['mechanism']} ({c.get('data','')})\n"
        plan_text += "\n"

    if interrupt_plan:
        plan_text += "━━━ 中断策略 ━━━\n"
        for ip in interrupt_plan:
            plan_text += f"  {ip['irq']}: {ip['strategy']}\n"
            if ip.get("note"):
                plan_text += f"    💡 {ip['note']}\n"
        plan_text += "\n"

    plan_text += "━━━ 资源估算 ━━━\n"
    plan_text += f"  任务总数: {resource_check['total_tasks']} (含 Idle + Timer)\n"
    plan_text += f"  栈总量: {total_stack_bytes}B\n"
    plan_text += f"  推荐堆: {recommended_heap}B\n"
    plan_text += f"  RAM 预估: {total_ram_est}B / {ram_k*1024}B ({ram_pct}%)\n"

    if detected_peripherals:
        plan_text += f"\n━━━ 使用外设 ━━━\n"
        plan_text += f"  {', '.join(detected_peripherals)}\n"

    if warnings:
        plan_text += f"\n━━━ 警告 ━━━\n"
        for w in warnings:
            plan_text += f"  {w}\n"

    return {
        "success": True,
        "plan": {
            "tasks": planned_tasks,
            "communication": comm_topology,
            "interrupts": interrupt_plan,
            "peripherals": detected_peripherals,
            "resources": resource_check,
        },
        "warnings": warnings,
        "uses_fpu": uses_float and has_fpu,
        "uses_printf": uses_printf,
        "message": plan_text,
    }


def stm32_flash(bin_path: str = None) -> dict:
    """烧录固件到 STM32（需要先 connect + compile）"""
    ctx = get_context()
    if not ctx.hw_connected:
        return {"success": False, "message": "硬件未连接，请先调用 stm32_connect"}
    path = bin_path or ctx.last_bin_path
    if not path or not Path(path).exists():
        return {"success": False, "message": f"固件文件不存在: {path}"}
    _get_serial().clear()
    return flash_via_swd(ctx, path)


def stm32_read_registers(regs: list = None) -> dict:
    """读取 STM32 硬件寄存器（RCC、GPIO、TIM、UART 等）"""
    if not get_context().hw_connected:
        return {"success": False, "message": "硬件未连接"}
    result = read_registers(get_context(), regs, debug_all=not regs)
    if result is not None:
        return {"success": True, "registers": result}
    return {"success": False, "message": "寄存器读取失败"}


def stm32_analyze_fault() -> dict:
    """读取并分析 HardFault 寄存器（SCB_CFSR / SCB_HFSR）"""
    if not get_context().hw_connected:
        return {"success": False, "message": "硬件未连接"}
    regs = _get_bridge().read_registers(["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC"])
    if not regs:
        return {"success": False, "message": "寄存器读取失败"}
    analysis = _get_bridge().analyze_fault(regs)
    return {"success": True, "registers": regs, "analysis": analysis}


def stm32_serial_read(timeout: float = 3.0, wait_for: str = None) -> dict:
    """读取 UART 串口输出（调试日志）"""
    return read_serial_output(get_context(), timeout=timeout, wait_for=wait_for)


def stm32_reset_debug_attempts() -> dict:
    """重置调试轮次计数器。开始一个全新需求时调用，确保计数从 1 开始。"""
    get_context().debug_attempt = 0
    return {"success": True, "message": "计数器已重置"}


def stm32_auto_flash_cycle(code: str, request: str = "") -> dict:
    """
    完整开发闭环（自动计轮次，最多 MAX_DEBUG_ATTEMPTS 轮）：
      编译 → 烧录（若已连接硬件）→ 等待启动 → 读串口 → 读寄存器
    返回每步结果 + 当前轮次 + 是否应放弃。
    """
    ctx = get_context()
    ctx.debug_attempt += 1
    attempt = ctx.debug_attempt
    steps = []

    # 超出最大轮次 → 直接告知 AI 放弃
    if attempt > MAX_DEBUG_ATTEMPTS:
        return {
            "success": False,
            "give_up": True,
            "attempt": attempt,
            "message": f"已达到最大调试轮次 ({MAX_DEBUG_ATTEMPTS})，自动调试无法解决，请检查硬件接线或手动排查",
            "steps": [],
        }

    remaining = MAX_DEBUG_ATTEMPTS - attempt
    CONSOLE.print(f"[dim]  第 {attempt}/{MAX_DEBUG_ATTEMPTS} 轮[/]")

    # 1. 编译
    comp = stm32_compile(code)
    steps.append({"step": "compile", "success": comp["success"], "msg": comp["message"][:300]})
    if not comp["success"]:
        return {
            "success": False,
            "attempt": attempt,
            "remaining": remaining,
            "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
            "steps": steps,
            "error": "编译失败，请根据错误信息修改代码",
            "compile_errors": comp["message"],
        }

    # 2. 烧录（失败时延迟重试，最多2次）
    if ctx.hw_connected and comp.get("bin_path"):
        fr = stm32_flash(comp["bin_path"])
        if not fr["success"]:
            for retry in range(2):
                wait_sec = 1.5 * (retry + 1)
                CONSOLE.print(
                    f"[yellow]  烧录失败（{fr['message'][:60]}），{wait_sec:.0f}s 后重连重试...[/]"
                )
                time.sleep(wait_sec)
                stm32_connect(ctx.chip)
                fr = stm32_flash(comp["bin_path"])
                if fr["success"]:
                    break
        steps.append({"step": "flash", "success": fr["success"], "msg": fr["message"]})
        if not fr["success"]:
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps,
                "error": f"烧录失败（重试2次后仍失败）: {fr['message']}",
            }
        # 3. 串口监控
        uart_out = ""
        sensor_errors = []
        if ctx.serial_connected:
            CONSOLE.print("[dim]  等待启动...[/]")
            uart_out = wait_serial_adaptive(
                _get_serial(),
                keyword="Gary:BOOT",
                min_wait=0.5,
                max_wait=POST_FLASH_DELAY + 4.0,
            )
            boot_ok = "Gary:BOOT" in uart_out
            # 检测传感器错误关键词
            sensor_errors = [
                line.strip()
                for line in uart_out.splitlines()
                if "ERR:" in line or "Error" in line or "not found" in line.lower()
            ]
            # 打印串口输出到终端供用户查看
            if uart_out.strip():
                CONSOLE.print(f"[dim]  串口输出:[/]\n[cyan]{uart_out.strip()[:400]}[/]")
            else:
                CONSOLE.print("[yellow]  串口无输出（程序未启动或卡死）[/]")
            steps.append(
                {
                    "step": "uart",
                    "output": uart_out[:500],
                    "boot_ok": boot_ok,
                    "sensor_errors": sensor_errors,
                }
            )
        else:
            time.sleep(POST_FLASH_DELAY + 1.0)
            boot_ok = True  # 无串口跳过验证

        # 4. 读寄存器（打印到终端）
        regs = _get_bridge().read_all_for_debug()
        has_fault = False
        if regs:
            KEY_SET = (
                "SCB_CFSR",
                "PC",
                "RCC_APB2ENR",
                "RCC_APB1ENR",
                "GPIOA_CRL",
                "GPIOA_CRH",
                "GPIOA_ODR",
                "GPIOA_IDR",
                "GPIOB_CRL",
                "GPIOB_CRH",
                "GPIOB_ODR",
                "GPIOB_IDR",
                "GPIOC_CRL",
                "GPIOC_CRH",
                "GPIOC_ODR",
                "GPIOC_IDR",
                "GPIOA_MODER",
                "GPIOB_MODER",
                "GPIOC_MODER",
                "TIM2_CR1",
                "I2C1_CR1",
                "I2C1_SR1",
                "I2C2_CR1",
                "I2C2_SR1",
            )
            key_regs = {k: v for k, v in regs.items() if k in KEY_SET}
            has_fault = regs.get("SCB_CFSR", "0x00000000") not in ("0x00000000", "0x0")
            # 打印寄存器到终端
            CONSOLE.print("[dim]  关键寄存器:[/]")
            for k, v in key_regs.items():
                color = "red" if (k == "SCB_CFSR" and has_fault) else "dim"
                CONSOLE.print(f"[{color}]    {k} = {v}[/]")
            steps.append(
                {
                    "step": "registers",
                    "key_regs": key_regs,
                    "has_hardfault": has_fault,
                    "fault_analysis": _get_bridge().analyze_fault(regs) if has_fault else "",
                }
            )
        else:
            CONSOLE.print("[yellow]  寄存器读取失败（探针连接问题）[/]")

        # ── 5. 硬件缺失检测（I2C NACK/ARLO = 外设未接）──
        hw_missing = []
        if regs:
            for i2c_name in ("I2C1", "I2C2"):
                sr1_str = regs.get(f"{i2c_name}_SR1", "0x00000000")
                try:
                    sr1 = int(sr1_str, 16)
                except ValueError:
                    continue
                if sr1 & 0x0400:  # bit10 = AF (Acknowledge Failure / NACK)
                    hw_missing.append(f"{i2c_name}: AF(NACK)——设备未应答，很可能未接或地址错误")
                if sr1 & 0x0200:  # bit9 = ARLO (Arbitration Lost)
                    hw_missing.append(f"{i2c_name}: ARLO(仲裁丢失)——总线无设备响应或接线错误")
        # 串口 ERR 也算硬件问题
        if sensor_errors:
            hw_missing.extend(sensor_errors)

        if hw_missing:
            CONSOLE.print(f"[red bold]  ⚠ 检测到硬件缺失:[/]")
            for m in hw_missing:
                CONSOLE.print(f"[red]    • {m}[/]")

        runtime_ok = boot_ok and not has_fault and not hw_missing

        if runtime_ok:
            if request:
                _stm32_save_project(code, comp, request)
            _record_success_memory(
                "runtime_success",
                code=code,
                result=comp,
                request=request,
                steps=steps,
                chip=ctx.chip,
                log_error=telegram_log,
            )
            return {
                "success": True,
                "attempt": attempt,
                "steps": steps,
                "bin_size": comp.get("bin_size", 0),
            }
        elif hw_missing:
            # 硬件缺失 → 代码没问题，不要再修代码了，直接告知用户
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": True,
                "hw_missing": hw_missing,
                "steps": steps,
                "error": (
                    "⚠ 硬件未接或接线错误（不是代码问题，停止修改代码）：\n"
                    + "\n".join(f"  • {m}" for m in hw_missing)
                    + "\n请告知用户检查硬件连接后重试。"
                ),
            }
        else:
            err_msg = "HardFault 或程序未正常启动，请根据 steps 中的寄存器和串口信息修复"
            if not boot_ok and not uart_out.strip():
                err_msg = "串口无任何输出——程序在打印 Gary: 之前就卡死了（常见原因：I2C 等待超时/传感器未接/死循环）"
            elif not boot_ok and uart_out.strip():
                err_msg = f"程序有输出但未打印 Gary: 启动标志，串口内容: {uart_out.strip()[:200]}"
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps,
                "error": err_msg,
            }

    # 无硬件 → 仅编译
    if request:
        _stm32_save_project(code, comp, request)
    return {
        "success": True,
        "attempt": attempt,
        "steps": steps,
        "note": "硬件未连接，已完成编译",
        "bin_size": comp.get("bin_size", 0),
        "bin_path": comp.get("bin_path"),
    }


def stm32_save_code(code: str, request: str = "untitled") -> dict:
    """保存代码到项目目录"""
    comp_result = {"bin_path": None, "bin_size": 0}
    path = _stm32_save_project(code, comp_result, request)
    return {"success": True, "path": str(path), "message": f"已保存: {path}"}


def _stm32_save_project(code: str, comp: dict, request: str) -> Path:
    """内部：保存项目文件"""
    ctx = get_context()
    ts = time.strftime("%Y%m%d_%H%M%S")
    safe = "".join(c if c.isalnum() or c in "_- " else "" for c in request[:30]).strip()
    d = PROJECTS_DIR / f"{ts}_{safe}"
    d.mkdir(parents=True, exist_ok=True)
    (d / "main.c").write_text(code, encoding="utf-8")
    if comp.get("bin_path") and Path(comp["bin_path"]).exists():
        shutil.copy2(comp["bin_path"], d / "firmware.bin")
    (d / "config.json").write_text(
        json.dumps(
            {
                "chip": ctx.chip,
                "request": request,
                "bin_size": comp.get("bin_size", 0),
                "timestamp": ts,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    ctx.last_code = code
    # 同步更新 latest_workspace，保证 stm32_recompile 始终能找到最新代码
    try:
        latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
        latest.mkdir(parents=True, exist_ok=True)
        (latest / "main.c").write_text(code, encoding="utf-8")
    except Exception:
        pass
    CONSOLE.print(f"[dim]  已保存: {d}[/]")
    return d


def stm32_list_projects() -> dict:
    """列出最近 15 个历史项目"""
    if not PROJECTS_DIR.exists():
        return {"success": True, "projects": [], "message": "暂无项目"}
    projects = []
    for p in sorted(PROJECTS_DIR.iterdir(), reverse=True)[:15]:
        cf = p / "config.json"
        if cf.exists():
            try:
                c = json.loads(cf.read_text(encoding="utf-8"))
                projects.append(
                    {
                        "name": p.name,
                        "chip": c.get("chip", "?"),
                        "request": c.get("request", ""),
                        "timestamp": c.get("timestamp", ""),
                    }
                )
            except Exception:
                pass
    return {"success": True, "projects": projects}


def stm32_read_project(project_name: str) -> dict:
    """读取指定项目的 main.c 代码"""
    p = PROJECTS_DIR / project_name / "main.c"
    if not p.exists():
        return {"success": False, "message": f"项目不存在: {project_name}"}
    code = p.read_text(encoding="utf-8")
    return {"success": True, "code": code, "path": str(p), "lines": len(code.splitlines())}


# ─────────────────────────────────────────────────────────────
# 通用文件/命令工具（来自 claude_terminal，简化版）
# ─────────────────────────────────────────────────────────────


def read_file(file_path: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        lines = content.splitlines()
        numbered = "\n".join(f"{i+1:4d} | {l}" for i, l in enumerate(lines[:800]))
        return {
            "success": True,
            "numbered_view": numbered,
            "raw_content": content[:40000],
            "total_lines": len(lines),
        }
    except Exception as e:
        return {"error": str(e)}


def create_or_overwrite_file(file_path: str, content: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return {"success": True, "path": str(p), "lines": len(content.splitlines())}
    except Exception as e:
        return {"error": str(e)}


def str_replace_edit(file_path: str, old_str: str, new_str: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        count = content.count(old_str)
        if count == 0:
            return {"error": "未找到 old_str，请检查空格/换行是否完全一致"}
        if count > 1:
            return {"error": f"找到 {count} 个匹配，请增加上下文使其唯一"}
        new_content = content.replace(old_str, new_str, 1)
        p.write_text(new_content, encoding="utf-8")
        return {"success": True, "message": "替换成功", "path": str(p)}
    except Exception as e:
        return {"error": str(e)}


def list_directory(path: str = ".") -> dict:
    try:
        p = Path(path).expanduser().resolve()
        items = [{"name": x.name, "type": "dir" if x.is_dir() else "file"} for x in p.iterdir()]
        return {
            "success": True,
            "path": str(p),
            "items": sorted(items, key=lambda x: (x["type"], x["name"])),
        }
    except Exception as e:
        return {"error": str(e)}


def execute_command(command: str) -> dict:
    if any(f in command for f in ["rm -rf /", ":(){ :|:& };:"]):
        return {"error": "命令被安全策略拒绝"}
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=60)
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"error": "命令超时（60s）"}
    except Exception as e:
        return {"error": str(e)}


def search_files(query: str, path: str = ".", file_type: str = None) -> dict:
    try:
        results = []
        for fp in Path(path).expanduser().resolve().rglob("*"):
            if not fp.is_file():
                continue
            if query.lower() not in fp.name.lower():
                continue
            if file_type and fp.suffix != file_type:
                continue
            results.append(str(fp))
            if len(results) >= 20:
                break
        return {"success": True, "files": results}
    except Exception as e:
        return {"error": str(e)}


def web_search(query: str) -> dict:
    try:
        import requests

        r = requests.get(
            "http://127.0.0.1:8080/search", params={"q": query, "format": "json"}, timeout=8
        )
        data = r.json()
        results = [
            {"title": x.get("title"), "url": x.get("url"), "snippet": x.get("content", "")[:200]}
            for x in data.get("results", [])[:5]
        ]
        return {"success": True, "results": results}
    except Exception as e:
        return {"error": f"搜索失败（需要本地 SearXNG）: {e}"}


# ─────────────────────────────────────────────────────────────
# 扩展工具（来自 claude_terminal / tool_schemas）
# ─────────────────────────────────────────────────────────────


def append_file_content(file_path: str, content: str) -> dict:
    """向文件末尾追加内容"""
    try:
        p = Path(file_path).expanduser().resolve()
        mode = "a" if p.exists() else "w"
        prefix = ""
        if mode == "a" and p.stat().st_size > 0:
            with open(p, "rb") as f:
                f.seek(-1, 2)
                if f.read(1) != b"\n":
                    prefix = "\n"
        with open(p, mode, encoding="utf-8") as f:
            f.write(prefix + content)
        return {"success": True, "path": str(p), "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def grep_search(
    pattern: str, path: str = ".", include_extension: str = None, recursive: bool = True
) -> dict:
    """使用正则搜索文件内容（递归）"""
    try:
        search_path = Path(path).expanduser().resolve()
        results = []
        count = 0
        max_results = 20
        glob_pattern = "**/*" if recursive else "*"
        for fp in search_path.glob(glob_pattern):
            if not fp.is_file():
                continue
            if include_extension and fp.suffix != include_extension:
                continue
            if fp.stat().st_size > 1024 * 1024:
                continue
            try:
                with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                    file_content = f.read()
                matches = list(re.finditer(pattern, file_content, re.MULTILINE))
                if matches:
                    file_matches = []
                    for m in matches[:5]:
                        line_num = file_content.count("\n", 0, m.start()) + 1
                        line_start = file_content.rfind("\n", 0, m.start()) + 1
                        line_end = file_content.find("\n", m.end())
                        if line_end == -1:
                            line_end = len(file_content)
                        line_content = file_content[line_start:line_end].strip()
                        file_matches.append(f"Line {line_num}: {line_content[:100]}")
                    results.append(
                        f"File: {fp.relative_to(search_path)}\n" + "\n".join(file_matches)
                    )
                    count += 1
                    if count >= max_results:
                        break
            except Exception:
                continue
        return {
            "success": True,
            "matches_found": count,
            "results": "\n\n".join(results) if results else "No matches found",
        }
    except Exception as e:
        return {"error": str(e)}


def execute_batch_commands(commands: list, stop_on_error: bool = True) -> dict:
    """批量顺序执行多条 Shell 命令，默认遇错停止"""
    results = []
    overall_success = True
    for cmd in commands:
        res = execute_command(cmd)
        results.append({"command": cmd, "result": res})
        if not res.get("success", False):
            overall_success = False
            if stop_on_error:
                break
    return {"success": overall_success, "executed_count": len(results), "results": results}


def fetch_url(url: str) -> dict:
    """抓取 URL 页面并返回纯文本内容"""
    try:
        import requests

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return {"error": "beautifulsoup4 未安装: pip install beautifulsoup4"}
        headers = {"User-Agent": "Mozilla/5.0 (compatible; STM32Agent/1.0)"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        return {"success": True, "url": url, "content": text[:5000], "truncated": len(text) > 5000}
    except Exception as e:
        return {"error": f"获取失败: {e}"}


def get_current_time() -> dict:
    """获取当前系统时间、星期和时区"""
    try:
        now = datetime.now()
        return {
            "success": True,
            "current_time": now.strftime("%Y-%m-%d %H:%M:%S"),
            "weekday": now.strftime("%A"),
            "timezone": str(now.astimezone().tzinfo),
        }
    except Exception as e:
        return {"error": str(e)}


def ask_human(question: str) -> dict:
    """向用户提问并等待输入"""
    try:
        CONSOLE.print(f"\n[cyan][❓ AI Question]: {question}[/]")
        answer = input(" > ")
        return {"success": True, "answer": answer}
    except Exception as e:
        return {"error": str(e)}


def git_status() -> dict:
    """执行 git status 查看修改状态"""
    return execute_command("git status")


def git_diff() -> dict:
    """执行 git diff 查看代码变更"""
    return execute_command("git diff")


def git_commit(message: str) -> dict:
    """执行 git commit -m <message>"""
    return execute_command(f"git commit -m {shlex.quote(message)}")


def edit_file_lines(
    file_path: str, operation: str, start_line: int, end_line: int = None, new_content: str = None
) -> dict:
    """基于行号编辑文件（replace/insert/delete）"""
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        total = len(lines)
        if start_line < 1 or start_line > total:
            return {"error": f"start_line {start_line} 超出范围 [1, {total}]"}
        if end_line is None:
            end_line = start_line
        if end_line < start_line or end_line > total:
            return {"error": f"end_line {end_line} 无效"}
        si, ei = start_line - 1, end_line
        if operation == "replace":
            if new_content is None:
                return {"error": "replace 需要 new_content"}
            if not new_content.endswith("\n"):
                new_content += "\n"
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[ei:]
        elif operation == "insert":
            if new_content is None:
                return {"error": "insert 需要 new_content"}
            if not new_content.endswith("\n"):
                new_content += "\n"
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[si:]
        elif operation == "delete":
            new_lines = lines[:si] + lines[ei:]
        else:
            return {"error": f"未知操作: {operation}"}
        with open(p, "w", encoding="utf-8") as f:
            f.writelines(new_lines)
        return {
            "success": True,
            "path": str(p),
            "operation": operation,
            "new_total_lines": len(new_lines),
        }
    except Exception as e:
        return {"error": str(e)}


def insert_content_by_regex(file_path: str, regex_pattern: str, content: str) -> dict:
    """在文件第一个正则匹配位置之后插入内容"""
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, "r", encoding="utf-8", errors="ignore") as f:
            file_content = f.read()
        m = re.search(regex_pattern, file_content, re.MULTILINE)
        if not m:
            return {"error": f"正则 '{regex_pattern}' 未匹配到内容"}
        new_content = file_content[: m.end()] + content + file_content[m.end() :]
        with open(p, "w", encoding="utf-8") as f:
            f.write(new_content)
        return {
            "success": True,
            "path": str(p),
            "match_found": m.group(0)[:50],
            "message": "内容已插入",
        }
    except Exception as e:
        return {"error": str(e)}


def check_python_code(file_path: str) -> dict:
    """检查 Python 文件语法和风格（flake8 / ast）"""
    import ast

    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        try:
            with open(p, "r", encoding="utf-8") as f:
                ast.parse(f.read())
        except SyntaxError as e:
            return {
                "success": False,
                "error_type": "SyntaxError",
                "line": e.lineno,
                "message": str(e),
            }
        lint_result = ""
        try:
            result = subprocess.run(
                f"flake8 {shlex.quote(str(p))}",
                shell=True,
                capture_output=True,
                text=True,
                timeout=10,
            )
            if result.returncode != 0 and result.stdout:
                lint_result = f"Flake8:\n{result.stdout}"
        except Exception:
            pass
        return {
            "success": True,
            "message": "语法检查通过",
            "linter_output": lint_result or "无问题",
        }
    except Exception as e:
        return {"error": str(e)}


def run_python_code(code: str) -> dict:
    """执行 Python 代码片段（临时文件沙箱）"""
    import tempfile

    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write(code)
            tmp_path = tmp.name
        result = subprocess.run(
            [sys.executable, tmp_path], capture_output=True, text=True, timeout=30
        )
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except Exception as e:
        return {"error": str(e)}


# ── Word 文档工具 ──────────────────────────────────────────────


def _get_docx_module():
    """懒加载 python-docx，未安装时返回 None"""
    try:
        import docx

        return docx
    except ImportError:
        return None


def read_docx(file_path: str) -> dict:
    """读取 Word 文档(.docx)的文本内容"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return {"success": True, "content": text, "total_paragraphs": len(doc.paragraphs)}
    except Exception as e:
        return {"error": str(e)}


def replace_docx_text(
    file_path: str, old_text: str, new_text: str, use_regex: bool = False
) -> dict:
    """替换 Word 文档中的文本（支持正则）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        count = 0
        for para in doc.paragraphs:
            if use_regex:
                if re.search(old_text, para.text):
                    replaced = re.sub(old_text, new_text, para.text)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = replaced
                    else:
                        para.add_run(replaced)
                    count += 1
            else:
                if old_text in para.text:
                    replaced_in_run = False
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
                            replaced_in_run = True
                    if not replaced_in_run:
                        para.text = para.text.replace(old_text, new_text)
                        count += 1
        doc.save(file_path)
        return {"success": True, "replaced_count": count, "message": f"已替换 {count} 处"}
    except Exception as e:
        return {"error": str(e)}


def append_docx_content(
    file_path: str, content: str, after_paragraph_index: int = None, style: str = None
) -> dict:
    """向 Word 文档追加内容（支持指定位置插入）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        paragraphs_text = [t for t in content.split("\n") if t.strip()]
        if after_paragraph_index is None:
            for p_text in paragraphs_text:
                p = doc.add_paragraph(p_text)
                if style:
                    try:
                        p.style = style
                    except Exception:
                        pass
        else:
            n = len(doc.paragraphs)
            if after_paragraph_index < 0 or after_paragraph_index >= n:
                return {"error": f"索引 {after_paragraph_index} 超出范围 (0-{n-1})"}
            if after_paragraph_index == n - 1:
                for p_text in paragraphs_text:
                    p = doc.add_paragraph(p_text)
                    if style:
                        try:
                            p.style = style
                        except Exception:
                            pass
            else:
                next_para = doc.paragraphs[after_paragraph_index + 1]
                base_style = doc.paragraphs[after_paragraph_index].style
                for p_text in paragraphs_text:
                    new_p = next_para.insert_paragraph_before(p_text)
                    if style:
                        try:
                            new_p.style = style
                        except Exception:
                            pass
                    else:
                        new_p.style = base_style
        doc.save(file_path)
        return {"success": True, "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def inspect_docx_structure(file_path: str, max_paragraphs: int = 50) -> dict:
    """查看 Word 文档段落结构（用于定位插入点）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        structure = []
        for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            if not preview.strip():
                preview = "[空段落]"
            structure.append(f"[{i}] {preview}")
        return {
            "success": True,
            "total_paragraphs": len(doc.paragraphs),
            "structure": "\n".join(structure),
        }
    except Exception as e:
        return {"error": str(e)}


def insert_docx_content_after_heading(
    file_path: str, heading_text: str, content: str, style: str = None
) -> dict:
    """在 Word 文档指定标题后插入内容（大小写不敏感）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        target_para = None
        for para in doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                target_para = para
                break
        if not target_para:
            return {"error": f"未找到标题: {heading_text}"}
        index = doc.paragraphs.index(target_para)
        return append_docx_content(file_path, content, index, style)
    except Exception as e:
        return {"error": str(e)}


# ── 电脑控制工具 ──────────────────────────────────────────────


def computer_screenshot() -> dict:
    """截取当前桌面截图并保存为 PNG"""
    try:
        import pyautogui

        ts = int(time.time())
        path = os.path.abspath(f"screenshot_{ts}.png")
        pyautogui.screenshot(path)
        return {"success": True, "image_path": path, "message": f"截图已保存: {path}"}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_move(x: int, y: int) -> dict:
    """移动鼠标到指定坐标"""
    try:
        import pyautogui

        pyautogui.moveTo(x, y)
        return {"success": True, "action": "move", "x": x, "y": y}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_click(button: str = "left") -> dict:
    """鼠标点击（left/right/double）"""
    try:
        import pyautogui

        if button == "double":
            pyautogui.doubleClick()
        else:
            pyautogui.click(button=button)
        return {"success": True, "button": button}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_keyboard_type(text: str) -> dict:
    """向焦点窗口输入文本"""
    try:
        import pyautogui

        pyautogui.write(text)
        return {"success": True, "typed": text}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


# ─────────────────────────────────────────────────────────────
# 工具注册表
# ─────────────────────────────────────────────────────────────
bind_tool_implementations(
    {
        # STM32 专属
        "stm32_list_probes": stm32_list_probes,
        "stm32_connect": stm32_connect,
        "stm32_disconnect": stm32_disconnect,
        "stm32_serial_connect": stm32_serial_connect,
        "stm32_serial_disconnect": stm32_serial_disconnect,
        "stm32_set_chip": stm32_set_chip,
        "stm32_hardware_status": stm32_hardware_status,
        "stm32_compile": stm32_compile,
        "stm32_compile_rtos": stm32_compile_rtos,
        "stm32_recompile": stm32_recompile,
        "stm32_regen_bsp": stm32_regen_bsp,
        "stm32_analyze_fault_rtos": stm32_analyze_fault_rtos,
        "stm32_rtos_check_code": stm32_rtos_check_code,
        "stm32_rtos_task_stats": stm32_rtos_task_stats,
        "stm32_rtos_suggest_config": stm32_rtos_suggest_config,
        "stm32_rtos_plan_project": stm32_rtos_plan_project,
        "stm32_flash": stm32_flash,
        "stm32_read_registers": stm32_read_registers,
        "stm32_analyze_fault": stm32_analyze_fault,
        "stm32_serial_read": stm32_serial_read,
        "stm32_auto_flash_cycle": stm32_auto_flash_cycle,
        "stm32_reset_debug_attempts": stm32_reset_debug_attempts,
        "stm32_generate_font": stm32_generate_font,
        "stm32_save_code": stm32_save_code,
        "stm32_list_projects": stm32_list_projects,
        "stm32_read_project": stm32_read_project,
        "gary_save_member_memory": gary_save_member_memory,
        # 通用
        "read_file": read_file,
        "create_or_overwrite_file": create_or_overwrite_file,
        "str_replace_edit": str_replace_edit,
        "list_directory": list_directory,
        "execute_command": execute_command,
        "search_files": search_files,
        "web_search": web_search,
        # 扩展通用工具
        "append_file_content": append_file_content,
        "grep_search": grep_search,
        "execute_batch_commands": execute_batch_commands,
        "fetch_url": fetch_url,
        "get_current_time": get_current_time,
        "ask_human": ask_human,
        "git_status": git_status,
        "git_diff": git_diff,
        "git_commit": git_commit,
        "edit_file_lines": edit_file_lines,
        "insert_content_by_regex": insert_content_by_regex,
        "check_python_code": check_python_code,
        "run_python_code": run_python_code,
        # Word 文档工具
        "read_docx": read_docx,
        "replace_docx_text": replace_docx_text,
        "append_docx_content": append_docx_content,
        "inspect_docx_structure": inspect_docx_structure,
        "insert_docx_content_after_heading": insert_docx_content_after_heading,
        # 电脑控制工具
        "computer_screenshot": computer_screenshot,
        "computer_mouse_move": computer_mouse_move,
        "computer_mouse_click": computer_mouse_click,
        "computer_keyboard_type": computer_keyboard_type,
    }
)

# ─────────────────────────────────────────────────────────────
# STM32 系统提示词
# ─────────────────────────────────────────────────────────────
# Gary config — CLI 内 AI 接口配置向导
# ─────────────────────────────────────────────────────────────
def configure_ai_cli(agent: "STM32Agent | None" = None):
    """交互式配置 AI 接口（可在 CLI 内调用，也可独立运行）"""
    import getpass as _gp

    CONSOLE.print()
    CONSOLE.rule("[bold cyan]  配置 AI 后端接口[/]")
    CONSOLE.print()

    cur_key, cur_url, cur_model = _read_ai_config()
    is_configured = bool(cur_key and not _api_key_is_placeholder(cur_key))

    if is_configured:
        CONSOLE.print(f"  [dim]当前 API Key :[/] {_mask_key(cur_key)}")
        CONSOLE.print(f"  [dim]当前 Base URL:[/] {cur_url}")
        CONSOLE.print(f"  [dim]当前 Model   :[/] {cur_model}")
        CONSOLE.print()

    # 服务商菜单
    CONSOLE.print("[bold cyan]  请选择 AI 服务提供商：[/]")
    for i, (name, url, _) in enumerate(_AI_PRESETS, 1):
        url_hint = f"  [dim]{url[:55]}[/]" if url else ""
        CONSOLE.print(f"    [yellow]{i}[/].  {name:<24}{url_hint}")
    CONSOLE.print()

    valid = [str(i) for i in range(1, len(_AI_PRESETS) + 1)]
    choice = ""
    while choice not in valid:
        try:
            choice = input(f"  输入序号 [1-{len(_AI_PRESETS)}] (回车取消): ").strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]")
            return
        if choice == "":
            CONSOLE.print("[dim]已取消[/]")
            return

    idx = int(choice) - 1
    preset_name, preset_url, preset_model = _AI_PRESETS[idx]

    # Base URL
    if preset_url:
        base_url = preset_url
        CONSOLE.print(f"  [dim]Base URL: {base_url}[/]")
    else:
        try:
            base_url = input("  Base URL: ").strip()
        except (EOFError, KeyboardInterrupt):
            base_url = cur_url

    # Model
    default_model = preset_model or cur_model or ""
    try:
        hint = f" (默认 {default_model})" if default_model else ""
        entered = input(f"  Model 名称{hint}: ").strip()
        model = entered if entered else default_model
    except (EOFError, KeyboardInterrupt):
        model = default_model

    # API Key
    CONSOLE.print()
    if preset_name == "Ollama (本地)":
        api_key = "ollama"
        CONSOLE.print("  [dim]Ollama 本地模式，API Key 自动设为 ollama[/]")
    else:
        CONSOLE.print(f"  [dim]请输入 {preset_name} API Key（不显示输入内容）[/]")
        try:
            api_key = _gp.getpass("  API Key: ")
        except Exception:
            try:
                api_key = input("  API Key: ").strip()
            except (EOFError, KeyboardInterrupt):
                api_key = ""
        if not api_key:
            if is_configured:
                CONSOLE.print("  [dim]未输入，保留原有 Key[/]")
                api_key = cur_key
            else:
                CONSOLE.print("[yellow]  未输入 API Key，配置取消[/]")
                return

    # 写入 config.py
    if _write_ai_config(api_key, base_url, model):
        _sync_ai_runtime_settings(reload_ai_config())
        CONSOLE.print()
        CONSOLE.print("[green]  ✓ 配置已保存到 config.py[/]")
        CONSOLE.print(f"  [green]✓[/] 服务商  {preset_name}")
        CONSOLE.print(f"  [green]✓[/] API Key {_mask_key(api_key)}")
        CONSOLE.print(f"  [green]✓[/] Model   {model}")
        # 重建当前会话的 AI 客户端
        if agent is not None:
            agent.client = get_ai_client(timeout=180.0, force_reload=True)
            CONSOLE.print("  [green]✓[/] AI 客户端已热重载，无需重启")
    else:
        CONSOLE.print("[red]  ✗ 写入 config.py 失败[/]")
    CONSOLE.print()


# ─────────────────────────────────────────────────────────────
# STM32 Agent（TUI + 流式对话 + 工具框架）
# ─────────────────────────────────────────────────────────────
class STM32Agent:
    def __init__(self, interactive: bool = True):
        self.interactive = interactive
        self.messages: List[Dict] = [{"role": "system", "content": self._compose_system_prompt()}]
        self.client = get_ai_client(timeout=180.0)
        self.command_completer = GaryCompleter(
            list_projects=stm32_list_projects,
            default_chip=DEFAULT_CHIP,
        )
        os.environ.setdefault("no_proxy", "localhost,127.0.0.1")
        os.environ.setdefault("NO_PROXY", "localhost,127.0.0.1")

    def refresh_ai_client(self):
        self.client = get_ai_client(timeout=180.0)

    def _compose_system_prompt(self) -> str:
        """Compose the runtime system prompt from modular prompt providers."""

        ctx = get_context()
        prompt = build_system_prompt(ctx.chip, ctx.cli_language, ctx.hw_connected)
        member_prompt = get_member_prompt_section()
        if member_prompt:
            prompt += "\n\n" + member_prompt
        for error_type in ("compile_error", "hardfault", "i2c_failure"):
            debug_prompt = get_debug_prompt(
                error_type,
                {"chip": ctx.chip, "hardware_connected": ctx.hw_connected},
            )
            if debug_prompt:
                prompt += "\n\n" + debug_prompt
        skill_prompt = _get_manager().get_all_prompt_additions()
        if skill_prompt:
            prompt += "\n" + skill_prompt.lstrip("\n")
        return prompt

    def reset_conversation(self):
        self.messages = [{"role": "system", "content": self._compose_system_prompt()}]

    def refresh_system_prompt(self):
        prompt = self._compose_system_prompt()
        if self.messages and self.messages[0].get("role") == "system":
            self.messages[0]["content"] = prompt
        else:
            self.messages.insert(0, {"role": "system", "content": prompt})

    def set_cli_language(self, language: str) -> dict:
        ctx = get_context()
        target = _normalize_cli_language(language)
        ctx.cli_language = target
        globals()["CLI_LANGUAGE"] = target
        saved = _write_cli_language_config(target)
        if saved:
            _sync_ai_runtime_settings(reload_ai_config())
        self.refresh_system_prompt()
        return {"success": True, "language": ctx.cli_language, "saved": saved}

    # ── Token 估算 ──────────────────────────────────────────
    # 替换原来的 _tokens 和 _truncate 方法
    def _tokens(self) -> int:
        return sum(len(str(m.get("content", ""))) // 3 for m in self.messages)

    def _truncate_result(self, s: str, tool_name: str = "") -> str:
        """针对不同工具结果使用不同截断策略"""
        if len(s) <= MAX_TOOL_RESULT_LEN:
            return s
        half = MAX_TOOL_RESULT_LEN // 2
        # 编译/串口结果：错误在末尾，保留末尾更多
        if tool_name in ("stm32_compile", "stm32_serial_read", "stm32_auto_flash_cycle"):
            head = MAX_TOOL_RESULT_LEN // 4
            tail = MAX_TOOL_RESULT_LEN - head
            return s[:head] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-tail:]
        return s[:half] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-half:]

    def _truncate_history(self):
        """滑动窗口：保留 system prompt + 最近消息，总字符不超限"""
        MAX_CHARS = 180_000
        total = sum(len(str(m.get("content", ""))) for m in self.messages)
        removed = 0
        while total > MAX_CHARS and len(self.messages) > 3:
            # 始终保留 messages[0]（system prompt）
            victim = self.messages.pop(1)
            victim_len = len(str(victim.get("content", "")))
            total -= victim_len
            removed += 1
        if removed and self.interactive:
            CONSOLE.print(
                f"[dim]  📦 {_cli_text(f'历史压缩：移除 {removed} 条旧消息', f'History trimmed: removed {removed} old messages')}[/]"
            )

    # 原来是直接传 self.messages，改为过滤后传
    def _messages_for_api(self) -> list:
        """发送给 API 前处理消息格式：
        - 若对话中出现过 reasoning_content（thinking 模式），则所有 assistant 消息都必须带该字段
        - 否则过滤掉该字段（避免不支持的 API 报错）
        """
        # 检测当前会话是否启用了 thinking 模式
        has_thinking = any(
            "reasoning_content" in m for m in self.messages if m.get("role") == "assistant"
        )

        result = []
        for m in self.messages:
            if m.get("role") == "assistant" and has_thinking:
                # thinking 模式：确保每条 assistant 消息都有 reasoning_content
                clean = dict(m)
                if "reasoning_content" not in clean:
                    clean["reasoning_content"] = ""
                result.append(clean)
            else:
                # 非 thinking 模式：过滤掉该字段
                clean = {k: v for k, v in m.items() if k != "reasoning_content"}
                result.append(clean)
        return result

    def _request_final_reply_after_tools(self, stream_to_console: bool = True) -> str:
        """部分模型在工具执行后会停在空回复，这里补一次只求最终答复的请求。"""
        if stream_to_console:
            CONSOLE.print(
                f"[dim]  ↺ {_cli_text('请求最终答复...', 'Requesting final reply...')}[/]"
            )
        telegram_log("chat final_reply_request start")
        try:
            stream = stream_chat(
                client=self.client,
                messages=self._messages_for_api()
                + [
                    {
                        "role": "system",
                        "content": _cli_text(
                            "请基于上面的工具结果直接给出最终答复，不要再调用工具，也不要只回复“已处理”。",
                            "Based on the tool results above, provide the final answer directly. Do not call more tools and do not reply with only 'done'.",
                        ),
                    }
                ],
                model=AI_MODEL,
                temperature=AI_TEMPERATURE,
            )
        except Exception as e:
            if stream_to_console:
                CONSOLE.print(
                    f"\n[red]{_cli_text('最终答复请求失败', 'Final reply request failed')}: {e}[/]"
                )
            telegram_log(f"chat final_reply_request error={str(e)[:160]}")
            return ""

        content = ""
        thinking = ""
        in_think = False
        try:
            for chunk in stream:
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta

                rc = getattr(delta, "reasoning_content", None)
                if rc:
                    if not in_think and stream_to_console:
                        CONSOLE.print(f"\n[dim {THEME}]💭 思考:[/]")
                    in_think = True
                    thinking += rc
                    if stream_to_console:
                        CONSOLE.print(rc, end="", style="dim")

                if delta.content:
                    if in_think and stream_to_console:
                        CONSOLE.print()
                    in_think = False
                    content += delta.content
                    if stream_to_console:
                        CONSOLE.print(delta.content, end="", style="white")

            if in_think and stream_to_console:
                CONSOLE.print()
            if content and stream_to_console:
                CONSOLE.print()
        except Exception as e:
            if stream_to_console:
                CONSOLE.print(
                    f"\n[red]{_cli_text('最终答复流式读取错误', 'Final reply stream error')}: {e}[/]"
                )
            telegram_log(f"chat final_reply_stream error={str(e)[:160]}")
            return ""

        telegram_log(f"chat final_reply_request done len={len(content.strip())}")
        return content.strip()

    def _summarize_tool_result(self, tool_name: str, result_obj, preview: str) -> str:
        text = (preview or "").replace("\n", " ").strip()
        if isinstance(result_obj, dict):
            pieces = []
            if result_obj.get("message"):
                pieces.append(str(result_obj["message"]).strip())
            if result_obj.get("error"):
                pieces.append(str(result_obj["error"]).strip())
            if result_obj.get("path"):
                pieces.append(f"path={result_obj['path']}")
            if result_obj.get("chip"):
                pieces.append(f"chip={result_obj['chip']}")
            if result_obj.get("attempt") is not None:
                pieces.append(f"attempt={result_obj['attempt']}")
            if not pieces and "success" in result_obj:
                pieces.append(f"success={result_obj.get('success')}")
            if pieces:
                text = " | ".join(pieces)
        text = text or "已执行"
        if len(text) > 180:
            text = text[:177] + "..."
        return f"{tool_name}: {text}"

    def _build_tool_only_reply(self, tool_summaries: list[str], reply_parts: list[str]) -> str:
        lines = [
            _cli_text(
                "模型没有输出最终总结，我根据本次执行结果整理如下：",
                "The model did not return a final summary. Based on this run:",
            )
        ]
        if reply_parts:
            preface = (reply_parts[-1] or "").strip()
            if preface:
                if len(preface) > 240:
                    preface = preface[:237] + "..."
                lines.append(preface)
        for item in tool_summaries[-5:]:
            lines.append(f"- {item}")
        return "\n".join(lines)

    # ── 流式响应 + 工具调用 ─────────────────────────────────
    def chat(
        self,
        user_input: str,
        stream_to_console: bool = True,
        text_callback=None,
        tool_callback=None,
    ) -> str:
        self._truncate_history()
        self.messages.append({"role": "user", "content": user_input})
        reply_parts: List[str] = []
        tool_summaries: List[str] = []
        used_tools = False

        while True:
            # API 调用
            try:
                stream = stream_chat(
                    client=self.client,
                    messages=self._messages_for_api(),
                    model=AI_MODEL,
                    tools=TOOL_SCHEMAS,
                    tool_choice="auto",
                    temperature=AI_TEMPERATURE,
                )
            except Exception as e:
                if stream_to_console:
                    CONSOLE.print(f"\n[red]{_cli_text('API 错误', 'API error')}: {e}[/]")
                return f"{_cli_text('API 错误', 'API error')}: {e}"

            # 收集流式输出
            content = ""
            tool_calls_raw: Dict[int, dict] = {}
            thinking = ""
            in_think = False

            try:
                for chunk in stream:
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta

                    # Reasoning (deepseek-r1 style)
                    rc = getattr(delta, "reasoning_content", None)
                    if rc:
                        if not in_think and stream_to_console:
                            CONSOLE.print(f"\n[dim {THEME}]💭 思考:[/]")
                        in_think = True
                        thinking += rc
                        if stream_to_console:
                            CONSOLE.print(rc, end="", style="dim")

                    # 文本内容
                    if delta.content:
                        if in_think and stream_to_console:
                            CONSOLE.print()
                        in_think = False
                        content += delta.content
                        if stream_to_console:
                            CONSOLE.print(delta.content, end="", style="white")
                        if text_callback:
                            preview_text = "\n\n".join(
                                part for part in [*reply_parts, content.strip()] if part
                            ).strip()
                            text_callback(preview_text)

                    # 工具调用
                    if delta.tool_calls:
                        # 用 model_dump() 获取含 Gemini extra_content 的完整 chunk 数据
                        try:
                            _chunk_dict = chunk.model_dump()
                            _raw_tcs = (_chunk_dict.get("choices") or [{}])[0].get("delta", {}).get(
                                "tool_calls"
                            ) or []
                        except Exception:
                            _raw_tcs = []
                        for i, tc in enumerate(delta.tool_calls):
                            idx = tc.index
                            if idx not in tool_calls_raw:
                                tool_calls_raw[idx] = {
                                    "id": "",
                                    "name": "",
                                    "args": "",
                                    "thought_signature": "",
                                }
                            if tc.id:
                                tool_calls_raw[idx]["id"] = tc.id
                            if tc.function and tc.function.name:
                                tool_calls_raw[idx]["name"] = tc.function.name
                            if tc.function and tc.function.arguments:
                                tool_calls_raw[idx]["args"] += tc.function.arguments
                            # Gemini thinking models 将签名放在 extra_content.google.thought_signature
                            # 必须原样回传到 function.thought_signature，否则下次请求报 400
                            _raw_tc = _raw_tcs[i] if i < len(_raw_tcs) else {}
                            sig = (
                                (_raw_tc.get("extra_content") or {})
                                .get("google", {})
                                .get("thought_signature")
                                or (_raw_tc.get("function") or {}).get("thought_signature")
                                or _raw_tc.get("thought_signature")
                                or getattr(tc, "thought_signature", None)
                                or (
                                    getattr(tc.function, "thought_signature", None)
                                    if tc.function
                                    else None
                                )
                            )
                            if sig:
                                tool_calls_raw[idx]["thought_signature"] += sig

                if in_think and stream_to_console:
                    CONSOLE.print()
                if content and stream_to_console:
                    CONSOLE.print()

            except Exception as e:
                if stream_to_console:
                    CONSOLE.print(f"\n[red]{_cli_text('流式读取错误', 'Streaming error')}: {e}[/]")
                return f"{_cli_text('流式读取错误', 'Streaming error')}: {e}"

            # 无工具调用 → 结束
            if not tool_calls_raw:
                if content.strip():
                    assistant_msg = {"role": "assistant", "content": content or ""}
                    if thinking:  # 如果有思考内容，带上它
                        assistant_msg["reasoning_content"] = thinking
                    self.messages.append(assistant_msg)
                    reply_parts.append(content.strip())
                    break
                if used_tools:
                    final_reply = self._request_final_reply_after_tools(
                        stream_to_console=stream_to_console
                    )
                    if final_reply:
                        self.messages.append({"role": "assistant", "content": final_reply})
                        reply_parts.append(final_reply)
                        break
                    fallback_reply = self._build_tool_only_reply(tool_summaries, reply_parts)
                    self.messages.append({"role": "assistant", "content": fallback_reply})
                    reply_parts.append(fallback_reply)
                    break
                assistant_msg = {"role": "assistant", "content": content or ""}
                if thinking:  # 如果有思考内容，带上它
                    assistant_msg["reasoning_content"] = thinking
                self.messages.append(assistant_msg)
                break

            # 构造 assistant tool_calls 消息
            tool_calls_list = []
            for idx in sorted(tool_calls_raw.keys()):
                tc = tool_calls_raw[idx]
                func_dict: dict = {"name": tc["name"], "arguments": tc["args"]}
                if tc.get("thought_signature"):  # Gemini 思考签名，必须原样回传
                    func_dict["thought_signature"] = tc["thought_signature"]
                tool_calls_list.append(
                    {
                        "id": tc["id"],
                        "type": "function",
                        "function": func_dict,
                    }
                )
            assistant_tool_msg = {
                "role": "assistant",
                "content": content or "",
                "tool_calls": tool_calls_list,
            }
            if thinking:  # 关键修复：把之前收集的 thinking 塞进来
                assistant_tool_msg["reasoning_content"] = thinking

            self.messages.append(assistant_tool_msg)
            used_tools = True

            # 执行工具
            tool_results = []
            for tc in tool_calls_list:
                func_name = tc["function"]["name"]
                args_str = tc["function"]["arguments"]
                result_obj = None
                telegram_log(f"chat tool_exec_start name={func_name}")

                if stream_to_console:
                    CONSOLE.print(f"[dim]  🔧 {func_name}[/]", end="")
                if tool_callback:
                    tool_callback({"phase": "start", "name": func_name, "arguments": args_str})
                try:
                    args = json.loads(args_str) if args_str.strip() else {}
                    result_obj = dispatch_tool_call(
                        func_name,
                        args,
                        get_context_fn=get_context,
                    )
                    result_str = json.dumps(result_obj, ensure_ascii=False, indent=2)
                except Exception as e:
                    result_str = f'{{"error": "{e}"}}'
                    result_obj = {"error": str(e)}
                    if tool_callback:
                        tool_callback({"phase": "error", "name": func_name, "error": str(e)})

                # 简短预览
                preview = result_str[:120].replace("\n", " ")
                if stream_to_console:
                    CONSOLE.print(f" → [dim green]{preview}[/]")
                if tool_callback:
                    tool_callback(
                        {
                            "phase": "finish",
                            "name": func_name,
                            "preview": preview,
                            "result": result_str,
                        }
                    )
                telegram_log(f"chat tool_exec_finish name={func_name} preview={preview[:80]}")
                tool_summaries.append(self._summarize_tool_result(func_name, result_obj, preview))

                tool_results.append(
                    {
                        "role": "tool",
                        "tool_call_id": tc["id"],
                        "content": self._truncate_result(result_str, func_name),
                    }
                )

            self.messages.extend(tool_results)
            self.refresh_system_prompt()
            # 继续循环，把工具结果发回 AI

            if content.strip():
                reply_parts.append(content.strip())

        return "\n\n".join(part for part in reply_parts if part).strip()

    # ── 内置命令处理 ────────────────────────────────────────
    def handle_builtin(self, cmd: str) -> bool:
        """Delegate slash command handling to the modular TUI command router."""

        return handle_slash_command(
            self,
            cmd,
            theme=THEME,
            cli_text=_cli_text,
            actions=_command_actions(),
            print_banner=_print_runtime_banner,
        )

    def run(self) -> None:
        """Run the interactive TUI loop."""

        run_interactive(
            self,
            handle_command=lambda agent, command: agent.handle_builtin(command),
            cli_text=_cli_text,
            theme=THEME,
            model=AI_MODEL,
            status_snapshot=_status_snapshot,
            ensure_cli_telegram_daemon=_ensure_cli_telegram_daemon,
            shutdown_runtime=_shutdown_cli_runtime,
            history_path=GARY_HOME / "prompt_history.txt",
            ensure_gary_home=_ensure_gary_home,
            completer=self.command_completer,
        )



def _status_snapshot() -> dict[str, Any]:
    """Return the current UI status snapshot for banner/status rendering."""

    ctx = get_context()
    return {
        "chip": ctx.chip,
        "hw_connected": ctx.hw_connected,
        "serial_connected": ctx.serial_connected,
    }



def _print_runtime_banner() -> None:
    """Render the current runtime banner using the modular TUI UI helpers."""

    print_banner(
        model=AI_MODEL,
        cli_text=_cli_text,
        theme=THEME,
        **_status_snapshot(),
    )



def _command_actions() -> dict[str, Callable[..., Any]]:
    """Build the slash-command action map for the TUI command router."""

    return {
        "connect": stm32_connect,
        "disconnect": stm32_disconnect,
        "serial_connect": stm32_serial_connect,
        "serial_disconnect": stm32_serial_disconnect,
        "set_chip": stm32_set_chip,
        "flash": stm32_flash,
        "hardware_status": stm32_hardware_status,
        "list_probes": stm32_list_probes,
        "list_projects": stm32_list_projects,
        "configure_ai_cli": configure_ai_cli,
        "shutdown_runtime": _shutdown_cli_runtime,
        "parse_cli_language": _parse_cli_language,
    }


def _configure_telegram_runtime() -> None:
    """Register stm32_agent runtime hooks for the Telegram integration."""

    configure_telegram_integration(
        console=CONSOLE,
        cli_text=_cli_text,
        is_ai_configured=_ai_is_configured,
        configure_ai_cli=lambda: configure_ai_cli(),
        agent_factory=lambda interactive=False: STM32Agent(interactive=interactive),
        hardware_status=stm32_hardware_status,
        connect=stm32_connect,
        disconnect=stm32_disconnect,
        set_chip=stm32_set_chip,
        list_projects=stm32_list_projects,
        detect_serial_ports=detect_serial_ports,
        serial_connect=stm32_serial_connect,
        get_current_chip=lambda: get_context().chip or "(未设置)",
        script_path=_PROJECT_ROOT / "stm32_agent.py",
        workdir=_PROJECT_ROOT,
    )


_configure_telegram_runtime()


def _print_startup_checks() -> None:
    """Render startup dependency checks before entering interactive mode."""

    ctx = get_context()
    CONSOLE.print(f"[dim]{_cli_text('检查环境...', 'Checking environment...')}[/]")

    compiler = _get_compiler()
    info = compiler.check(ctx.chip)
    if info.get("gcc"):
        CONSOLE.print(f"[green]  GCC: {info['gcc_version']}[/]")
    else:
        CONSOLE.print(
            f"[yellow]  GCC: {_cli_text('未找到 arm-none-eabi-gcc', 'arm-none-eabi-gcc not found')}[/]"
        )

    if info.get("hal"):
        CONSOLE.print(f"[green]  HAL: {_cli_text('已就绪', 'ready')}[/]")
    else:
        CONSOLE.print(
            f"[yellow]  HAL: {_cli_text('未找到，请运行 setup.sh', 'not found, run setup.sh')}[/]"
        )

    try:
        import pyocd

        CONSOLE.print(f"[green]  pyocd: {pyocd.__version__}[/]")
    except ImportError:
        CONSOLE.print(
            f"[yellow]  pyocd: {_cli_text('未安装（pip install pyocd）', 'not installed (pip install pyocd)')}[/]"
        )

    import glob as _glob
    import platform as _platform

    serial_candidates = detect_serial_ports(verbose=False)
    if serial_candidates:
        CONSOLE.print(
            f"[green]  {_cli_text('串口', 'Serial')}: {_cli_text(f'检测到 {serial_candidates}（连接时自动选择）', f'detected {serial_candidates} (auto-selected on connect)')}[/]"
        )
    else:
        platform_name = _platform.system()
        if platform_name == "Linux":
            import grp as _grp

            no_perm = [
                path
                for pattern in ["/dev/ttyUSB*", "/dev/ttyACM*", "/dev/ttyS[4-9]*"]
                for path in _glob.glob(pattern)
                if os.path.exists(path) and not os.access(path, os.R_OK | os.W_OK)
            ]
            if no_perm:
                try:
                    grp_name = _grp.getgrgid(os.stat(no_perm[0]).st_gid).gr_name
                except Exception:
                    grp_name = "dialout"
                CONSOLE.print(
                    f"[yellow]  {_cli_text('串口', 'Serial')}: {_cli_text(f'发现 {no_perm} 但权限不足', f'found {no_perm} but permissions are insufficient')}[/]"
                )
                CONSOLE.print(
                    f"[yellow]    → sudo usermod -aG {grp_name} $USER && newgrp {grp_name}[/]"
                )
            else:
                CONSOLE.print(
                    f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备（连接硬件后重试）', 'no serial device detected (connect hardware and retry)')}[/]"
                )
        elif platform_name == "Darwin":
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备', 'no serial device detected')}[/]"
            )
            CONSOLE.print(
                f"[dim]    {_cli_text('插上 USB 转串口后重启程序，或运行 /serial list 查看', 'Reconnect your USB serial adapter and restart, or run /serial list')}[/]"
            )
        elif platform_name == "Windows":
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到 COM 口', 'no COM ports detected')}[/]"
            )
            CONSOLE.print(
                f"[dim]    {_cli_text('请在设备管理器确认驱动已安装（CH340/CP210x），或运行 /serial list', 'Check the driver in Device Manager (CH340/CP210x), or run /serial list')}[/]"
            )
        else:
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备（连接硬件后重试）', 'no serial device detected (connect hardware and retry)')}[/]"
            )

    CONSOLE.print()


def run(args: argparse.Namespace) -> None:
    """Run the Gary CLI using parsed command-line arguments."""

    ctx = get_context()
    command = str(getattr(args, "command", "") or "").lower()
    command_args = list(getattr(args, "command_args", []) or [])
    telegram_arg = str(getattr(args, "telegram", "") or "").strip()

    if command == "telegram":
        handle_telegram_command(" ".join(command_args), source="cli")
        return
    if telegram_arg:
        handle_telegram_command(telegram_arg, source="cli")
        return

    if getattr(args, "doctor", False):
        run_doctor(cli_text=_cli_text)
        return

    if getattr(args, "config", False):
        configure_ai_cli()
        return

    chip_arg = str(getattr(args, "chip", "") or "").strip()
    if chip_arg:
        ctx.chip = chip_arg.upper()

    _print_startup_checks()

    task = str(getattr(args, "task", "") or "").strip()
    if task:
        if getattr(args, "connect", False):
            stm32_connect(chip_arg or None)
        agent = STM32Agent(interactive=False)
        run_oneshot(agent, task)
        stm32_disconnect()
        return

    if getattr(args, "connect", False):
        stm32_connect(chip_arg or None)

    agent = STM32Agent()
    agent.run()


__all__ = ["STM32Agent", "run"]
