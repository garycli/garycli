"""Word document tool helpers."""

from __future__ import annotations

import re


def _get_docx_module():
    """懒加载 python-docx，未安装时返回 None。"""

    try:
        import docx
    except ImportError:
        return None
    return docx


def read_docx(file_path: str) -> dict:
    """读取 Word 文档(.docx)的文本内容。"""

    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return {"success": True, "content": text, "total_paragraphs": len(doc.paragraphs)}
    except Exception as exc:
        return {"error": str(exc)}


def replace_docx_text(file_path: str, old_text: str, new_text: str, use_regex: bool = False) -> dict:
    """替换 Word 文档中的文本（支持正则）。"""

    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        count = 0
        for paragraph in doc.paragraphs:
            if use_regex:
                if not re.search(old_text, paragraph.text):
                    continue
                replaced = re.sub(old_text, new_text, paragraph.text)
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = replaced
                else:
                    paragraph.add_run(replaced)
                count += 1
                continue

            if old_text not in paragraph.text:
                continue
            replaced_in_run = False
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
                    replaced_in_run = True
            if not replaced_in_run:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                count += 1
        doc.save(file_path)
        return {"success": True, "replaced_count": count, "message": f"已替换 {count} 处"}
    except Exception as exc:
        return {"error": str(exc)}


def append_docx_content(
    file_path: str,
    content: str,
    after_paragraph_index: int = None,
    style: str = None,
) -> dict:
    """向 Word 文档追加内容（支持指定位置插入）。"""

    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        paragraphs_text = [text for text in content.split("\n") if text.strip()]
        if after_paragraph_index is None:
            for paragraph_text in paragraphs_text:
                paragraph = doc.add_paragraph(paragraph_text)
                if style:
                    try:
                        paragraph.style = style
                    except Exception:
                        pass
        else:
            total = len(doc.paragraphs)
            if after_paragraph_index < 0 or after_paragraph_index >= total:
                return {"error": f"索引 {after_paragraph_index} 超出范围 (0-{total - 1})"}
            if after_paragraph_index == total - 1:
                for paragraph_text in paragraphs_text:
                    paragraph = doc.add_paragraph(paragraph_text)
                    if style:
                        try:
                            paragraph.style = style
                        except Exception:
                            pass
            else:
                next_para = doc.paragraphs[after_paragraph_index + 1]
                base_style = doc.paragraphs[after_paragraph_index].style
                for paragraph_text in paragraphs_text:
                    new_paragraph = next_para.insert_paragraph_before(paragraph_text)
                    if style:
                        try:
                            new_paragraph.style = style
                        except Exception:
                            pass
                    else:
                        new_paragraph.style = base_style
        doc.save(file_path)
        return {"success": True, "message": "内容已追加"}
    except Exception as exc:
        return {"error": str(exc)}


def inspect_docx_structure(file_path: str, max_paragraphs: int = 50) -> dict:
    """查看 Word 文档段落结构（用于定位插入点）。"""

    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        structure = []
        for index, paragraph in enumerate(doc.paragraphs[:max_paragraphs]):
            preview = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
            if not preview.strip():
                preview = "[空段落]"
            structure.append(f"[{index}] {preview}")
        return {
            "success": True,
            "total_paragraphs": len(doc.paragraphs),
            "structure": "\n".join(structure),
        }
    except Exception as exc:
        return {"error": str(exc)}


def insert_docx_content_after_heading(
    file_path: str,
    heading_text: str,
    content: str,
    style: str = None,
) -> dict:
    """在 Word 文档指定标题后插入内容（大小写不敏感）。"""

    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if heading_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                break
        if not target_paragraph:
            return {"error": f"未找到标题: {heading_text}"}
        index = doc.paragraphs.index(target_paragraph)
        return append_docx_content(file_path, content, index, style)
    except Exception as exc:
        return {"error": str(exc)}


__all__ = [
    "append_docx_content",
    "inspect_docx_structure",
    "insert_docx_content_after_heading",
    "read_docx",
    "replace_docx_text",
]
